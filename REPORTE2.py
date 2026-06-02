import os
import socket
import subprocess
import json
import re
import shutil
import threading
import tempfile
import time
import tkinter as tk
from tkinter import simpledialog, messagebox, ttk, filedialog
from datetime import datetime
import html # Para escapar strings en el HTML

# Importar librería para pandas
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Importar librería para crear Word (docx)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ================== CONFIGURACIÓN ==================
CARPETA_DESTINO = r"\\192.168.1.68\transformadores"
# ===================================================

PS_SCRIPT = """
$ErrorActionPreference = 'SilentlyContinue'
$bios   = Get-CimInstance Win32_Bios
$sys    = Get-CimInstance Win32_ComputerSystem
$cpu    = Get-CimInstance Win32_Processor
$svc    = Get-CimInstance SoftwareLicensingService
$winKey = if ($svc.OA3xOriginalProductKey) { $svc.OA3xOriginalProductKey } else { "No encontrada en BIOS" }

$ramInfo = @(Get-CimInstance Win32_PhysicalMemory | ForEach-Object {
    $speed = if ($_.Speed) { "$($_.Speed)" } else { "Desconocida" }
    $Partnumber = if ($_.PartNumber) { $_.PartNumber.Trim() } else { "Desconocida" }
    $Manufacturer = if ($_.Manufacturer) { $_.Manufacturer.Trim() } else { "Desconocida" }
    $typeNum = if ($_.SMBIOSMemoryType) { $_.SMBIOSMemoryType } else { $_.MemoryType }
    $typeStr = switch ($typeNum) { 24 {"PC3"} 26 {"PC4"} 30 {"LPDDR4"} 34 {"PC5"} 35 {"LPDDR5"} default {"PC4"} }
    "{0}GB | {1} | {2} | {3} | {4}" -f ([math]::Round($_.Capacity/1GB)), $speed, $typeStr, $Manufacturer, $Partnumber
})

$disksData = @(Get-PhysicalDisk | Where-Object { $_.BusType -notin @('USB','File Backed Virtual') } | ForEach-Object {
    $size  = [math]::Round($_.Size / 1000000000)
    $media = if ([string]::IsNullOrWhiteSpace($_.MediaType) -or $_.MediaType -eq 'Unspecified') { 'SSD' } else { [string]$_.MediaType }
    $bus   = if ([string]$_.BusType -match 'NVMe') { 'M.2' } else { [string]$_.BusType }
    @{ desc = "{0} {1}GB {2} {3}" -f $_.FriendlyName.Trim(), $size, $media, $bus
       serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
})
if ($disksData.Count -eq 0) {
    $disksData = @(Get-CimInstance Win32_DiskDrive | Where-Object { $_.InterfaceType -ne 'USB' } | ForEach-Object {
        @{ desc = "{0} {1}GB SSD" -f $_.Model.Trim(), ([math]::Round($_.Size/1000000000))
           serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
    })
}

$netInfo = @(Get-CimInstance Win32_NetworkAdapterConfiguration | Where-Object { $_.IPEnabled } | ForEach-Object {
    "{0} | IP: {1} | MAC: {2}" -f $_.Description, $_.IPAddress[0], $_.MACAddress
})

@{ serial=($bios.SerialNumber); model=($sys.Model); win_key=($winKey)
   cpu=($cpu.Name); ram_rows=$ramInfo; disks_data=$disksData; net_rows=$netInfo
} | ConvertTo-Json -Compress -Depth 3
"""

# ─────────────────────────────────────────────────────────
#  HTML / CSS / JS  (plantilla del reporte)
# ─────────────────────────────────────────────────────────
HTML_START = r"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Inventario Maestro</title>
<style>
  *{box-sizing:border-box;margin:0;padding:0}
  body{font-family:'Segoe UI',Tahoma,sans-serif;background:#eef2f7;color:#1e293b;padding:30px}
  .container{max-width:1050px;margin:auto}
  .header{text-align:center;margin-bottom:28px}
  h1{font-size:26px;color:#1a3c6e;margin-bottom:6px}
  .subtitle{color:#64748b;font-size:13px;margin-bottom:14px}
  .counter-box{background:#dbeafe;color:#1d4ed8;padding:5px 18px;border-radius:20px;
    font-weight:700;font-size:13px;display:inline-block;border:1px solid #bfdbfe;margin-bottom:12px}

  /* Nuevos estilos para las píldoras de estadísticas */
  .stats-container { display:flex; justify-content:center; gap:10px; margin-bottom:22px; flex-wrap:wrap; }
  .stat-pill { padding:5px 16px; border-radius:20px; font-weight:700; font-size:12px; border:1px solid; }
  .stat-ram { background:#fce7f3; color:#be185d; border-color:#fbcfe8; }
  .stat-cpu { background:#fef3c7; color:#b45309; border-color:#fde68a; }
  .stat-disk { background:#dcfce7; color:#15803d; border-color:#bbf7d0; }

  .button-group{display:flex;justify-content:center;gap:12px;margin-bottom:32px;flex-wrap:wrap}
  .btn{color:#fff;border:none;padding:10px 22px;font-size:13px;font-weight:700;border-radius:8px;
    cursor:pointer;box-shadow:0 2px 6px rgba(0,0,0,.18);transition:transform .1s,box-shadow .1s}
  .btn:hover{transform:translateY(-1px);box-shadow:0 4px 12px rgba(0,0,0,.22)}
  .btn:active{transform:translateY(0)}
  .btn-copy{background:linear-gradient(135deg,#0078d4,#005a9e)}
  .btn-export{background:linear-gradient(135deg,#16a34a,#15803d)}
  .btn-delete{background:#ef4444;padding:5px 12px;font-size:11px;border-radius:6px}
  .btn-delete:hover{background:#dc2626}

  details{background:#fff;margin-bottom:12px;border-radius:12px;
    box-shadow:0 2px 8px rgba(0,0,0,.08);overflow:hidden;border:1px solid #e2e8f0}
  summary{padding:14px 18px;font-weight:700;cursor:pointer;
    background:linear-gradient(135deg,#1a3c6e,#0f2d5a);color:#fff;
    list-style:none;display:flex;justify-content:space-between;align-items:center;gap:10px}
  summary:hover{background:linear-gradient(135deg,#1e4a84,#142f60)}
  summary::-webkit-details-marker{display:none}

  .tag-entrega{background:#22c55e;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-retiro{background:#f97316;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-falla{background:#ef4444;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700;margin-left:8px}
  .logistics-row{background:#f8fafc}

  .content{padding:20px 22px;border-top:1px solid #f1f5f9}
  .content-header{display:flex;justify-content:flex-end;margin-bottom:10px; gap: 8px;}
  table{border-collapse:collapse;width:100%}
  th,td{border:1px solid #e2e8f0;padding:9px 12px;text-align:left;font-size:13px}
  th{background:#f8fafc;width:28%;color:#475569;font-weight:600}
  tr:hover td,tr:hover th{background:#f0f9ff}
  .serial-tag{color:#dc2626;font-weight:700;font-family:Consolas,monospace}
  .net-tag{background:#dcfce7;color:#166534;padding:2px 7px;border-radius:4px;font-weight:600;font-size:12px}
  .win-key{color:#1d4ed8;font-family:Consolas,monospace;font-weight:700;font-size:12px}
  .obs-row{background:#fef9c3!important}
  .obs-row td,.obs-row th{color:#854d0e;font-weight:600;font-style:italic}
  .falla-row{background:#fee2e2!important}
  .falla-row td,.falla-row th{color:#991b1b;font-weight:600}
  .office-row{background:#eff6ff}
  .office-row td,.office-row th{color:#1e40af}
</style>
<script>
// --- Lógica de Estadísticas Dinámicas ---
function actualizarEstadisticas() {
  let eqs = document.querySelectorAll('details');
  document.getElementById('total-count').innerText = eqs.length;

  let cpus = {}, rams = {}, discos = {};

  eqs.forEach(eq => {
    let d = extraerDatos(eq);

    // Agrupar CPU
    let cpuU = d.cpu.toUpperCase();
    let cpuKey = "Otro";
    if(cpuU.includes("I3")) cpuKey = "i3";
    else if(cpuU.includes("I5")) cpuKey = "i5";
    else if(cpuU.includes("I7")) cpuKey = "i7";
    else if(cpuU.includes("I9")) cpuKey = "i9";
    else if(cpuU.includes("RYZEN 3")) cpuKey = "Ryzen 3";
    else if(cpuU.includes("RYZEN 5")) cpuKey = "Ryzen 5";
    else if(cpuU.includes("RYZEN 7")) cpuKey = "Ryzen 7";
    cpus[cpuKey] = (cpus[cpuKey] || 0) + 1;

    // Agrupar RAM
    let ramMatch = d.specs.match(/^(\d+GB)/);
    let ramKey = ramMatch ? ramMatch[1] : "Otra";
    rams[ramKey] = (rams[ramKey] || 0) + 1;

    // Agrupar Discos
    d.discos.forEach(dk => {
      let desc = dk.desc.toUpperCase();
      let tipo = "Desconocido";
      if(desc.includes("M.2") || desc.includes("NVME")) tipo = "M.2 NVMe";
      else if(desc.includes("SSD")) tipo = "SSD 2.5\"";
      else if(desc.includes("HDD")) tipo = "HDD";
      discos[tipo] = (discos[tipo] || 0) + 1;
    });
  });

  let mkHTML = (obj, clase, prefijo) => {
    let arr = Object.entries(obj).map(([k,v]) => `${k} (${v})`);
    return arr.length ? `<span class="stat-pill ${clase}">${prefijo}: ${arr.join(' | ')}</span>` : '';
  };

  let statsDiv = document.getElementById('stats-container');
  if(statsDiv) {
    statsDiv.innerHTML = mkHTML(rams, 'stat-ram', 'RAM') +
                         mkHTML(cpus, 'stat-cpu', 'CPU') +
                         mkHTML(discos, 'stat-disk', 'Discos');
  }
}

document.addEventListener("DOMContentLoaded", actualizarEstadisticas);

function eliminarEquipo(id){
  if(confirm("¿Quitar este equipo del reporte actual?\n(No afecta los archivos ya guardados)")){
    let n=document.getElementById(id);
    if(n){
      n.remove();
      actualizarEstadisticas();
    }
  }
}

function extraerDatos(eq){
  let model  = eq.querySelector('.model-name').innerText.trim();
  let serial = eq.querySelector('.serial-tag').innerText.trim();
  let tds=eq.querySelectorAll('th,td');
  let obs="",cpu="",office="",guia="",mov="",ramSpeed="",ramTech="PC4",fallas="",Manufacturer="",Partnumber="Desconocida";
  let totalRam=0,ramArr=[],discos=[],tempDesc="";

  for(let i=0;i<tds.length;i++){
    if(tds[i].tagName !== 'TH') continue;

    let t=tds[i].innerText.trim();
    if(t==='N° de Guía')          guia   =tds[i+1].innerText.trim();
    if(t==='Tipo Movimiento')     mov    =tds[i+1].innerText.trim();
    if(t.includes('Obs. General'))obs    =tds[i+1].innerText.trim();
    if(t==='Fallas Detectadas')   fallas =tds[i+1].innerText.trim();
    if(t==='Procesador')          cpu    =tds[i+1].innerText.trim();
    if(t.includes('Office'))      office =tds[i+1].innerText.trim();
    if(t==='Discos Internos')     tempDesc=tds[i+1].innerText.trim();
    if(t==='Serie Disco Duro')    discos.push({desc:tempDesc,serial:tds[i+1].innerText.trim()});
    if(t==='Módulo RAM'){
      let p=tds[i+1].innerText.split('|');
      let n=parseInt(p[0].replace(/\D/g,''));
      if(!isNaN(n)){totalRam+=n;ramArr.push(n)}
      if(p[1]&&ramSpeed==="")ramSpeed=p[1].trim();
      if(p[2]&&ramTech==="PC4")ramTech=p[2].trim();
      if(p[3]&&Manufacturer==="")Manufacturer=p[3].trim()
      if(p[4]&&Partnumber==="Desconocida")Partnumber=p[4].trim()
    }
  }

  let ramFinal="RAM Desconocida";
  if(totalRam>0){
    let det=ramArr.length>1?(ramArr.every(v=>v===ramArr[0])?`(${ramArr[0]}X${ramArr.length})`:`(${ramArr.join('+')})`):"";
    let ts=ramSpeed!=="Desconocida"?`${ramTech}-${ramSpeed}Mhz`:ramTech;
    ramFinal=`${totalRam}GB${det} (${ts})`;
  }
  let diskSimple=discos.map(d=>{let m=d.desc.match(/(\d+)\s*GB\s*(.*)/i);return m?m[1]+" GB "+m[2]:d.desc});
  let specs=ramFinal+(diskSimple.length>0?", "+diskSimple.join(" + "):"");
  if(obs&&obs!=="Sin observaciones")specs+=", OBS: "+obs;
  if(fallas&&fallas!=="Ninguna")specs+=", FALLAS: "+fallas;

  return {model,serial,cpu,obs,office,guia,mov,discos,specs,equipoFull:model+" "+cpu, fallas, Manufacturer, Partnumber};
}

function copiarFilas(){
  let txt="";
  document.querySelectorAll('details').forEach(eq=>{
    let d=extraerDatos(eq);
    txt+=`${d.equipoFull}\t${d.specs}\t${d.serial}\n`;
    d.discos.forEach(dk=>{
      let lbl=dk.desc.toUpperCase().includes("HDD")?"HDD":"SSD";
      txt+=`${lbl}\t${dk.desc}\t${dk.serial}\n`;
    });
    if(d.office&&d.office.includes('Key:')){
      let ver=(d.office.match(/Office\s+([A-Za-z0-9]+)/)||[])[1]||"2016";
      let key=(d.office.match(/Key:\s*([^)]+)/)||[])[1]||"";
      txt+=`OFFICE ${ver}\tHOME AND BUSINESS\t${key.trim()}\n`;
    }
  });

  // API Moderna del portapapeles con fallback
  if (navigator.clipboard && window.isSecureContext) {
      navigator.clipboard.writeText(txt).then(() => {
          alert("✅ ¡Filas copiadas!\n\nPega directamente en tu Excel / Valida.");
      }).catch(err => {
          alert("Error al copiar al portapapeles.");
      });
  } else {
      let ta=document.createElement("textarea");ta.value=txt;document.body.appendChild(ta);ta.select();
      try{document.execCommand('copy');alert("✅ ¡Filas copiadas!\n\nPega directamente en tu Excel / Valida.");}
      catch(e){alert("Error al copiar.")}
      document.body.removeChild(ta);
  }
}

function exportarValidaCSV(){
  let csv="\uFEFFTIPO;GUIA;NOMBRE DEL EQUIPO;DESCRIPCION;NUMERO DE SERIE;OBSERVACIONES\n";
  document.querySelectorAll('details').forEach(eq=>{
    let d=extraerDatos(eq);
    csv+=`"${d.mov}";"${d.guia}";"${d.equipoFull}";"${d.specs}";"${d.serial}";"${d.obs}"\n`;
    d.discos.forEach(dk=>{
      let lbl=dk.desc.toUpperCase().includes("HDD")?"HDD":"SSD";
      csv+=`"${d.mov}";"${d.guia}";"${lbl}";"${dk.desc}";"${dk.serial}";""\n`;
    });
    if(d.office&&d.office.includes('Key:')){
      let ver=(d.office.match(/Office\s+([A-Za-z0-9]+)/)||[])[1]||"2016";
      let key=(d.office.match(/Key:\s*([^)]+)/)||[])[1]||"";
      csv+=`"${d.mov}";"${d.guia}";"OFFICE ${ver}";"HOME AND BUSINESS";"${key.trim()}";""\n`;
    }
  });
  let blob=new Blob([csv],{type:'text/csv;charset=utf-8;'}),a=document.createElement("a");
  a.href=URL.createObjectURL(blob);
  let titulo=document.querySelector('title').innerText;
  let tipo=titulo.includes("Entregas")?"Entregas":(titulo.includes("Retiros")?"Retiros":"General");
  let cliente=titulo.replace(tipo+" - ","").replace("Inventario Maestro","General").replace(/ /g,"_");
  a.download=`Reporte_${tipo}_${cliente}_${new Date().toISOString().slice(0,10).replace(/-/g,"")}.csv`;
  a.style.visibility='hidden';document.body.appendChild(a);a.click();document.body.removeChild(a);
}

/* --- LOGICA DE IMPRESIÓN CENTRALIZADA --- */

function getPrintStyle() {
  return `
    <style>
      @media print {
        @page { margin: 15mm; size: auto; }
        body { margin: 0; }
        .no-print { display: none; }
        .page-break { page-break-after: always; }
      }
      body { font-family: 'Arial', sans-serif; font-size: 11px; color: #000; max-width: 700px; margin: auto; padding: 20px; line-height: 1.4; }
      .header-title { display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 20px; font-size: 13px; }
      .row { margin-bottom: 10px; font-size: 11px;}
      .row span { margin-right: 25px; }
      table { width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 20px; }
      th, td { border: 1px solid #000; padding: 5px 8px; text-align: left; }
      .center { text-align: center; }
      .checkboxes { font-family: 'Segoe UI Symbol', sans-serif; }
      .bold { font-weight: bold; }
    </style>
  `;
}

function generarHtmlEtiqueta(d, comps) {
  let fecha = new Date().toLocaleDateString('es-CL');
  let cpuLimpia = d.cpu.replace(/Intel\\(R\\)|Core\\(TM\\)/gi, "").trim();

  let trs = comps.map(c => {
     let iconOK   = c.estado === "OK" ? "☑" : "□";
     let iconOBS  = c.estado === "OBS" ? "☑" : "□";
     let iconMalo = c.estado === "MALO" ? "☑" : "□";
     return `<tr>
       <td>${c.nombre}</td>
       <td class="center">${iconOK} OK &nbsp;&nbsp;&nbsp; ${iconOBS} OBS &nbsp;&nbsp;&nbsp; ${iconMalo} MALO</td>
       <td>${c.obs}</td>
     </tr>`;
  }).join("");

  return `
    <div class="header-title">
      <span>ETIQUETAS NOTEBOOK</span>
      <span>FECHA: ${fecha}</span>
    </div>
    <div class="row">
      <span><span class="bold">MODELO:</span> ${d.model}</span>
      <span><span class="bold">PROCESADOR:</span> ${cpuLimpia}</span>
      <span><span class="bold">GUIA:</span> ${d.guia}</span>
      <span><span class="bold">SERIAL:</span> ${d.serial}</span>
    </div>
    <div class="row checkboxes" style="margin-bottom:20px;">
      <span class="bold">ESTADO GLOBAL:</span> &nbsp;&nbsp;&nbsp;&nbsp;
      □ BUENA &nbsp;&nbsp;&nbsp; ☑ POR REPARAR &nbsp;&nbsp;&nbsp; □ MALA &nbsp;&nbsp;&nbsp; □ REPUESTO
    </div>
    <table>
      <thead>
        <tr>
          <th width="20%">REVISIÓN DE COMPONENTES</th>
          <th width="35%" class="center">ESTADO</th>
          <th width="45%">OBSERVACIÓN</th>
        </tr>
      </thead>
      <tbody class="checkboxes">
        ${trs}
      </tbody>
    </table>
    <div class="row checkboxes bold" style="margin-top:20px;">
      CHECKLIST: &nbsp;&nbsp;&nbsp;&nbsp; □ PRUEBA S/O &nbsp;&nbsp;&nbsp;&nbsp; □ HYDRA &nbsp;&nbsp;&nbsp;&nbsp; □ HARDWARE DEFAULT
    </div>
  `;
}

function imprimirEtiqueta(id) {
  let eq = document.getElementById(id);
  if(!eq) return;
  let d = extraerDatos(eq);

  let rawComps = eq.getAttribute('data-comps');
  let comps = [];
  try { comps = JSON.parse(rawComps); } catch(e) {}

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Etiqueta ${d.serial}</title>${getPrintStyle()}</head><body>${generarHtmlEtiqueta(d, comps)}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 350);
}

function imprimirTodasEtiquetas() {
  let etiquetasHtml = [];

  document.querySelectorAll('details').forEach(eq => {
    if (eq.querySelector('.tag-falla')) {
      let d = extraerDatos(eq);
      let rawComps = eq.getAttribute('data-comps');
      let comps = [];
      try { comps = JSON.parse(rawComps); } catch(e) {}

      etiquetasHtml.push(generarHtmlEtiqueta(d, comps));
    }
  });

  if (etiquetasHtml.length === 0) {
    alert("No hay equipos con fallas en este reporte para imprimir.");
    return;
  }

  let contenido = etiquetasHtml.join('<div class="page-break"></div>');

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Lote de Etiquetas</title>${getPrintStyle()}</head><body>${contenido}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 500);
}
</script>
</head>
<body>
<div class="container">
<div class="header">
  <h1>📋 Inventario Maestro de Equipos</h1>
  <p class="subtitle">Sistema de Control de Entregas y Retiros</p>
  <span class="counter-box">Total Equipos: <span id="total-count">0</span></span>

  <div id="stats-container" class="stats-container"></div>

  <div class="button-group">
    <button class="btn btn-copy"   onclick="copiarFilas()">📋 Copiar Filas (Pegar en Valida)</button>
    <button class="btn btn-export" onclick="exportarValidaCSV()">📥 Descargar CSV</button>
    <button class="btn" style="background:#334155;" onclick="imprimirTodasEtiquetas()">🖨️ Imprimir Lote de Etiquetas</button>
  </div>
</div>
<div id="main-list">"""

HTML_END = "\n</div></div></body></html>"

# ─────────────────────────────────────────────────────────
#  Helpers de formato para python-docx
# ─────────────────────────────────────────────────────────

def _set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'),  str(int(space_after  * 20)))
    if line_spacing:
        spacing.set(qn('w:line'),      str(int(line_spacing * 240)))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_paragraph_spacing(p, space_before=0, space_after=4)
    return p

def _add_campo(doc, etiqueta, valor, etiqueta_bold=True, valor_bold=False, font_size=11):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=0, space_after=2)
    run_lbl = p.add_run(etiqueta)
    run_lbl.bold = etiqueta_bold
    run_lbl.font.size = Pt(font_size)
    run_lbl.font.name = 'Arial'
    run_val = p.add_run(valor)
    run_val.bold = valor_bold
    run_val.font.size = Pt(font_size)
    run_val.font.name = 'Arial'
    return p

def _add_titulo_seccion(doc, texto, font_size=11, uppercase=True, space_before=10, space_after=4):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
    run = p.add_run(texto.upper() if uppercase else texto)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return p

def _add_falla_item(doc, texto, font_size=11):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=0, space_after=1)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    run = p.add_run(f"- {texto.upper()}")
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    return p

# ─────────────────────────────────────────────────────────
#  Generación de Word (Docx)
# ─────────────────────────────────────────────────────────

def generar_informe_word(cliente, equipos_malos, ruta_cliente):
    if not DOCX_AVAILABLE:
        messagebox.showwarning("Falta Librería", "No se pudo generar el Word porque falta 'python-docx'.")
        return

    try:
        doc = Document()

        # Configurar márgenes
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        logo_path = os.path.join(os.getcwd(), "lg1.png")

        # 1. CABECERA ÚNICA PARA TODO EL LOTE
        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            r_logo = p_logo.add_run()
            r_logo.add_picture(logo_path, width=Cm(6.5))
            _set_paragraph_spacing(p_logo, space_before=0, space_after=4)

        p_titulo = doc.add_paragraph()
        _set_paragraph_spacing(p_titulo, space_before=0, space_after=8)
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run("INFORME TÉCNICO DE RETIRO - RESUMEN DE LOTE")
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

        _add_horizontal_rule(doc)

        _add_campo(doc, "CLIENTE:             ", cliente.upper())
        _add_campo(doc, "FECHA DE REVISIÓN:   ", fecha_hoy)
        _add_campo(doc, "EQUIPOS DEFECTUOSOS: ", str(len(equipos_malos)))

        _add_horizontal_rule(doc)
        _add_titulo_seccion(doc, "DETALLE DE EQUIPOS Y FALLAS", space_before=8, space_after=8)

        # 2. CREACIÓN DE LA TABLA RESUMEN
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = 'Table Grid' # Le pone bordes a la tabla

        # Títulos de las columnas
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'GUÍA'
        hdr_cells[1].text = 'MODELO'
        hdr_cells[2].text = 'N° DE SERIE'
        hdr_cells[3].text = 'DETALLE DE FALLAS'

        # Poner los títulos de la tabla en negrita
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # 3. LLENAR LA TABLA CON LOS EQUIPOS
        for equipo in equipos_malos:
            guia   = equipo.get("GUIA_ID", "S/N")
            modelo = equipo.get("MODELO",  "Desconocido")
            serie  = equipo.get("SERIAL",  "Desconocido")
            obs    = equipo.get("OBS", "")

            detalle_comps = equipo.get("DETALLE_COMPONENTES", [])
            fallas_list = []
            for c in detalle_comps:
                if c["estado"] != "OK":
                    texto = f"- {c['nombre']}: {c['estado']}"
                    if c["obs"]: texto += f" ({c['obs']})"
                    fallas_list.append(texto)

            # Unimos todas las fallas en un solo bloque de texto
            texto_fallas = "\n".join(fallas_list)
            if obs and obs not in ("Sin observaciones", ""):
                texto_fallas += f"\nOBS GENERAL: {obs}"

            # Añadir fila a la tabla
            row_cells = tabla.add_row().cells
            row_cells[0].text = str(guia)
            row_cells[1].text = modelo
            row_cells[2].text = serie
            row_cells[3].text = texto_fallas

        # 4. CONCLUSIÓN Y FIRMA ÚNICA AL FINAL
        _add_titulo_seccion(doc, "CONCLUSIÓN:", space_before=15, space_after=4)
        p_conclu = doc.add_paragraph()
        run_c = p_conclu.add_run(
            "Los equipos listados en la tabla superior presentan fallas físicas o de hardware. "
            "Se requiere revisión detallada, cambio de componentes afectados o envío a bodega "
            "a la espera de repuestos o reparación."
        )

        for _ in range(3):
            doc.add_paragraph()

        _add_horizontal_rule(doc)
        p_firma = doc.add_paragraph()
        run_f1 = p_firma.add_run("TÉCNICO ENCARGADO: ")
        run_f1.bold = True
        run_f2 = p_firma.add_run("TOMAS GAC\n")
        run_f2.bold = True
        run_r1 = p_firma.add_run("RUT: ")
        run_r1.bold = True
        p_firma.add_run("21.790.634-2")

        # Guardar el documento
        fecha_str      = datetime.now().strftime("%Y%m%d")
        nombre_archivo = f"Informe_Tecnico_{cliente}_{fecha_str}.docx"
        ruta_guardado  = os.path.join(ruta_cliente, nombre_archivo)
        doc.save(ruta_guardado)

    except Exception as e:
        print(f"Error generando Word: {e}")
        messagebox.showerror("Error al generar Word", str(e))

def get_inventory_fast(max_retries=3, timeout_per_attempt=45):
    """
    Obtiene información de hardware via PowerShell.
    - Ejecuta el script desde archivo temporal (más confiable que -Command inline).
    - Extrae el JSON con regex para tolerar advertencias extra en la salida.
    - Reintenta automáticamente hasta max_retries veces si falla.
    """
    last_error = None

    for attempt in range(1, max_retries + 1):
        tmp_path = None
        try:
            # ── 1. Escribir el script a un .ps1 temporal ──────────────────────
            # Pasar el script inline con -Command puede fallar por parsing de args.
            # Ejecutarlo como archivo es siempre más fiable.
            with tempfile.NamedTemporaryFile(
                mode='w', suffix='.ps1', delete=False, encoding='utf-8'
            ) as tmp:
                tmp.write(PS_SCRIPT)
                tmp_path = tmp.name

            # ── 2. Configurar proceso oculto ───────────────────────────────────
            si = subprocess.STARTUPINFO()
            si.dwFlags  |= subprocess.STARTF_USESHOWWINDOW
            si.wShowWindow = 0  # SW_HIDE

            cmd = [
                "powershell",
                "-NoProfile",
                "-NonInteractive",           # evita prompts interactivos que cuelgan
                "-ExecutionPolicy", "Bypass",
                "-File", tmp_path            # ejecutar como archivo, no como -Command
            ]

            # ── 3. Ejecutar con timeout extendido ─────────────────────────────
            result = subprocess.check_output(
                cmd, text=True, startupinfo=si,
                stderr=subprocess.PIPE,      # capturar stderr para diagnóstico
                timeout=timeout_per_attempt
            )

            # ── 4. Extraer el bloque JSON de la salida ────────────────────────
            # PowerShell puede emitir advertencias o líneas extra ANTES del JSON.
            # Buscamos el primer { ... } que abarque toda la línea JSON.
            json_match = re.search(r'\{.*\}', result, flags=re.DOTALL)
            if not json_match:
                raise ValueError(
                    f"Salida de PowerShell no contiene JSON válido "
                    f"(intento {attempt}/{max_retries}). "
                    f"Salida recibida: {result[:200]!r}"
                )

            data = json.loads(json_match.group(0))

            # ── 5. Limpiar archivo temporal ────────────────────────────────────
            try:
                os.unlink(tmp_path)
                tmp_path = None
            except Exception:
                pass

            # ── 6. Procesar los datos del hardware ─────────────────────────────
            raw_model = data.get("model", "Desconocido")
        cln_model = re.sub(r'(?i)\b(11\.6|12\.5|13\.3|14|15\.6|16|17\.3)\s*(inch|")?\\b', '', raw_model)
        cln_model = re.sub(r'(?i)\bnotebook\s*pc\b', '', cln_model)
        cln_model = re.sub(r'\s+', ' ', cln_model).strip()

        ram_list = data.get("ram_rows", [])
        if isinstance(ram_list, str): ram_list = [ram_list]
        ram_html = "".join(f"<tr><th>Módulo RAM</th><td>{r}</td></tr>" for r in ram_list) \
                   or "<tr><th>RAM</th><td>No detectada</td></tr>"

        fabricantes_ram = []
        for r in ram_list:
            partes = r.split("|")
            if len(partes) >= 4:
                fabricantes_ram.append(partes[3].strip())
                # Eliminar duplicados por si tiene 2 RAMs de la misma marca
        fabricantes_unicos = list(set(fabricantes_ram))
        fabricante_final = " / ".join(fabricantes_unicos) if fabricantes_unicos else "Desconocida"

        partnumbers_ram = []
        for r in ram_list:
            partes = r.split("|")
            if len(partes) >= 5:
                partnumbers_ram.append(partes[4].strip())
        partnumber_final = " / ".join(set(partnumbers_ram)) if partnumbers_ram else "Desconocida"

        disks = data.get("disks_data", [])
        if isinstance(disks, dict): disks = [disks]
        disk_html = ""
        for d in disks:
            desc  = d.get("desc", "Desconocido")
            serie = re.sub(r"[^a-zA-Z0-9]", "", d.get("serial", "NA"))[-16:]
            disk_html += (f"<tr><th>Discos Internos</th><td>{desc}</td></tr>"
                          f"<tr><th>Serie Disco Duro</th><td class='serial-tag'>{serie}</td></tr>")

        net_list = data.get("net_rows", [])
        if isinstance(net_list, str): net_list = [net_list]
        net_html = "".join(f"<tr><td>Red Activa</td><td><span class='net-tag'>{n}</span></td></tr>" for n in net_list) \
                   or "<tr><td colspan='2'>Sin red activa</td></tr>"

            # ── 7. Retornar datos — éxito ──────────────────────────────────────
            return {
                "fecha":    datetime.now().strftime("%d/%m/%Y %H:%M"),
                "host":     socket.gethostname(),
                "model":    cln_model,
                "serial":   data.get("serial",  "SN_DESCONOCIDO").strip(),
                "key":      data.get("win_key", "N/A"),
                "cpu":      data.get("cpu",     "Desconocido"),
                "ram_html": ram_html,
                "ram_raw":  ram_list,
                "disk_html":disk_html,
                "net_rows": net_html,
                "Manufacturer": fabricante_final,
                "Partnumber": partnumber_final
            }

        except subprocess.TimeoutExpired:
            last_error = f"Timeout tras {timeout_per_attempt}s (intento {attempt}/{max_retries})"
            print(f"⚠️  {last_error}")
        except json.JSONDecodeError as e:
            last_error = f"JSON inválido (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        except Exception as e:
            last_error = f"Error inesperado (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        finally:
            # Garantizar limpieza del archivo temporal en cualquier caso
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

        # Espera breve entre reintentos para dar tiempo a WMI a recuperarse
        if attempt < max_retries:
            time.sleep(2)

    print(f"❌ Escaneo falló después de {max_retries} intentos. Último: {last_error}")
    return None

_GUIA_CONFIG = os.path.join(os.getcwd(), "Reportes_Guardados", ".guia_retiro.json")

def guardar_guia_config(guia: str):
    os.makedirs(os.path.dirname(_GUIA_CONFIG), exist_ok=True)
    with open(_GUIA_CONFIG, "w", encoding="utf-8") as f:
        json.dump({"guia": guia}, f)

def cargar_guia_config() -> str:
    try:
        with open(_GUIA_CONFIG, "r", encoding="utf-8") as f:
            return json.load(f).get("guia", "")
    except Exception:
        return ""

def limpiar_guia_config():
    try:
        if os.path.exists(_GUIA_CONFIG):
            os.remove(_GUIA_CONFIG)
    except Exception:
        pass


def guardar_equipo_general(data, obs, tiene_office, office_key, version_office, mov, guia, detalle_componentes):
    ruta_base = os.path.join(os.getcwd(), "Reportes_Guardados")
    os.makedirs(ruta_base, exist_ok=True)

    obs_final  = obs  or "Sin observaciones"
    guia_final = guia or "Sin Guía"
    tipo       = "Entregas" if mov == "Entrega" else "Retiros"
    report_file= os.path.join(ruta_base, f"Reporte_{tipo}.html")
    json_file  = os.path.join(ruta_base, f"data_{re.sub(r'[^a-zA-Z0-9]','',data['serial'])}.json")
    safe_id = f"dev-{re.sub(r'[^a-zA-Z0-9]','',data['serial'])}"

    office_html = (f'<tr class="office-row"><th>🔑 Office Instalado</th>'
                   f'<td><b>Office {version_office}</b> (Key: {office_key})</td></tr>') if tiene_office else ""

    tag_cls = "tag-entrega" if mov == "Entrega" else "tag-retiro"

    tiene_fallas = any(c["estado"] != "OK" for c in detalle_componentes)
    fallas_str = ", ".join([f"{c['nombre']} ({c['estado']})" for c in detalle_componentes if c["estado"] != "OK"])

    falla_html = ""
    tag_falla_html = ""
    btn_imprimir = ""

    if tiene_fallas:
        tag_falla_html = '<span class="tag-falla">CON FALLAS</span>'
        falla_html = f'<tr class="falla-row"><th>Fallas Detectadas</th><td>{fallas_str}</td></tr>'
        btn_imprimir = f'<button class="btn" style="background:#475569; padding:5px 12px; font-size:11px;" onclick="imprimirEtiqueta(\'{safe_id}\')">🖨️ Etiqueta Individual</button>'

    comps_json = html.escape(json.dumps(detalle_componentes))

    # --- AQUÍ ESTABA EL ERROR DE SINTAXIS (F-STRING ROTO) ---
    new_entry = f"""<details id="{safe_id}" data-comps='{comps_json}' open>
<summary>
  <span class="model-name">{data['model']}</span>
  <span style="font-size:12px;opacity:.85">S/N: {data['serial']} &nbsp;|&nbsp; 📅 {data['fecha']} {tag_falla_html}</span>
</summary>
<div class="content">
  <div class="content-header">
    {btn_imprimir}
    <button class="btn btn-delete" onclick="eliminarEquipo('{safe_id}')">🗑️ Quitar del Reporte</button>
  </div>
  <table>
    <tr class="logistics-row"><th>Tipo Movimiento</th><td><b>{mov}</b></td></tr>
    <tr class="logistics-row"><th>N° de Guía</th><td><b>{guia_final}</b></td></tr>
    {falla_html}
    <tr class="obs-row"><th>⚠️ Obs. General</th><td>{obs_final}</td></tr>
    {office_html}
    <tr><th>Número de Serie</th><td class="serial-tag">{data['serial']}</td></tr>
    <tr><th>Licencia Windows (OA3)</th><td class="win-key">{data['key']}</td></tr>
    <tr><th>Procesador</th><td>{data['cpu']}</td></tr>
    {data['ram_html']}{data['disk_html']}{data['net_rows']}
  </table>
</div></details>"""

    # --- AQUÍ HABÍA UN ERROR DE INDENTACIÓN ---
    if os.path.exists(report_file):
        content = open(report_file, "r", encoding="utf-8").read()
    else:
        content = HTML_START.replace("Inventario Maestro de Equipos", f"Reporte de {tipo}") + HTML_END

    content = re.sub(rf'<details id="{safe_id}".*?</details>', "", content, flags=re.DOTALL)
    content = content.replace('<div id="main-list">', '<div id="main-list">\n' + new_entry)

    total = len(re.findall(r"<details ", content))
    content = re.sub(r'<span id="total-count">\d+</span>', f'<span id="total-count">{total}</span>', content)

    with open(report_file, "w", encoding="utf-8") as f:
        f.write(content)

    with open(json_file, "w", encoding="utf-8") as f:
        json.dump({
            "CLIENTE": "PENDIENTE", "SERIAL": data["serial"], "MODELO": data["model"],
            "TIPO_MOVIMIENTO": mov, "GUIA_ID": guia_final, "OBS": obs_final,
            "DETALLE_COMPONENTES": detalle_componentes, "TIENE_FALLAS": tiene_fallas,
            "DATA": data, "OFFICE": f"Office {version_office}" if tiene_office else "No",
        }, f, ensure_ascii=False)


def accion_agregar_lote(ventana, data, obs, tiene_office, office_key, version_office, mov, guia, comps_data):
    detalle_componentes = []
    for c in comps_data:
        est = c["estado"].get()
        ob  = c["obs"].get().strip()
        detalle_componentes.append({"nombre": c["nombre"], "estado": est, "obs": ob})

    guardar_equipo_general(data, obs, tiene_office, office_key, version_office, mov, guia, detalle_componentes)

    if mov == "Retiro" and guia:
        guardar_guia_config(guia)

    tipo = "Entregas" if mov == "Entrega" else "Retiros"
    # --- AQUÍ HABÍA OTRO ERROR DE INDENTACIÓN ---
    messagebox.showinfo("✅ Guardado", f"Equipo agregado al archivo 'Reporte_{tipo}'.")
    mostrar_menu_principal(ventana)


def accion_finalizar_lote(ventana):
    ruta_base    = os.path.join(os.getcwd(), "Reportes_Guardados")
    html_e       = os.path.join(ruta_base, "Reporte_Entregas.html")
    html_r       = os.path.join(ruta_base, "Reporte_Retiros.html")

    if not os.path.exists(html_e) and not os.path.exists(html_r):
        messagebox.showwarning("Lote Vacío", "No hay ningún reporte pendiente.\n\nPrimero escanea y agrega equipos.")
        return

    cliente = simpledialog.askstring("Empacar Lote", "¿A qué CLIENTE pertenecen estos equipos?", parent=ventana)
    if not cliente:
        return

    cliente_clean = re.sub(r"[^a-zA-Z0-9\s_\-]", "", cliente).strip().replace(" ", "_").upper() or "GENERAL"
    ruta_cliente  = os.path.join(ruta_base, cliente_clean)
    os.makedirs(ruta_cliente, exist_ok=True)
    fecha_hoy = datetime.now().strftime("%Y%m%d")

    for archivo, tipo_doc in [(html_e, "Entregas"), (html_r, "Retiros")]:
        if not os.path.exists(archivo):
            continue
        content = open(archivo, "r", encoding="utf-8").read()
        content = content.replace(f"Reporte de {tipo_doc}", f"{tipo_doc} - {cliente_clean}")
        content = content.replace(f"<title>Reporte de {tipo_doc}</title>", f"<title>{tipo_doc} - {cliente_clean}</title>")
        nuevo   = os.path.join(ruta_cliente, f"Reporte_{tipo_doc}_{cliente_clean}_{fecha_hoy}.html")
        with open(nuevo, "w", encoding="utf-8") as f:
            f.write(content)
        os.remove(archivo)

    equipos_malos = []
    for archivo in os.listdir(ruta_base):
        if not (archivo.endswith(".json") and archivo.startswith("data_")):
            continue
        ruta_orig = os.path.join(ruta_base, archivo)
        serial_part = archivo[5:-5]
        nuevo_json  = os.path.join(ruta_cliente, f"Reporte_{cliente_clean}_{fecha_hoy}_{serial_part}.json")
        try:
            jdata = json.load(open(ruta_orig, "r", encoding="utf-8"))
            jdata["CLIENTE"] = cliente_clean
            if jdata.get("TIENE_FALLAS", False):
                equipos_malos.append(jdata)
            json.dump(jdata, open(nuevo_json, "w", encoding="utf-8"), ensure_ascii=False)
            os.remove(ruta_orig)
        except Exception:
            shutil.move(ruta_orig, nuevo_json)

    if equipos_malos:
        generar_informe_word(cliente_clean, equipos_malos, ruta_cliente)
        aviso_extra = f"\n\n📄 Se generó el 'Informe_Tecnico_{cliente_clean}.docx' con {len(equipos_malos)} equipo(s) defectuoso(s)."
    else:
        aviso_extra = ""

    messagebox.showinfo("📦 Lote Finalizado", f"¡Lote empacado con éxito!\n\nCarpeta: {cliente_clean}{aviso_extra}")
    limpiar_guia_config()


def accion_enviar_red(ventana):
    carpeta_local = os.path.join(os.getcwd(), "Reportes_Guardados")
    if not os.path.exists(carpeta_local):
        messagebox.showwarning("Aviso", "No hay reportes locales para enviar.")
        return

    carpetas = [d for d in os.listdir(carpeta_local) if os.path.isdir(os.path.join(carpeta_local, d))]
    if not carpetas:
        pendiente = any(
            os.path.exists(os.path.join(carpeta_local, f))
            for f in ("Reporte_Entregas.html", "Reporte_Retiros.html")
        )
        if pendiente:
            messagebox.showwarning("Lote Incompleto", "Hay equipos sin empacar.\n\nPresiona '📦 EMPACAR LOTE' antes de enviar.")
        else:
            messagebox.showwarning("Aviso", "No hay carpetas de clientes listas para enviar.")
        return

    try:
        os.makedirs(CARPETA_DESTINO, exist_ok=True)
    except Exception as e:
        messagebox.showerror("Error de Red", f"No hay conexión a:\n{CARPETA_DESTINO}\n\nError: {e}")
        return

    try:
        count = 0
        for cliente_dir in carpetas:
            src = os.path.join(carpeta_local, cliente_dir)
            dst = os.path.join(CARPETA_DESTINO, cliente_dir)
            os.makedirs(dst, exist_ok=True)
            for archivo in os.listdir(src):
                ruta_arch = os.path.join(src, archivo)
                if os.path.isfile(ruta_arch):
                    shutil.copy2(ruta_arch, dst)
            count += 1
        messagebox.showinfo("🚀 Enviado", f"¡Sincronización Exitosa!\n\n{count} carpeta(s) transferidas a:\n{CARPETA_DESTINO}")
    except Exception as e:
        messagebox.showerror("Error al Copiar", f"No se pudo copiar a la red:\n{e}")

# ─────────────────────────────────────────────────────────
#  Módulos de Auditoría
# ─────────────────────────────────────────────────────────
def abrir_modulo_auditoria(ventana_padre):
    win = tk.Toplevel(ventana_padre)
    win.title("Auditoría — Entregas vs Retiros")
    win.geometry("860x520")
    win.configure(bg="#eef2f7")
    win.grab_set()

    tk.Label(win, text="🔍 Auditoría Diferencial (Entregas vs Retiros)", font=("Segoe UI", 14, "bold"), bg="#eef2f7", fg="#1a3c6e").pack(pady=12)

    frame_rutas = tk.Frame(win, bg="#eef2f7")
    frame_rutas.pack(fill="x", padx=20, pady=4)
    ruta_entrega = tk.StringVar()
    ruta_retiro  = tk.StringVar()

    def sel_carpeta(var, titulo):
        c = filedialog.askdirectory(title=titulo, parent=win)
        if c: var.set(c)

    style_btn = dict(bg="#1a3c6e", fg="white", font=("Segoe UI", 9), cursor="hand2", relief="flat")
    tk.Button(frame_rutas, text="📂 Carpeta Entregas", command=lambda: sel_carpeta(ruta_entrega, "Carpeta de ENTREGAS"), **style_btn).grid(row=0, column=0, padx=5, pady=3, sticky="ew")
    tk.Entry(frame_rutas, textvariable=ruta_entrega, width=55, state="readonly", font=("Segoe UI", 9)).grid(row=0, column=1, padx=5)
    tk.Button(frame_rutas, text="📂 Carpeta Retiros", command=lambda: sel_carpeta(ruta_retiro, "Carpeta de RETIROS"), **style_btn).grid(row=1, column=0, padx=5, pady=3, sticky="ew")
    tk.Entry(frame_rutas, textvariable=ruta_retiro, width=55, state="readonly", font=("Segoe UI", 9)).grid(row=1, column=1, padx=5)

    style = ttk.Style()
    style.configure("Audit.Treeview", rowheight=24, font=("Segoe UI", 9))
    style.configure("Audit.Treeview.Heading", font=("Segoe UI", 9, "bold"))

    cols = ("estado", "serial", "modelo", "detalle")
    tree = ttk.Treeview(win, columns=cols, show="headings", height=13, style="Audit.Treeview")
    for col, w, txt in [("estado",105,"ESTADO"),("serial",130,"S/N"),("modelo",190,"MODELO"),("detalle",400,"DETALLE")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="center" if col in ("estado","serial") else "w")

    tree.tag_configure("ok",      foreground="#16a34a")
    tree.tag_configure("falta",   foreground="#dc2626", background="#fee2e2")
    tree.tag_configure("sobra",   foreground="#d97706", background="#fef3c7")
    tree.tag_configure("alterado",foreground="#b91c1c", background="#fecaca", font=("Segoe UI", 9, "bold"))

    sb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=sb.set)
    frame_tree = tk.Frame(win, bg="#eef2f7")
    frame_tree.pack(fill="both", expand=True, padx=20, pady=8)
    tree.pack(side="left", fill="both", expand=True, in_=frame_tree)
    sb.pack(side="right",  fill="y",   in_=frame_tree)

    def cargar_jsons(ruta):
        datos = {}
        for arch in os.listdir(ruta):
            if arch.endswith(".json"):
                try:
                    j = json.load(open(os.path.join(ruta, arch), "r", encoding="utf-8"))
                    if "SERIAL" in j: datos[j["SERIAL"]] = j
                except Exception: pass
        return datos

    def ejecutar():
        tree.delete(*tree.get_children())
        if not ruta_entrega.get() or not ruta_retiro.get(): return messagebox.showwarning("Faltan Datos", "Selecciona ambas carpetas.", parent=win)
        entregas = cargar_jsons(ruta_entrega.get())
        retiros  = cargar_jsons(ruta_retiro.get())
        if not entregas: return messagebox.showerror("Error", "Carpeta de Entregas sin JSONs válidos.", parent=win)
        if not retiros:  return messagebox.showerror("Error", "Carpeta de Retiros sin JSONs válidos.",  parent=win)

        for serial, d_e in entregas.items():
            modelo = d_e.get("MODELO", "Desconocido")
            if serial not in retiros:
                tree.insert("", tk.END, values=("❌ FALTANTE", serial, modelo, "Equipo de entrega no está en retiro."), tags=("falta",))
                continue
            d_r = retiros[serial]
            alertas = []
            if len(d_e["DATA"].get("ram_raw", [])) != len(d_r["DATA"].get("ram_raw", [])):
                alertas.append(f"RAM alterada.")
            s_e = {d.get("serial","NA") for d in d_e["DATA"].get("disks_data",[]) if isinstance(d, dict)}
            s_r = {d.get("serial","NA") for d in d_r["DATA"].get("disks_data",[]) if isinstance(d, dict)}
            if s_e != s_r: alertas.append("DISCO serial distinto.")
            tag = "alterado" if alertas else "ok"
            tree.insert("", tk.END, values=("⚠️ ALTERADO" if alertas else "✅ OK", serial, modelo, " | ".join(alertas) if alertas else "Hardware coincide."), tags=(tag,))
        for serial, d_r in retiros.items():
            if serial not in entregas:
                tree.insert("", tk.END, values=("❓ SOBRANTE", serial, d_r.get("MODELO","Desconocido"), "Equipo en retiro no figura en entregas."), tags=("sobra",))
    tk.Button(win, text="⚙️  INICIAR AUDITORÍA", command=ejecutar, bg="#16a34a", fg="white", font=("Segoe UI", 10, "bold"), padx=18, pady=7, cursor="hand2", relief="flat").pack(pady=10)

def abrir_modulo_auditoria_mixta(ventana_padre):
    win = tk.Toplevel(ventana_padre)
    win.title("Auditoría — Sistema vs Físico")
    win.geometry("900x550")
    win.configure(bg="#eef2f7")
    win.grab_set()
    tk.Label(win, text="📊 Auditoría Cruzada (Teórico vs Físico)", font=("Segoe UI", 14, "bold"), bg="#eef2f7", fg="#1a3c6e").pack(pady=12)

    frame_rutas = tk.Frame(win, bg="#eef2f7")
    frame_rutas.pack(fill="x", padx=20, pady=4)
    ruta_base = tk.StringVar()
    ruta_fisico = tk.StringVar()
    def sel_archivo(var, titulo, tipos):
        arch = filedialog.askopenfilename(title=titulo, parent=win, filetypes=tipos)
        if arch: var.set(arch)

    style_btn = dict(bg="#1a3c6e", fg="white", font=("Segoe UI", 9), cursor="hand2", relief="flat")
    tk.Button(frame_rutas, text="📂 Cargar Excel Base", command=lambda: sel_archivo(ruta_base, "Seleccionar Excel Base", [("Archivos Excel", "*.xlsx *.xls")]), **style_btn).grid(row=0, column=0, padx=5, pady=3, sticky="ew")
    tk.Entry(frame_rutas, textvariable=ruta_base, width=55, state="readonly", font=("Segoe UI", 9)).grid(row=0, column=1, padx=5)
    tk.Button(frame_rutas, text="📂 Cargar Físico", command=lambda: sel_archivo(ruta_fisico, "Seleccionar Reporte Físico", [("HTML o Excel", "*.html *.xlsx *.xls")]), **style_btn).grid(row=1, column=0, padx=5, pady=3, sticky="ew")
    tk.Entry(frame_rutas, textvariable=ruta_fisico, width=55, state="readonly", font=("Segoe UI", 9)).grid(row=1, column=1, padx=5)

    style = ttk.Style()
    style.configure("Audit.Treeview", rowheight=24, font=("Segoe UI", 9))
    style.configure("Audit.Treeview.Heading", font=("Segoe UI", 9, "bold"))
    cols = ("estado", "serial", "modelo", "origen")
    tree = ttk.Treeview(win, columns=cols, show="headings", height=12, style="Audit.Treeview")
    for col, w, txt in [("estado",110,"ESTADO"),("serial",150,"N° SERIE"),("modelo",250,"MODELO / INFO"),("origen",280,"DETALLE")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="center" if col in ("estado","serial") else "w")
    tree.tag_configure("ok", foreground="#16a34a")
    tree.tag_configure("falta", foreground="#dc2626", background="#fee2e2")
    tree.tag_configure("sobra", foreground="#d97706", background="#fef3c7")

    sb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=sb.set)
    frame_tree = tk.Frame(win, bg="#eef2f7")
    frame_tree.pack(fill="both", expand=True, padx=20, pady=8)
    tree.pack(side="left", fill="both", expand=True, in_=frame_tree)
    sb.pack(side="right", fill="y",   in_=frame_tree)

    def extraer_excel(ruta):
        equipos = {}
        try:
            df = pd.read_excel(ruta)
            df.columns = df.columns.str.strip().str.upper()
            col_serie = None
            for col in df.columns:
                if col in ["NUMERO DE SERIE", "NÚMERO DE SERIE", "SERIAL", "S/N", "SN", "SERIE"]: col_serie = col; break
            if not col_serie:
                for col in df.columns:
                    if "SERIE" in str(col) and "DISCO" not in str(col): col_serie = col; break
            if col_serie:
                col_modelo = next((c for c in df.columns if "MODEL" in str(c) or "EQUIPO" in str(c)), None)
                for _, row in df.iterrows():
                    sn = str(row[col_serie]).strip()
                    if sn and sn.lower() != "nan": equipos[sn.upper()] = str(row[col_modelo]).strip() if col_modelo else "Excel"
        except Exception as e: messagebox.showerror("Error Excel", f"No se pudo leer {ruta}\n{e}", parent=win)
        return equipos

    def extraer_html(ruta):
        equipos = {}
        try:
            with open(ruta, "r", encoding="utf-8") as f: content = f.read()
            bloques = re.findall(r'<details.*?</details>', content, flags=re.DOTALL)
            for bloque in bloques:
                s_match = re.search(r'<tr><th>Número de Serie</th><td class="serial-tag">(.*?)</td></tr>', bloque)
                m_match = re.search(r'<span class="model-name">(.*?)</span>', bloque)
                if s_match: equipos[s_match.group(1).strip().upper()] = m_match.group(1).strip() if m_match else "HTML"
        except Exception as e: messagebox.showerror("Error HTML", f"No se pudo leer {ruta}\n{e}", parent=win)
        return equipos

    def ejecutar():
        tree.delete(*tree.get_children())
        r_base, r_fis = ruta_base.get(), ruta_fisico.get()
        if not r_base or not r_fis: return messagebox.showwarning("Faltan Datos", "Selecciona ambos archivos.", parent=win)
        dict_base = extraer_excel(r_base)
        if not dict_base: return
        dict_fis = extraer_html(r_fis) if r_fis.endswith('.html') else extraer_excel(r_fis)
        if not dict_fis: return

        for sn, mod in dict_base.items():
            if sn in dict_fis: tree.insert("", tk.END, values=("✅ MATCH", sn, mod, "Encontrado en ambos."), tags=("ok",))
            else: tree.insert("", tk.END, values=("❌ FALTANTE", sn, mod, "Está en la Base, pero NO en el Físico."), tags=("falta",))
        for sn, mod in dict_fis.items():
            if sn not in dict_base: tree.insert("", tk.END, values=("❓ SOBRANTE", sn, mod, "Físico no figura en el Excel Base."), tags=("sobra",))
    tk.Button(win, text="⚙️  CRUZAR INVENTARIOS", command=ejecutar, bg="#16a34a", fg="white", font=("Segoe UI", 10, "bold"), padx=18, pady=7, cursor="hand2", relief="flat").pack(pady=10)

# ─────────────────────────────────────────────────────────
#  Navegación e Interfaz Principal
# ─────────────────────────────────────────────────────────
def limpiar_ventana(ventana):
    for widget in ventana.winfo_children():
        widget.destroy()

def mostrar_menu_principal(ventana):
    limpiar_ventana(ventana)
    ventana.geometry("520x480")

    tk.Label(ventana, text="🖥️ Sistema de Control de Inventario", font=("Segoe UI", 16, "bold"), bg="#eef2f7", fg="#1a3c6e").pack(pady=(30, 20))
    frame_btns = tk.Frame(ventana, bg="#eef2f7")
    frame_btns.pack(fill="both", expand=True, padx=40)

    def hacer_btn_menu(texto, comando, color, size=11, bold="bold"):
        tk.Button(frame_btns, text=texto, command=comando, bg=color, fg="white", font=("Segoe UI", size, bold), pady=10, cursor="hand2", relief="flat").pack(fill="x", pady=6)

    hacer_btn_menu("🔍 ESCANEAR ESTE EQUIPO\n(Crear registro)", lambda: iniciar_escaneo(ventana), "#0078d4", size=12)
    hacer_btn_menu("📦 EMPACAR LOTE PENDIENTE", lambda: accion_finalizar_lote(ventana), "#16a34a")
    hacer_btn_menu("🚀 ENVIAR CARPETAS A RED", lambda: accion_enviar_red(ventana), "#d97706")

    ttk.Separator(frame_btns, orient='horizontal').pack(fill='x', pady=15)
    tk.Label(frame_btns, text="Herramientas de Auditoría", font=("Segoe UI", 10, "bold"), bg="#eef2f7", fg="#64748b").pack(pady=(0, 5))
    hacer_btn_menu("🔍 AUDITORÍA (Entregas vs Retiros)", lambda: abrir_modulo_auditoria(ventana), "#334155", size=9, bold="normal")
    hacer_btn_menu("📊 AUDITORÍA MIXTA (Excel vs Físico)", lambda: abrir_modulo_auditoria_mixta(ventana), "#0f2d5a", size=9, bold="normal")

def iniciar_escaneo(ventana):
    limpiar_ventana(ventana)
    ventana.geometry("750x850")

    frame_carga = tk.Frame(ventana, bg="#eef2f7")
    frame_carga.place(relx=.5, rely=.5, anchor="center")

    tk.Label(frame_carga, text="⏳", font=("Segoe UI", 48), bg="#eef2f7").pack()
    lbl_scan = tk.Label(frame_carga, text="Escaneando hardware...", font=("Segoe UI", 15, "bold"), bg="#eef2f7", fg="#1a3c6e")
    lbl_scan.pack(pady=6)
    tk.Label(frame_carga, text="Esto tomará unos segundos.", font=("Segoe UI", 10), bg="#eef2f7", fg="#64748b").pack()

    dots = ["", ".", "..", "..."]
    dot_idx = [0]
    def animar():
        if frame_carga.winfo_exists():
            lbl_scan.config(text=f"Escaneando hardware{dots[dot_idx[0] % 4]}")
            dot_idx[0] += 1
            ventana.after(450, animar)
    animar()

    resultado = [None]
    def escanear():
        resultado[0] = get_inventory_fast()
        ventana.after(0, lambda: construir_ui_formulario(ventana, resultado[0]))
    threading.Thread(target=escanear, daemon=True).start()


def construir_ui_formulario(ventana, data):
    limpiar_ventana(ventana)

    if not data:
        messagebox.showerror("Error", "No se pudo obtener información del hardware.")
        mostrar_menu_principal(ventana)
        return

    top_frame = tk.Frame(ventana, bg="#eef2f7")
    top_frame.pack(fill="x", pady=(14, 4), padx=30)

    tk.Button(top_frame, text="⬅ Volver", command=lambda: mostrar_menu_principal(ventana), bg="#cbd5e1", fg="#334155", font=("Segoe UI", 9, "bold"), cursor="hand2", relief="flat", padx=10).pack(side="left")
    tk.Label(top_frame, text="💻 Asistente de Entrega / Retiro", font=("Segoe UI", 15, "bold"), bg="#eef2f7", fg="#1a3c6e").pack(side="left", padx=50)

    info_frame = tk.Frame(ventana, bg="#dbeafe", bd=1, relief="solid")
    info_frame.pack(padx=30, fill="x")
    tk.Label(info_frame, text=f"  {data['model']}   |   S/N: {data['serial']}  ", font=("Consolas", 10), bg="#dbeafe", fg="#1d4ed8", pady=6).pack()

    marco = tk.Frame(ventana, bg="#eef2f7")
    marco.pack(pady=10, fill="x", padx=30)

    # ── LOGÍSTICA ──
    lf_log = ttk.LabelFrame(marco, text=" 🚚 Datos Logísticos ")
    lf_log.pack(fill="x", pady=(0, 10))
    frame_log_int = tk.Frame(lf_log, bg="#eef2f7")
    frame_log_int.pack(pady=5)

    tk.Label(frame_log_int, text="Movimiento:", bg="#eef2f7", font=("Segoe UI", 9)).grid(row=0, column=0, padx=8, pady=6, sticky="e")
    combo_mov = ttk.Combobox(frame_log_int, values=["Entrega", "Retiro"], width=13, state="readonly", font=("Segoe UI", 9))
    combo_mov.set("Entrega")
    combo_mov.grid(row=0, column=1, padx=6, pady=6, sticky="w")

    tk.Label(frame_log_int, text="N° de Guía:", bg="#eef2f7", font=("Segoe UI", 9)).grid(row=0, column=2, padx=8, sticky="e")
    entry_guia = tk.Entry(frame_log_int, width=13, font=("Segoe UI", 9), state=tk.DISABLED, fg="#aaa")
    entry_guia.grid(row=0, column=3, padx=6, pady=6, sticky="w")

    # ── CHECKLIST DETALLADO DE COMPONENTES ──
    lf_fallas = ttk.LabelFrame(marco, text=" ⚠️ Revisión de Componentes (Solo Retiros) ")

    componentes_etiqueta = [
        "Carcaza", "Pantalla", "Teclado", "Touchpad", "Puertos USB",
        "Puertos Video", "Ethernet/Wi-Fi", "Placa base", "Memoria",
        "Disco duro", "Batería"
    ]

    comps_data = []

    frame_grilla = tk.Frame(lf_fallas, bg="#eef2f7")
    frame_grilla.pack(padx=10, pady=10)

    tk.Label(frame_grilla, text="Componente", font=("Segoe UI", 9, "bold"), bg="#eef2f7").grid(row=0, column=0, sticky="w", padx=5)
    tk.Label(frame_grilla, text="Estado", font=("Segoe UI", 9, "bold"), bg="#eef2f7").grid(row=0, column=1, padx=5)
    tk.Label(frame_grilla, text="Observación (Obligatorio si es OBS/MALO)", font=("Segoe UI", 9, "bold"), bg="#eef2f7").grid(row=0, column=2, sticky="w", padx=5)

    for i, comp_name in enumerate(componentes_etiqueta, start=1):
        tk.Label(frame_grilla, text=comp_name, bg="#eef2f7", font=("Segoe UI", 9)).grid(row=i, column=0, sticky="w", padx=5, pady=2)

        cmb_estado = ttk.Combobox(frame_grilla, values=["OK", "OBS", "MALO"], width=6, state="readonly", font=("Segoe UI", 9))
        cmb_estado.set("OK")
        cmb_estado.grid(row=i, column=1, padx=5, pady=2)

        ent_obs = tk.Entry(frame_grilla, width=40, font=("Segoe UI", 9), state=tk.DISABLED, bg="#f1f5f9")
        ent_obs.grid(row=i, column=2, padx=5, pady=2, sticky="w")

        def toggle_entry(event, e=ent_obs, c=cmb_estado):
            if c.get() != "OK":
                e.config(state=tk.NORMAL, bg="#ffffff")
                e.focus()
            else:
                e.delete(0, tk.END)
                e.config(state=tk.DISABLED, bg="#f1f5f9")

        cmb_estado.bind("<<ComboboxSelected>>", toggle_entry)

        comps_data.append({"nombre": comp_name, "estado": cmb_estado, "obs": ent_obs})

    def toggle_guia(event=None):
        if combo_mov.get() == "Retiro":
            entry_guia.config(state=tk.NORMAL, fg="#000")
            lf_fallas.pack(fill="x", pady=(0, 10))
            guia_guardada = cargar_guia_config()
            if guia_guardada:
                entry_guia.delete(0, tk.END)
                entry_guia.insert(0, guia_guardada)
        else:
            entry_guia.delete(0, tk.END)
            entry_guia.config(state=tk.DISABLED, fg="#aaa")
            lf_fallas.pack_forget()
            for c in comps_data:
                c["estado"].set("OK")
                c["obs"].config(state=tk.NORMAL)
                c["obs"].delete(0, tk.END)
                c["obs"].config(state=tk.DISABLED, bg="#f1f5f9")

    combo_mov.bind("<<ComboboxSelected>>", toggle_guia)

    # ── OBSERVACIÓN GENERAL ──
    tk.Label(marco, text="Observación General para el Acta:", bg="#eef2f7", font=("Segoe UI", 10, "bold"), fg="#334155").pack(anchor="w", pady=(4, 2))
    entry_obs_gen = tk.Entry(marco, width=50, font=("Segoe UI", 10))
    entry_obs_gen.pack(fill="x", ipady=4, pady=(0, 8))

    # ── OFFICE ──
    frame_off = tk.Frame(marco, bg="#eef2f7")
    frame_off.pack(fill="x", pady=(4, 0))
    var_office = tk.IntVar()
    def toggle_office():
        state = tk.NORMAL if var_office.get() else tk.DISABLED
        entry_office.config(state=state)
        combo_ver.config(state="readonly" if var_office.get() else tk.DISABLED)
        if not var_office.get(): entry_office.delete(0, tk.END)

    tk.Checkbutton(frame_off, text="📦 Registrar Office", variable=var_office, command=toggle_office, bg="#eef2f7", font=("Segoe UI", 10, "bold"), fg="#1e40af", activebackground="#eef2f7", cursor="hand2").pack(side="left")
    combo_ver = ttk.Combobox(frame_off, values=["2013","2016","2019","2021","2024","365"], width=8, state=tk.DISABLED, font=("Segoe UI", 10))
    combo_ver.set("2016")
    combo_ver.pack(side="left", padx=6)

    entry_office = tk.Entry(marco, font=("Consolas", 10), justify="center", state=tk.DISABLED, fg="#555")
    entry_office.pack(fill="x", ipady=4, pady=6)

    # ── BOTONES ──
    frame_bts = tk.Frame(ventana, bg="#eef2f7")
    frame_bts.pack(pady=10)

    tk.Button(frame_bts, text="💾 GUARDAR EQUIPO EN LOTE",
              command=lambda: accion_agregar_lote(ventana, data, entry_obs_gen.get().strip(), var_office.get()==1, entry_office.get().strip(), combo_ver.get(), combo_mov.get(), entry_guia.get().strip(), comps_data),
              bg="#0078d4", fg="white", font=("Segoe UI", 12, "bold"), padx=20, pady=10, cursor="hand2", relief="flat").pack()

def iniciar_interfaz_principal():
    ventana = tk.Tk()
    ventana.title("Asistente de Entrega — Arrienda.cl")
    ventana.resizable(False, False)
    ventana.attributes("-topmost", True)
    ventana.configure(bg="#eef2f7")

    style = ttk.Style()
    style.theme_use("clam")
    style.configure("TCombobox",   fieldbackground="#fff", background="#fff", font=("Segoe UI", 9))
    style.configure("TEntry",      fieldbackground="#fff", font=("Segoe UI", 9))
    style.configure("TLabelframe", background="#eef2f7")
    style.configure("TLabelframe.Label", background="#eef2f7", font=("Segoe UI", 10, "bold"), foreground="#334155")

    mostrar_menu_principal(ventana)
    ventana.mainloop()

if __name__ == "__main__":
    iniciar_interfaz_principal()
