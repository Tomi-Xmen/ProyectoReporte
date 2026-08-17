"""
Sistema de Control de Inventario — Arrienda.cl
==============================================

Aplicación de escritorio (Tkinter) para escanear el hardware de un equipo,
registrar entregas/retiros y generar reportes HTML, JSON y Word.

ÍNDICE DEL ARCHIVO (buscá el banner "==== NOMBRE ====" para saltar):

  1. CONFIGURACIÓN            — rutas y constantes globales.
  2. TEMA Y HELPERS DE UI      — paleta (COLORS), fuentes y widgets reutilizables
                                 de Tkinter (titulo_ui, ventana_modal, boton_menu,
                                 boton_accion, marco_scrolleable, treeview_auditoria).
  3. SCANNER (PowerShell)      — PS_SCRIPT: script embebido que lee el hardware.
  4. PLANTILLA HTML DEL REPORTE — HTML_START / HTML_END (CSS + JS del reporte).
  5. EXPORT WORD               — helpers docx + generar_informe_word (informe de
                                 fallas; lo dispara cerrar_lote, ver sección 7).
  6. SCANNER (ejecución)       — get_inventory_fast: corre PS_SCRIPT y parsea.
  7. PERSISTENCIA              — modelo de LOTES por OT/guía, construir_entry_html,
                                 guardar_equipo_general, cerrar_lote, crear_backup,
                                 envío al servidor por SCP.
  8. RECONSTRUCCIÓN / MERGE    — rearmar HTML desde los JSON guardados.
  9. MÓDULOS (Toplevel)        — etiquetas manuales, transformadores, panel de
                                 clonación y cambio de nombre del equipo.
 10. UI PRINCIPAL              — menú, pantalla de escaneo y formulario.
 11. MODO CLONACIÓN            — flujo headless --clonacion para el post-script
                                 de FOG (sin menú ni elección de OT).
 12. ENTRYPOINT                — iniciar_interfaz_principal / __main__.

DISPOSICIÓN EN DISCO (todo cuelga de ./Reportes_Guardados):

    ENTREGA/<CLIENTE>/OT-553/      equipos, HTML y FUSIONADO de esa OT
    RETIROS/<CLIENTE>/GUIA-8891/   ídem para retiros (+ informe Word de fallas,
                                   que se escribe recién al CERRAR la OT)
    _BACKUPS/Backup_<fecha>_<hora>_<motivo>.zip
    .lote_activo.json              puntero al lote que recibe lo escaneado

Los marcadores dentro de cada carpeta de OT mandan sobre su estado:
'.cerrado' = finalizada (los recorridos la saltan sin abrirla) y
'.enviado' = ya subida al servidor (control local; no viaja).

El envío al servidor va por SCP (SSH) y sube SOLO lo que es fuente de verdad o
entregable: los JSON individuales, el FUSIONADO, el Word de fallas y los
marcadores. El HTML no viaja: es un derivado y allá se rearma con
"RECONSTRUIR DESDE JSONs", que produce un reporte idéntico y con todas sus
funcionalidades (copiar filas, CSV de Valida, fusionado, etiquetas).

NOTA: la apariencia del REPORTE HTML vive en HTML_START/HTML_END y en
construir_entry_html; no se debe alterar sin querer cambiar el reporte.
"""

import html  # Para escapar strings en el HTML
import json
import os
import re
import socket
import subprocess
import sys
import tempfile
import threading
import webbrowser
import time
import zipfile
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

# --- SOLUCIÓN PYINSTALLER: Asegurar que el directorio de trabajo sea siempre el correcto ---
if getattr(sys, "frozen", False):
    # Si el programa está compilado (.exe) con PyInstaller
    ruta_base_proyecto = os.path.dirname(sys.executable)
else:
    # Si se está ejecutando como script de Python normal (.py)
    ruta_base_proyecto = os.path.dirname(os.path.abspath(__file__))

# Cambiamos el directorio de trabajo a la carpeta real del .exe
os.chdir(ruta_base_proyecto)
# -------------------------------------------------------------------------------------------

# Importar librerías para el envío por SCP (SSH). Si faltan, el botón de envío
# avisa en vez de reventar: el resto de la app (escaneo, reportes) no depende
# de ellas.
try:
    import paramiko
    from scp import SCPClient

    SCP_AVAILABLE = True
except ImportError:
    SCP_AVAILABLE = False

# Importar librería para crear Word (docx)
try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# 1. CONFIGURACIÓN
# ----------------------------------------------------------------------------
# ENVÍO AL SERVIDOR — por SCP sobre SSH (reemplaza el copiado a la carpeta
# compartida SMB). Cambiar cualquiera de estas constantes exige recompilar el
# .exe: son la única fuente de verdad de la conexión.
#
# Al servidor viajan SOLO los archivos que son FUENTE DE VERDAD o entregables:
# los JSON individuales, el JSON FUSIONADO, el informe Word y el marcador
# '.cerrado'. El HTML NO se sube: es un derivado y se rearma allá con
# "RECONSTRUIR DESDE JSONs", que produce un reporte idéntico y 100% funcional.
# ============================================================================
SCP_HOST = "192.168.10.15"         # IP o nombre del servidor SSH
SCP_PUERTO = 22                    # puerto SSH
SCP_USUARIO = "Reporte"            # usuario SSH
# Clave privada SSH. Si queda vacía se intenta con las claves del agente o de
# ~/.ssh del usuario de Windows.
SCP_CLAVE = r"C:\Users\Tomas\.ssh\id_clonado"
SCP_CLAVE_PASSPHRASE = ""          # vacío si la clave no tiene passphrase
# Carpeta base REMOTA (ruta estilo Linux). Debajo se replica el árbol completo
# <GRUPO>/<CLIENTE>/<OT-553>, igual que en local: así lo que se baja del
# servidor se reconstruye como OT canónica y no como "carpeta suelta".
SCP_DESTINO = (
    "/srv/dev-disk-by-uuid-50e34ae9-5743-420d-b093-89dfd29299cd"
    "/DocumentosLab/Reportes_Guardados"
)
SCP_TIMEOUT = 20                   # segundos para conectar antes de fallar


# ============================================================================
# 2. TEMA Y HELPERS DE UI (solo para la ventana Tkinter de escritorio)
# ----------------------------------------------------------------------------
# Paleta y helpers centralizados para la interfaz Tkinter. NO afectan al HTML
# del reporte (ese usa sus propios colores dentro de HTML_START/HTML_END y de
# las plantillas f-string). Cambiar el look de la app se hace desde aquí.
# ============================================================================

# Color de marca (azul corporativo) como fuente única para Tkinter y Word.
MARCA_RGB = (0x1A, 0x3C, 0x6E)
MARCA_HEX = "1A3C6E"

COLORS = {
    "fondo": "#eef2f7",          # fondo general de ventanas
    "azul": "#1a3c6e",           # azul de marca (títulos, botones de ruta)
    "azul_btn": "#0078d4",       # botón principal "escanear"
    "verde": "#16a34a",          # acción / OK
    "verde_esmeralda": "#059669",
    "naranja": "#d97706",        # advertencia / enviar red
    "morado": "#7c3aed",         # reconstruir
    "celeste": "#0ea5e9",        # manejo de la OT (nueva / cambiar / refrescar)
    "pizarra": "#475569",        # etiquetas
    "gris": "#64748b",           # texto secundario
    "gris_oscuro": "#334155",
    "blanco": "white",
}


def fuente(size=9, bold=False):
    """Tupla de fuente estándar de la app (Segoe UI)."""
    return ("Segoe UI", size, "bold") if bold else ("Segoe UI", size)


def titulo_ui(parent, texto, size=14, pady=12, fg=None):
    """Label de título estándar (azul de marca sobre fondo claro)."""
    lbl = tk.Label(
        parent,
        text=texto,
        font=fuente(size, True),
        bg=COLORS["fondo"],
        fg=fg or COLORS["azul"],
    )
    lbl.pack(pady=pady)
    return lbl


def ventana_modal(padre, titulo_ventana, geometria):
    """Crea un Toplevel modal con el fondo y grab_set estándar."""
    win = tk.Toplevel(padre)
    win.title(titulo_ventana)
    win.geometry(geometria)
    win.configure(bg=COLORS["fondo"])
    win.grab_set()
    return win


def marco_scrolleable(padre, bg=None):
    """
    Área con scroll vertical. Devuelve (contenedor, interior, canvas): el
    contenedor se empaqueta como cualquier widget y los hijos van en 'interior'.

    Un Frame de Tk no scrollea solo; el truco estándar es meterlo dentro de un
    Canvas y mover el Canvas. La rueda del mouse se engancha al entrar y se
    suelta al salir, porque bind_all es global y si no queda secuestrada
    después de que el área desaparezca.
    """
    fondo = bg or COLORS["fondo"]
    contenedor = tk.Frame(padre, bg=fondo)
    canvas = tk.Canvas(contenedor, bg=fondo, highlightthickness=0)
    barra = tk.Scrollbar(contenedor, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=barra.set)
    barra.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    interior = tk.Frame(canvas, bg=fondo)
    ventana_interna = canvas.create_window((0, 0), window=interior, anchor="nw")
    interior.bind(
        "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )
    canvas.bind(
        "<Configure>", lambda e: canvas.itemconfigure(ventana_interna, width=e.width)
    )

    def _rueda(event):
        canvas.yview_scroll(-1 * (event.delta // 120), "units")

    canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _rueda))
    canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
    return contenedor, interior, canvas


def boton_menu(parent, texto, comando, color, size=11, bold="bold"):
    """Botón ancho del menú/pantallas (relleno horizontal)."""
    tk.Button(
        parent,
        text=texto,
        command=comando,
        bg=color,
        fg="white",
        font=("Segoe UI", size, bold),
        pady=10,
        cursor="hand2",
        relief="flat",
    ).pack(fill="x", pady=6)


def boton_accion(parent, texto, comando, color=None, padx=18, pady=7):
    """Botón de acción destacado (por defecto verde). El caller lo empaqueta."""
    return tk.Button(
        parent,
        text=texto,
        command=comando,
        bg=color or COLORS["verde"],
        fg="white",
        font=fuente(10, True),
        padx=padx,
        pady=pady,
        cursor="hand2",
        relief="flat",
    )


def treeview_auditoria(win, cols_spec, height):
    """
    Tabla estándar de resultados (estilo + columnas + tags de color + scrollbar
    dentro de su frame), lista para insertar filas.

    cols_spec: lista de (col_id, ancho, encabezado). Las columnas 'estado' y
    'serial' se centran; el resto va alineado a la izquierda.

    Tags de color disponibles al insertar: 'ok' (verde), 'falta' (rojo) y
    'sobra' (ámbar). Hoy su único consumidor es el panel de clonación; el
    nombre viene de los módulos de auditoría que ya no existen.
    """
    style = ttk.Style()
    style.configure("Audit.Treeview", rowheight=24, font=fuente(9))
    style.configure("Audit.Treeview.Heading", font=fuente(9, True))

    cols = tuple(c[0] for c in cols_spec)
    tree = ttk.Treeview(
        win, columns=cols, show="headings", height=height, style="Audit.Treeview"
    )
    for col, w, txt in cols_spec:
        tree.heading(col, text=txt)
        tree.column(
            col, width=w, anchor="center" if col in ("estado", "serial") else "w"
        )

    tree.tag_configure("ok", foreground="#16a34a")
    tree.tag_configure("falta", foreground="#dc2626", background="#fee2e2")
    tree.tag_configure("sobra", foreground="#d97706", background="#fef3c7")

    sb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=sb.set)
    frame_tree = tk.Frame(win, bg=COLORS["fondo"])
    frame_tree.pack(fill="both", expand=True, padx=20, pady=8)
    tree.pack(side="left", fill="both", expand=True, in_=frame_tree)
    sb.pack(side="right", fill="y", in_=frame_tree)
    return tree

# ============================================================================
# 3. SCANNER (PowerShell)
# ----------------------------------------------------------------------------
# Script embebido que recolecta el hardware vía CIM/WMI + powercfg y lo emite
# como JSON. Se ejecuta desde get_inventory_fast() (sección 6).
# ============================================================================
PS_SCRIPT = """
$ErrorActionPreference = 'SilentlyContinue'
$bios   = Get-CimInstance Win32_Bios
$sys    = Get-CimInstance Win32_ComputerSystem
$cpu    = Get-CimInstance Win32_Processor
$svc    = Get-CimInstance SoftwareLicensingService
$winKey = if ($svc.OA3xOriginalProductKey) { $svc.OA3xOriginalProductKey } else { "No encontrada en BIOS" }

# --- NUEVO: LÓGICA DE BATERÍA CON POWERCFG XML ---
$battInfo = "No detectada"
$xmlPath = "$env:TEMP\\batt_report_temp.xml"
if (Test-Path $xmlPath) { Remove-Item $xmlPath -Force -ErrorAction SilentlyContinue }

# Ejecutamos powercfg silenciosamente para que genere el XML
& powercfg /batteryreport /xml /output $xmlPath | Out-Null

if (Test-Path $xmlPath) {
    try {
        [xml]$battXml = Get-Content $xmlPath
        $baterias = $battXml.BatteryReport.Batteries.Battery
        if ($baterias) {
            # Tomar la primera batería por si el equipo tiene dos
            $bat = if ($baterias.Count -gt 1) { $baterias[0] } else { $baterias }

            $design = [int]$bat.DesignCapacity
            $full = [int]$bat.FullChargeCapacity

            if ($design -gt 0) {
                $salud = [math]::Round(($full / $design) * 100)
                if ($salud -gt 100) { $salud = 100 } # Tope visual al 100%
                $battInfo = "$salud% | Diseño: $design mWh | Actual: $full mWh"
            }
        }
    } catch {
        $battInfo = "Error al leer datos XML"
    }
    # Autodestruir el reporte temporal para no dejar basura
    Remove-Item $xmlPath -Force -ErrorAction SilentlyContinue
}
# ------------------------------------------------

$ramInfo = @(Get-CimInstance Win32_PhysicalMemory | ForEach-Object {
    $speed = if ($_.Speed) { "$($_.Speed)" } else { "Desconocida" }
    $Partnumber = if ($_.PartNumber) { $_.PartNumber.Trim() } else { "Desconocida" }
    $Manufacturer = if ($_.Manufacturer) { $_.Manufacturer.Trim() } else { "Desconocida" }
    $typeNum = if ($_.SMBIOSMemoryType) { $_.SMBIOSMemoryType } else { $_.MemoryType }
    $typeStr = switch ($typeNum) { 24 {"PC3"} 26 {"PC4"} 30 {"LPDDR4"} 34 {"PC5"} 35 {"LPDDR5"} default {"PC4"} }
    "{0}GB | {1} | {2} | {3} | {4}" -f ([math]::Round($_.Capacity/1GB)), $speed, $typeStr, $Manufacturer, $Partnumber
})

$disksData = @(Get-PhysicalDisk | Where-Object { $_.BusType -notin @('USB','File Backed Virtual') } | ForEach-Object {
    $size  = [math]::Round($_.Size / 1000000000)
    $media = if ([string]::IsNullOrWhiteSpace($_.MediaType) -or $_.MediaType -eq 'Unspecified') { 'SSD' } else { [string]$_.MediaType }
    $bus   = if ([string]$_.BusType -match 'NVMe') { 'M.2' } else { [string]$_.BusType }
    @{ desc = "{0} {1}GB {2} {3}" -f $_.FriendlyName.Trim(), $size, $media, $bus
       serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
})
if ($disksData.Count -eq 0) {
    $disksData = @(Get-CimInstance Win32_DiskDrive | Where-Object { $_.InterfaceType -ne 'USB' } | ForEach-Object {
        @{ desc = "{0} {1}GB SSD" -f $_.Model.Trim(), ([math]::Round($_.Size/1000000000))
           serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
    })
}

$netInfo = @(Get-CimInstance Win32_NetworkAdapterConfiguration | Where-Object { $_.IPEnabled } | ForEach-Object {
    "{0} | IP: {1} | MAC: {2}" -f $_.Description, $_.IPAddress[0], $_.MACAddress
})

@{ serial=($bios.SerialNumber); model=($sys.Model); win_key=($winKey)
   cpu=($cpu.Name); ram_rows=$ramInfo; disks_data=$disksData; net_rows=$netInfo; battery=($battInfo)
} | ConvertTo-Json -Compress -Depth 3
"""

# ─────────────────────────────────────────────────────────
#  HTML / CSS / JS  (plantilla del reporte)
# ─────────────────────────────────────────────────────────
# ============================================================================
# 4. PLANTILLA HTML DEL REPORTE  (⚠️ apariencia del reporte — NO modificar)
# ----------------------------------------------------------------------------
# HTML_START trae toda la página (CSS + JS de estadísticas, etiquetas, export
# CSV, etc.) y termina en <div id="main-list">. HTML_END cierra las etiquetas.
# Las entradas de cada equipo se insertan entre ambos (ver construir_entry_html).
# ============================================================================
HTML_START = r"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Inventario Maestro</title>
<style>
  *{box-sizing:border-box;margin:0;padding:0}
  body{font-family:'Segoe UI',Tahoma,sans-serif;background:#eef2f7;color:#1e293b;padding:30px}
  .container{max-width:1050px;margin:auto}
  .header{text-align:center;margin-bottom:28px}
  h1{font-size:26px;color:#1a3c6e;margin-bottom:6px}
  .subtitle{color:#64748b;font-size:13px;margin-bottom:14px}
  .counter-box{background:#dbeafe;color:#1d4ed8;padding:5px 18px;border-radius:20px;
    font-weight:700;font-size:13px;display:inline-block;border:1px solid #bfdbfe;margin-bottom:12px}

  .stats-container { display:flex; justify-content:center; gap:10px; margin-bottom:22px; flex-wrap:wrap; }
  .stat-pill { padding:5px 16px; border-radius:20px; font-weight:700; font-size:12px; border:1px solid; }
  .stat-ram { background:#fce7f3; color:#be185d; border-color:#fbcfe8; }
  .stat-cpu { background:#fef3c7; color:#b45309; border-color:#fde68a; }
  .stat-disk { background:#dcfce7; color:#15803d; border-color:#bbf7d0; }

  .button-group{display:flex;justify-content:center;gap:12px;margin-bottom:32px;flex-wrap:wrap}
  .btn{color:#fff;border:none;padding:10px 22px;font-size:13px;font-weight:700;border-radius:8px;
    cursor:pointer;box-shadow:0 2px 6px rgba(0,0,0,.18);transition:transform .1s,box-shadow .1s}
  .btn:hover{transform:translateY(-1px);box-shadow:0 4px 12px rgba(0,0,0,.22)}
  .btn:active{transform:translateY(0)}
  .btn-copy{background:linear-gradient(135deg,#0078d4,#005a9e)}
  .btn-export{background:linear-gradient(135deg,#16a34a,#15803d)}
  .btn-delete{background:#ef4444;padding:5px 12px;font-size:11px;border-radius:6px}
  .btn-delete:hover{background:#dc2626}

  details{background:#fff;margin-bottom:12px;border-radius:12px;
    box-shadow:0 2px 8px rgba(0,0,0,.08);overflow:hidden;border:1px solid #e2e8f0}
  summary{padding:14px 18px;font-weight:700;cursor:pointer;
    background:linear-gradient(135deg,#1a3c6e,#0f2d5a);color:#fff;
    list-style:none;display:flex;justify-content:space-between;align-items:center;gap:10px}
  summary:hover{background:linear-gradient(135deg,#1e4a84,#142f60)}
  summary::-webkit-details-marker{display:none}

  .tag-entrega{background:#22c55e;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-retiro{background:#f97316;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-falla{background:#ef4444;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700;margin-left:8px}
  .logistics-row{background:#f8fafc}

  .content{padding:20px 22px;border-top:1px solid #f1f5f9}
  .content-header{display:flex;justify-content:flex-end;margin-bottom:10px; gap: 8px;}
  table{border-collapse:collapse;width:100%}
  th,td{border:1px solid #e2e8f0;padding:9px 12px;text-align:left;font-size:13px}
  th{background:#f8fafc;width:28%;color:#475569;font-weight:600}
  tr:hover td,tr:hover th{background:#f0f9ff}
  .serial-tag{color:#dc2626;font-weight:700;font-family:Consolas,monospace}
  .net-tag{background:#dcfce7;color:#166534;padding:2px 7px;border-radius:4px;font-weight:600;font-size:12px}
  .win-key{color:#1d4ed8;font-family:Consolas,monospace;font-weight:700;font-size:12px}
  .obs-row{background:#fef9c3!important}
  .obs-row td,.obs-row th{color:#854d0e;font-weight:600;font-style:italic}
  .falla-row{background:#fee2e2!important}
  .falla-row td,.falla-row th{color:#991b1b;font-weight:600}
  .office-row{background:#eff6ff}
  .office-row td,.office-row th{color:#1e40af}
</style>
<script>
function actualizarEstadisticas() {
  let eqs = document.querySelectorAll('details');
  document.getElementById('total-count').innerText = eqs.length;

  let cpus = {}, rams = {}, discos = {};

  eqs.forEach(eq => {
    let d = extraerDatos(eq);

    let cpuU = d.cpu.toUpperCase();
    let cpuKey = "Otro";
    if(cpuU.includes("I3")) cpuKey = "i3";
    else if(cpuU.includes("I5")) cpuKey = "i5";
    else if(cpuU.includes("I7")) cpuKey = "i7";
    else if(cpuU.includes("I9")) cpuKey = "i9";
    else if(cpuU.includes("RYZEN 3")) cpuKey = "Ryzen 3";
    else if(cpuU.includes("RYZEN 5")) cpuKey = "Ryzen 5";
    else if(cpuU.includes("RYZEN 7")) cpuKey = "Ryzen 7";
    cpus[cpuKey] = (cpus[cpuKey] || 0) + 1;

    let ramMatch = d.specs.match(/^(\d+GB)/);
    let ramKey = ramMatch ? ramMatch[1] : "Otra";
    rams[ramKey] = (rams[ramKey] || 0) + 1;

    d.discos.forEach(dk => {
      let desc = dk.desc.toUpperCase();
      let tipo = "Desconocido";
      if(desc.includes("M.2") || desc.includes("NVME")) tipo = "M.2 NVMe";
      else if(desc.includes("SSD")) tipo = "SSD 2.5\"";
      else if(desc.includes("HDD")) tipo = "HDD";
      discos[tipo] = (discos[tipo] || 0) + 1;
    });
  });

  let mkHTML = (obj, clase, prefijo) => {
    let arr = Object.entries(obj).map(([k,v]) => `${k} (${v})`);
    return arr.length ? `<span class="stat-pill ${clase}">${prefijo}: ${arr.join(' | ')}</span>` : '';
  };

  let statsDiv = document.getElementById('stats-container');
  if(statsDiv) {
    statsDiv.innerHTML = mkHTML(rams, 'stat-ram', 'RAM') +
                         mkHTML(cpus, 'stat-cpu', 'CPU') +
                         mkHTML(discos, 'stat-disk', 'Discos');
  }
}

document.addEventListener("DOMContentLoaded", actualizarEstadisticas);

function eliminarEquipo(id){
  if(confirm("¿Quitar este equipo del reporte actual?\n(No afecta los archivos ya guardados)")){
    let n=document.getElementById(id);
    if(n){
      n.remove();
      actualizarEstadisticas();
    }
  }
}

function extraerDatos(eq){
  let model  = eq.querySelector('.model-name').innerText.trim();
  let serial = eq.querySelector('.serial-tag').innerText.trim();
  let tds=eq.querySelectorAll('th,td');
  let obs="",cpu="",office="",guia="",mov="",ramSpeed="",ramTech="PC4",fallas="",Manufacturer="",Partnumber="Desconocida",ot="",transf="";
  let totalRam=0,ramArr=[],discos=[],tempDesc="";

  for(let i=0;i<tds.length;i++){
    if(tds[i].tagName !== 'TH') continue;

    let t=tds[i].innerText.trim();
    if(t==='N° de Guía')          guia   =tds[i+1].innerText.trim();
    if(t==='N° de OT')            ot     =tds[i+1].innerText.trim();
    if(t==='Tipo Movimiento')     mov    =tds[i+1].innerText.trim();
    if(t.includes('Obs. General'))obs    =tds[i+1].innerText.trim();
    if(t==='Fallas Detectadas')   fallas =tds[i+1].innerText.trim();
    if(t==='Procesador')          cpu    =tds[i+1].innerText.trim();
    if(t.includes('Office'))      office =tds[i+1].innerText.trim();
    if(t==='Discos Internos')     tempDesc=tds[i+1].innerText.trim();
    if(t==='Serie Disco Duro')    discos.push({desc:tempDesc,serial:tds[i+1].innerText.trim()});
    /* includes y no ===: el th lleva el emoji 🔌 adelante. */
    if(t.includes('Serie Transformador')) transf=tds[i+1].innerText.trim();
    if(t==='Módulo RAM'){
      let p=tds[i+1].innerText.split('|');
      let n=parseInt(p[0].replace(/\D/g,''));
      if(!isNaN(n)){totalRam+=n;ramArr.push(n)}
      if(p[1]&&ramSpeed==="")ramSpeed=p[1].trim();
      if(p[2]&&ramTech==="PC4")ramTech=p[2].trim();
      if(p[3]&&Manufacturer==="")Manufacturer=p[3].trim()
      if(p[4]&&Partnumber==="Desconocida")Partnumber=p[4].trim()
    }
  }

  let ramFinal="RAM Desconocida";
  if(totalRam>0){
    let det=ramArr.length>1?(ramArr.every(v=>v===ramArr[0])?`(${ramArr[0]}X${ramArr.length})`:`(${ramArr.join('+')})`):"";
    let ts=ramSpeed!=="Desconocida"?`${ramTech}-${ramSpeed}Mhz`:ramTech;
    ramFinal=`${totalRam}GB${det} (${ts})`;
  }
  let diskSimple=discos.map(d=>{let m=d.desc.match(/(\d+)\s*GB\s*(.*)/i);return m?m[1]+" GB "+m[2]:d.desc});
  let specs=ramFinal+(diskSimple.length>0?", "+diskSimple.join(" + "):"");
  if(obs&&obs!=="Sin observaciones")specs+=", OBS: "+obs;
  if(fallas&&fallas!=="Ninguna")specs+=", FALLAS: "+fallas;

  return {model,serial,cpu,obs,office,guia,mov,discos,specs,equipoFull:model+" "+cpu, fallas, Manufacturer, Partnumber, ot, transf};
}

/* --- EXPORT A VALIDA: COPIAR FILAS Y CSV --- */

const MSG_COPIADO="✅ ¡Filas copiadas!\n\nPega directamente en tu Excel / Valida.";

/* Filas que un equipo aporta a Valida, como [nombre, descripción, serie, obs].
   Fuente ÚNICA de copiarFilas y exportarValidaCSV: las dos salidas describen el
   mismo equipo y tienen que decir lo mismo. Cuando cada una armaba sus filas por
   su cuenta terminaron divergiendo (la fila del transformador salía distinta en
   cada export). Accesorio nuevo = una línea acá y aparece en las dos. */
function filasValida(d){
  let filas=[[d.equipoFull, d.specs, d.serial, d.obs]];

  d.discos.forEach(dk=>{
    let lbl=dk.desc.toUpperCase().includes("HDD")?"HDD":"SSD";
    filas.push([lbl, dk.desc, dk.serial, ""]);
  });

  /* Sin 'Key:' no hay nada que cargar: la fila de Office existe en el HTML
     aunque SN_APP esté PENDIENTE (ver construir_entry_html), y exportarla daría
     una fila con la serie en blanco. */
  if(d.office && d.office.includes('Key:')){
    let ver=(d.office.match(/Office\s+([A-Za-z0-9]+)/)||[])[1]||"2016";
    let key=(d.office.match(/Key:\s*([^)]+)/)||[])[1]||"";
    filas.push([`OFFICE ${ver}`, "HOME AND BUSINESS", key.trim(), ""]);
  }

  /* El transformador va sin nombre: en Valida se carga como accesorio y con la
     descripción alcanza. */
  if(d.transf) filas.push(["", "TRANSFORMADOR C/ADAPTADOR", d.transf, ""]);

  return filas;
}

function copiarFilas(){
  let txt="";
  document.querySelectorAll('details').forEach(eq=>{
    /* Al portapapeles van solo las 3 primeras columnas: la obs ya viaja dentro
       de specs y la cuarta es para el CSV. */
    filasValida(extraerDatos(eq)).forEach(f=>{ txt+=f.slice(0,3).join("\t")+"\n" });
  });

  if(!txt){ alert("No hay equipos en el reporte para copiar."); return; }

  if(navigator.clipboard && window.isSecureContext){
      navigator.clipboard.writeText(txt)
        .then(()=>alert(MSG_COPIADO))
        .catch(()=>copiarConTextarea(txt));
  } else {
      copiarConTextarea(txt);
  }
}

/* Respaldo del portapapeles. Se usa en dos casos: el reporte abierto con
   file://, donde navigator.clipboard no existe, y el rebote de la API cuando la
   ventana perdió el foco —de ahí que el catch caiga acá y no en un alert de
   error. */
function copiarConTextarea(txt){
  let ta=document.createElement("textarea");
  ta.value=txt;
  ta.style.cssText="position:fixed;left:-9999px;top:0";  /* que select() no salte el scroll */
  document.body.appendChild(ta);
  ta.select();
  try{ document.execCommand('copy'); alert(MSG_COPIADO); }
  catch(e){ alert("Error al copiar al portapapeles.") }
  document.body.removeChild(ta);
}

/* Nombre base compartido por las descargas: Reporte_<tipo>_<cliente>_<AAAAMMDD> */
function nombreBaseReporte(){
  let titulo=document.querySelector('title').innerText;
  let tipo=titulo.includes("Entregas")?"Entregas":(titulo.includes("Retiros")?"Retiros":"General");
  let cliente=titulo.replace(tipo+" - ","").replace("Inventario Maestro","General").replace(/ /g,"_");
  return `Reporte_${tipo}_${cliente}_${new Date().toISOString().slice(0,10).replace(/-/g,"")}`;
}

function descargarArchivo(nombre, contenido, mime){
  let blob=new Blob([contenido],{type:mime}),a=document.createElement("a");
  a.href=URL.createObjectURL(blob);
  a.download=nombre;
  a.style.visibility='hidden';document.body.appendChild(a);a.click();document.body.removeChild(a);
}

function exportarValidaCSV(){
  let csv="\uFEFFTIPO;GUIA;NOMBRE DEL EQUIPO;DESCRIPCION;NUMERO DE SERIE;OBSERVACIONES\n";
  document.querySelectorAll('details').forEach(eq=>{
    let d=extraerDatos(eq);
    /* Movimiento y gu\u00EDa son del equipo, no de la fila: se repiten en todas las
       filas que aporta, que es como Valida espera el CSV. */
    filasValida(d).forEach(f=>{
      csv+=`"${d.mov}";"${d.guia}";"${f[0]}";"${f[1]}";"${f[2]}";"${f[3]}"\n`;
    });
  });
  descargarArchivo(nombreBaseReporte()+".csv", csv, 'text/csv;charset=utf-8;');
}

/* --- EXPORT DEL JSON FUSIONADO --- */

/* Respaldo para entradas sin data-fusion (reportes armados antes de que
   existiera el atributo): se rearma el resumen leyendo la tabla. */
function resumenFusionadoDesdeDOM(eq){
  let val=sel=>{let n=eq.querySelector(sel);return n?n.innerText.trim():""};
  let hdd=[];
  eq.querySelectorAll('th').forEach(th=>{
    if(th.innerText.trim()==='Serie Disco Duro'&&th.nextElementSibling)
      hdd.push(th.nextElementSibling.innerText.trim());
  });
  let obs=val('.obs-row td');
  let key=((val('.office-row td').match(/Key:\s*([^)]+)/)||[])[1]||"").trim();
  return {
    SERIAL: val('.serial-tag'),
    SN_Win: val('.win-key'),
    SN_HDD: hdd.join("_")||"PENDIENTE",
    SN_Transf: val('.transf-tag')||"PENDIENTE",
    SN_APP: key||null,
    OBS: (obs&&obs!=="Sin observaciones")?obs:null
  };
}

function exportarJsonFusionado(){
  let equipos=[];
  document.querySelectorAll('details').forEach(eq=>{
    let raw=eq.getAttribute('data-fusion');
    if(raw){
      try{ equipos.push(JSON.parse(raw)); return; }catch(e){}
    }
    equipos.push(resumenFusionadoDesdeDOM(eq));
  });
  if(!equipos.length){ alert("No hay equipos en el reporte para exportar."); return; }
  descargarArchivo(
    nombreBaseReporte()+"_FUSIONADO.json",
    JSON.stringify(equipos,null,4),
    'application/json;charset=utf-8;'
  );
}

/* --- LOGICA DE IMPRESIÓN CENTRALIZADA --- */

function getPrintStyle() {
  return `
    <style>
      @media print {
        @page { margin: 15mm; size: auto; }
        body { margin: 0; }
        .no-print { display: none; }
        .page-break { page-break-after: always; }
      }
      body { font-family: 'Arial', sans-serif; font-size: 11px; color: #000; max-width: 700px; margin: auto; padding: 20px; line-height: 1.4; }
      .header-title { display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 20px; font-size: 13px; }
      .row { margin-bottom: 10px; font-size: 11px;}
      .row span { margin-right: 25px; }
      table { width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 20px; }
      th, td { border: 1px solid #000; padding: 5px 8px; text-align: left; }
      .center { text-align: center; }
      .checkboxes { font-family: 'Segoe UI Symbol', sans-serif; }
      .bold { font-weight: bold; }
    </style>
  `;
}

function generarHtmlEtiqueta(d, comps) {
  let fecha = new Date().toLocaleDateString('es-CL');
  let cpuLimpia = d.cpu.replace(/Intel\(R\)|Core\(TM\)/gi, "").trim();

  let trs = comps.map(c => {
     let iconOK   = c.estado === "OK" ? "☑" : "□";
     let iconOBS  = c.estado === "OBS" ? "☑" : "□";
     let iconMalo = c.estado === "MALO" ? "☑" : "□";
     return `<tr>
       <td>${c.nombre}</td>
       <td class="center">${iconOK} OK &nbsp;&nbsp;&nbsp; ${iconOBS} OBS &nbsp;&nbsp;&nbsp; ${iconMalo} MALO</td>
       <td>${c.obs}</td>
     </tr>`;
  }).join("");

  return `
    <div class="header-title">
      <span>ETIQUETAS NOTEBOOK</span>
      <span>FECHA: ${fecha}</span>
    </div>
    <div class="row">
      <span><span class="bold">MODELO:</span> ${d.model}</span>
      <span><span class="bold">PROCESADOR:</span> ${cpuLimpia}</span>
      <span><span class="bold">GUIA:</span> ${d.guia}</span>
      <span><span class="bold">SERIAL:</span> ${d.serial}</span>
    </div>
    <div class="row checkboxes" style="margin-bottom:20px;">
      <span class="bold">ESTADO GLOBAL:</span> &nbsp;&nbsp;&nbsp;&nbsp;
      □ BUENA &nbsp;&nbsp;&nbsp; ☑ POR REPARAR &nbsp;&nbsp;&nbsp; □ MALA &nbsp;&nbsp;&nbsp; □ REPUESTO
    </div>
    <table>
      <thead>
        <tr>
          <th width="20%">REVISIÓN DE COMPONENTES</th>
          <th width="35%" class="center">ESTADO</th>
          <th width="45%">OBSERVACIÓN</th>
        </tr>
      </thead>
      <tbody class="checkboxes">
        ${trs}
      </tbody>
    </table>
    <div class="row checkboxes bold" style="margin-top:20px;">
      CHECKLIST: &nbsp;&nbsp;&nbsp;&nbsp; □ PRUEBA S/O &nbsp;&nbsp;&nbsp;&nbsp; □ HYDRA &nbsp;&nbsp;&nbsp;&nbsp; □ HARDWARE DEFAULT
    </div>
  `;
}

function imprimirEtiqueta(id) {
  let eq = document.getElementById(id);
  if(!eq) return;
  let d = extraerDatos(eq);

  let rawComps = eq.getAttribute('data-comps');
  let comps = [];
  try { comps = JSON.parse(rawComps); } catch(e) {}

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Etiqueta ${d.serial}</title>${getPrintStyle()}</head><body>${generarHtmlEtiqueta(d, comps)}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 350);
}

function imprimirTodasEtiquetas() {
  let etiquetasHtml = [];

  document.querySelectorAll('details').forEach(eq => {
    if (eq.querySelector('.tag-falla')) {
      let d = extraerDatos(eq);
      let rawComps = eq.getAttribute('data-comps');
      let comps = [];
      try { comps = JSON.parse(rawComps); } catch(e) {}

      etiquetasHtml.push(generarHtmlEtiqueta(d, comps));
    }
  });

  if (etiquetasHtml.length === 0) {
    alert("No hay equipos con fallas en este reporte para imprimir.");
    return;
  }

  let contenido = etiquetasHtml.join('<div class="page-break"></div>');

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Lote de Etiquetas</title>${getPrintStyle()}</head><body>${contenido}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 500);
}
</script>
</head>
<body>
<div class="container">
<div class="header">
  <h1>📋 Inventario Maestro de Equipos</h1>
  <p class="subtitle">Sistema de Control de Entregas y Retiros</p>
  <span class="counter-box">Total Equipos: <span id="total-count">0</span></span>

  <div id="stats-container" class="stats-container"></div>

  <div class="button-group">
    <button class="btn btn-copy"   onclick="copiarFilas()">📋 Copiar Filas (Pegar en Valida)</button>
    <button class="btn btn-export" onclick="exportarValidaCSV()">📥 Descargar CSV</button>
    <button class="btn" style="background:#0f766e;" onclick="exportarJsonFusionado()">🧩 Descargar JSON Fusionado</button>
    <button class="btn" style="background:#334155;" onclick="imprimirTodasEtiquetas()">🖨️ Imprimir Lote de Etiquetas</button>
  </div>
</div>
<div id="main-list">"""

HTML_END = "\n</div></div></body></html>"

# ─────────────────────────────────────────────────────────
#  Helpers de formato para python-docx
# ─────────────────────────────────────────────────────────


# ============================================================================
# 5. EXPORT WORD
# ----------------------------------------------------------------------------
# Helpers de formato docx + generar_informe_word: informe técnico de equipos
# defectuosos. Requiere python-docx (DOCX_AVAILABLE).
#
# CUÁNDO SE CREA EL INFORME: nadie llama acá al escanear ni al guardar un
# equipo. El informe nace en un solo punto del flujo normal —cerrar_lote()—
# y solo si se cumplen las TRES condiciones a la vez:
#   1) la OT es de RETIRO  (una Entrega nunca genera informe; de hecho el panel
#      de "Revisión de Componentes" del formulario está oculto en las entregas),
#   2) al menos un equipo quedó con TIENE_FALLAS=True, o sea con algún
#      componente marcado OBS o MALO, y
#   3) la OT se cierra de verdad (botón CERRAR OT).
# Si falta cualquiera de las tres no hay .docx y tampoco hay error: es el
# comportamiento esperado.
#
# La otra vía es _reconstruir_carpeta_suelta() (sección 8), que sí lo regenera
# al rearmar una carpeta bajada del servidor. OJO con la asimetría: regenerar
# una OT canónica con "RECONSTRUIR DESDE JSONs" rehace el HTML y el FUSIONADO
# pero NO vuelve a escribir el Word.
# ============================================================================
def _set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(space_before * 20)))
    spacing.set(qn("w:after"), str(int(space_after * 20)))
    if line_spacing:
        spacing.set(qn("w:line"), str(int(line_spacing * 240)))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MARCA_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_paragraph_spacing(p, space_before=0, space_after=4)
    return p


def _add_campo(
    doc, etiqueta, valor, etiqueta_bold=True, valor_bold=False, font_size=11
):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=0, space_after=2)
    run_lbl = p.add_run(etiqueta)
    run_lbl.bold = etiqueta_bold
    run_lbl.font.size = Pt(font_size)
    run_lbl.font.name = "Arial"
    run_val = p.add_run(valor)
    run_val.bold = valor_bold
    run_val.font.size = Pt(font_size)
    run_val.font.name = "Arial"
    return p


def _add_titulo_seccion(
    doc, texto, font_size=11, uppercase=True, space_before=10, space_after=4
):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
    run = p.add_run(texto.upper() if uppercase else texto)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(*MARCA_RGB)
    return p


# ─────────────────────────────────────────────────────────
#  Generación de Word (Docx)
# ─────────────────────────────────────────────────────────


def generar_informe_word(cliente, equipos_malos, ruta_cliente, id_lote=""):
    if not DOCX_AVAILABLE:
        messagebox.showwarning(
            "Falta Librería", "No se pudo generar el Word porque falta 'python-docx'."
        )
        return

    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        logo_path = os.path.join(os.getcwd(), "lg1.png")

        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            r_logo = p_logo.add_run()
            r_logo.add_picture(logo_path, width=Cm(6.5))
            _set_paragraph_spacing(p_logo, space_before=0, space_after=4)

        p_titulo = doc.add_paragraph()
        _set_paragraph_spacing(p_titulo, space_before=0, space_after=8)
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run("INFORME TÉCNICO DE RETIRO - RESUMEN DE LOTE")
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(*MARCA_RGB)

        _add_horizontal_rule(doc)

        _add_campo(doc, "CLIENTE:             ", cliente.upper())
        _add_campo(doc, "FECHA DE REVISIÓN:   ", fecha_hoy)
        _add_campo(doc, "EQUIPOS DEFECTUOSOS: ", str(len(equipos_malos)))

        _add_horizontal_rule(doc)
        _add_titulo_seccion(
            doc, "DETALLE DE EQUIPOS Y FALLAS", space_before=8, space_after=8
        )

        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Table Grid"

        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "GUÍA"
        hdr_cells[1].text = "MODELO"
        hdr_cells[2].text = "N° DE SERIE"
        hdr_cells[3].text = "DETALLE DE FALLAS"

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for equipo in equipos_malos:
            guia = equipo.get("GUIA_ID", "S/N")
            modelo = equipo.get("MODELO", "Desconocido")
            serie = equipo.get("SERIAL", "Desconocido")
            obs = equipo.get("OBS", "")

            detalle_comps = equipo.get("DETALLE_COMPONENTES", [])
            fallas_list = []
            for c in detalle_comps:
                if c["estado"] != "OK":
                    texto = f"- {c['nombre']}: {c['estado']}"
                    if c["obs"]:
                        texto += f" ({c['obs']})"
                    fallas_list.append(texto)

            texto_fallas = "\n".join(fallas_list)
            if obs and obs not in ("Sin observaciones", ""):
                texto_fallas += f"\nOBS GENERAL: {obs}"

            row_cells = tabla.add_row().cells
            row_cells[0].text = str(guia)
            row_cells[1].text = modelo
            row_cells[2].text = serie
            row_cells[3].text = texto_fallas

        _add_titulo_seccion(doc, "CONCLUSIÓN:", space_before=15, space_after=4)
        p_conclu = doc.add_paragraph()
        p_conclu.add_run(
            "Los equipos listados en la tabla superior presentan fallas físicas o de hardware. "
            "Se requiere revisión detallada, cambio de componentes afectados o envío a bodega "
            "a la espera de repuestos o reparación."
        )

        for _ in range(3):
            doc.add_paragraph()

        _add_horizontal_rule(doc)
        p_firma = doc.add_paragraph()
        run_f1 = p_firma.add_run("TÉCNICO ENCARGADO: ")
        run_f1.bold = True
        run_f2 = p_firma.add_run("TOMAS GAC\n")
        run_f2.bold = True
        run_r1 = p_firma.add_run("RUT: ")
        run_r1.bold = True
        p_firma.add_run("21.790.634-2")

        fecha_str = datetime.now().strftime("%Y%m%d")
        id_part = f"_{id_lote}" if id_lote and id_lote != "SN" else ""
        nombre_archivo = f"Informe_Tecnico_{cliente}{id_part}_{fecha_str}.docx"
        ruta_guardado = _ruta_no_pisar(os.path.join(ruta_cliente, nombre_archivo))
        doc.save(ruta_guardado)

    except Exception as e:
        print(f"Error generando Word: {e}")
        messagebox.showerror("Error al generar Word", str(e))


# ============================================================================
# 6. SCANNER (ejecución y parseo)
# ----------------------------------------------------------------------------
# Corre PS_SCRIPT (sección 3) con reintentos/timeout y devuelve el dict `data`
# usado por todo el resto (formulario, guardado, HTML).
# ============================================================================

# ─────────────────────────────────────────────────────────
#  Filas de hardware del reporte HTML — ÚNICA fuente
#
#  El JS del reporte (extraerDatos) lee la RAM, los discos y la red de estas
#  filas, no de los JSON: los th 'Módulo RAM', 'Discos Internos' y
#  'Serie Disco Duro' son el contrato del que salen "Copiar Filas" y el CSV
#  para Valida. Por eso las arma un solo lugar, usado tanto por el escaneo
#  como por la reconstrucción: si divergen, el reporte reconstruido pierde
#  specs y exporta filas incompletas.
# ─────────────────────────────────────────────────────────
def render_filas_ram(ram_list):
    """Filas 'Módulo RAM' desde la lista cruda '8GB|3200|PC4|Fabricante|PN'."""
    if isinstance(ram_list, str):
        ram_list = [ram_list]
    ram_list = [r for r in (ram_list or []) if r]
    return (
        "".join(f"<tr><th>Módulo RAM</th><td>{r}</td></tr>" for r in ram_list)
        or "<tr><th>RAM</th><td>No detectada</td></tr>"
    )


def render_filas_discos(discos):
    """
    Filas de discos desde [{'desc':…, 'serial':…}]. Devuelve (html, seriales).
    El serial se normaliza igual que en el escaneo (sin símbolos, 16 últimos)
    para que el mismo disco dé la misma cadena vino de donde vino.
    """
    if isinstance(discos, dict):
        discos = [discos]
    html_discos = ""
    seriales = []
    for d in discos or []:
        desc = (d.get("desc") or "Desconocido").strip()
        serie = re.sub(r"[^a-zA-Z0-9]", "", str(d.get("serial", "NA")))[-16:]
        seriales.append(serie)
        html_discos += (
            f"<tr><th>Discos Internos</th><td>{desc}</td></tr>"
            f"<tr><th>Serie Disco Duro</th><td class='serial-tag'>{serie}</td></tr>"
        )
    return html_discos, seriales


def render_filas_red(net_list):
    """Filas 'Red Activa' desde la lista cruda de adaptadores."""
    if isinstance(net_list, str):
        net_list = [net_list]
    net_list = [n for n in (net_list or []) if n]
    return (
        "".join(
            f"<tr><td>Red Activa</td><td><span class='net-tag'>{n}</span></td></tr>"
            for n in net_list
        )
        or "<tr><td colspan='2'>Sin red activa</td></tr>"
    )


def get_inventory_fast(max_retries=3, timeout_per_attempt=45):
    last_error = None

    for attempt in range(1, max_retries + 1):
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".ps1", delete=False, encoding="utf-8"
            ) as tmp:
                tmp.write(PS_SCRIPT)
                tmp_path = tmp.name

            si = subprocess.STARTUPINFO()
            si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            si.wShowWindow = 0  # SW_HIDE

            cmd = [
                "powershell",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                tmp_path,
            ]

            result = subprocess.check_output(
                cmd,
                text=True,
                startupinfo=si,
                stderr=subprocess.PIPE,
                timeout=timeout_per_attempt,
            )

            json_match = re.search(r"\{.*\}", result, flags=re.DOTALL)
            if not json_match:
                raise ValueError(
                    f"Salida de PowerShell no contiene JSON válido "
                    f"(intento {attempt}/{max_retries}). "
                    f"Salida recibida: {result[:200]!r}"
                )

            data = json.loads(json_match.group(0))

            try:
                os.unlink(tmp_path)
                tmp_path = None
            except Exception:
                pass

            raw_model = data.get("model", "Desconocido")
            cln_model = re.sub(
                r'(?i)\b(11\.6|12\.5|13\.3|14|15\.6|16|17\.3)\s*(inch|")?\b',
                "",
                raw_model,
            )
            cln_model = re.sub(r"(?i)\bnotebook\s*pc\b", "", cln_model)
            cln_model = re.sub(r"\s+", " ", cln_model).strip()

            ram_list = data.get("ram_rows", [])
            if isinstance(ram_list, str):
                ram_list = [ram_list]
            ram_html = render_filas_ram(ram_list)

            fabricantes_ram = []
            for r in ram_list:
                partes = r.split("|")
                if len(partes) >= 4:
                    fabricantes_ram.append(partes[3].strip())
            fabricantes_unicos = list(set(fabricantes_ram))
            fabricante_final = (
                " / ".join(fabricantes_unicos) if fabricantes_unicos else "Desconocida"
            )

            partnumbers_ram = []
            for r in ram_list:
                partes = r.split("|")
                if len(partes) >= 5:
                    partnumbers_ram.append(partes[4].strip())
            partnumber_final = (
                " / ".join(set(partnumbers_ram)) if partnumbers_ram else "Desconocida"
            )

            disks = data.get("disks_data", [])
            if isinstance(disks, dict):
                disks = [disks]
            disk_html, disk_serials = render_filas_discos(disks)

            net_list = data.get("net_rows", [])
            if isinstance(net_list, str):
                net_list = [net_list]
            net_html = render_filas_red(net_list)

            return {
                "fecha": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "host": socket.gethostname(),
                "model": cln_model,
                "serial": data.get("serial", "SN_DESCONOCIDO").strip(),
                "key": data.get("win_key", "N/A"),
                "cpu": data.get("cpu", "Desconocido"),
                "ram_html": ram_html,
                "ram_raw": ram_list,
                "disk_html": disk_html,
                # Crudos además del HTML: si el HTML se pierde o cambia de
                # formato, la reconstrucción rearma las filas desde acá sin
                # perder la descripción del disco (que el serial solo no trae).
                "disk_raw": disks,
                "disk_serials": disk_serials,
                "net_html": net_html,
                "net_raw": net_list,
                # Se conserva el nombre viejo: los JSON ya guardados lo usan y
                # _reconstruir_entry_desde_json sigue leyéndolo primero.
                "net_rows": net_html,
                "Manufacturer": fabricante_final,
                "Partnumber": partnumber_final,
                "Battery": data.get("battery", "No detectada"),
            }

        except subprocess.TimeoutExpired:
            last_error = (
                f"Timeout tras {timeout_per_attempt}s (intento {attempt}/{max_retries})"
            )
            print(f"⚠️  {last_error}")
        except json.JSONDecodeError as e:
            last_error = f"JSON inválido (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        except Exception as e:
            last_error = f"Error inesperado (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

        if attempt < max_retries:
            time.sleep(2)

    print(f"❌ Escaneo falló después de {max_retries} intentos. Último: {last_error}")
    return None


# ============================================================================
# 7. PERSISTENCIA — MODELO DE LOTES POR OT / GUÍA
# ----------------------------------------------------------------------------
# Un LOTE es (cliente, movimiento, número de documento) y vive en su propia
# carpeta desde el PRIMER escaneo, no al final:
#
#     Reportes_Guardados/ENTREGA/<CLIENTE>/OT-553/      (entregas: manda la OT)
#     Reportes_Guardados/RETIROS/<CLIENTE>/GUIA-8891/   (retiros: manda la guía)
#
# El movimiento lo define el lote, no cada equipo. Los JSON individuales se
# escriben directo en esa carpeta; el HTML y el JSON FUSIONADO son DERIVADOS:
# se regeneran completos desde los JSON en cada guardado. Por eso una OT puede
# trabajarse durante varios días y el fusionado siempre queda con TODO adentro,
# listo para subir al sistema.
#
# Una OT está ABIERTA mientras no tenga el marcador '.cerrado'. Ese marcador es
# lo que permite ignorar las miles de carpetas ya finalizadas sin abrir ni un
# archivo: basta un stat por carpeta.
#
# RESPALDO: crear_backup() comprime ENTREGA/ y RETIROS/ enteras en un ZIP con
# fecha y hora bajo _BACKUPS/. Se dispara solo en los dos momentos críticos —
# al cerrar una OT y antes de subir al NAS — y va fuera de ENTREGA/RETIROS para
# no aparecer como un cliente/OT falso en los recorridos.
# ============================================================================

# Configuración por tipo de movimiento. Cada uno tiene su carpeta raíz, el
# rótulo del documento y el prefijo de la carpeta del lote.
_MOV_CFG = {
    "Entrega": {
        "tipo_doc": "Entregas",
        "grupo": "ENTREGA",
        "prefijo": "OT",  # carpeta del lote: OT-553
        "etiqueta_id": "Orden",  # rótulo del ID en el nombre del HTML
        "etiqueta_campo": "N° de Orden de Trabajo (OT)",
        # El otro documento del lote, el que NO da nombre a la carpeta.
        "etiqueta_campo_sec": "N° de Guía de Despacho",
    },
    "Retiro": {
        "tipo_doc": "Retiros",
        "grupo": "RETIROS",
        "prefijo": "GUIA",
        "etiqueta_id": "Guia",
        "etiqueta_campo": "N° de Guía de Retiro",
        "etiqueta_campo_sec": "N° de Orden de Trabajo (OT)",
    },
}

# Marcador de lote finalizado. Su sola presencia significa "aquí no hay nada
# que procesar": los recorridos automáticos se saltan la carpeta sin leerla.
_MARCADOR_CERRADO = ".cerrado"

# Marcador local (no se sube al NAS) que indica que la carpeta YA fue enviada a
# la red. Va por OT, no por cliente: así agregar un equipo a una OT no obliga a
# re-subir las miles de carpetas del cliente.
_MARCADOR_ENVIADO = ".enviado"

# Puntero al lote activo (el que recibe los equipos escaneados).
_ARCHIVO_LOTE_ACTIVO = ".lote_activo.json"

# Metadatos guardados DENTRO de la carpeta del lote. Existen porque la ruta
# codifica un solo número (el que nombra la carpeta) y el segundo documento
# —la guía de una entrega, la OT de un retiro— no se puede deducir de ella.
# Sin extensión .json a propósito: contar_equipos() cuenta los .json sueltos
# de la carpeta y este archivo no es un equipo.
_ARCHIVO_META_LOTE = ".lote"

_RESERVADOS_WIN = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def ruta_reportes():
    return os.path.join(os.getcwd(), "Reportes_Guardados")


def normalizar_numero(valor, mov="Entrega"):
    """
    Sanea el número de OT/guía para usarlo como nombre de carpeta: mayúsculas,
    sin caracteres inválidos de ruta y sin el prefijo si el usuario lo escribió
    ("OT 553", "ot-553" y "553" dan todos "553"). Devuelve "" si no queda nada
    usable o si cae en un nombre reservado de Windows (CON, AUX, COM1…).
    """
    n = str(valor).strip().upper()
    n = re.sub(r"^(OT|GUIA|GUÍA)[\s\-_.:#]*", "", n)
    n = re.sub(r"[^A-Z0-9\-_]", "-", n)
    n = re.sub(r"-{2,}", "-", n).strip("-_")[:30]
    if not n or n in _RESERVADOS_WIN:
        return ""
    return n


def carpeta_lote(mov, numero):
    return f"{_MOV_CFG[mov]['prefijo']}-{numero}"


def ruta_lote(mov, cliente, numero):
    return os.path.join(
        ruta_reportes(), _MOV_CFG[mov]["grupo"], cliente, carpeta_lote(mov, numero)
    )


def hacer_lote(mov, cliente, numero, ruta=None, numero_sec=""):
    """
    Descriptor de lote que circula por toda la app.

    'numero' es el documento que identifica el lote y da nombre a la carpeta:
    la OT en una entrega, la guía en un retiro. 'numero_sec' es el OTRO
    documento, el que la ruta no codifica; puede ir vacío si aún no se conoce.
    """
    return {
        "mov": mov,
        "cliente": cliente,
        "numero": numero,
        "numero_sec": (numero_sec or "").strip(),
        "ruta": ruta or ruta_lote(mov, cliente, numero),
    }


def numeros_lote(lote):
    """
    (OT, guía) del lote como par explícito. Cuál de los dos nombra la carpeta
    depende del movimiento: una entrega se archiva por OT y un retiro por guía,
    así que el otro sale de 'numero_sec' y puede no estar cargado todavía.
    """
    otro = (lote.get("numero_sec") or "").strip()
    if lote["mov"] == "Entrega":
        return lote["numero"], otro or "Sin Guía"
    return otro or "Sin Orden", lote["numero"]


def guardar_meta_lote(lote):
    """Deja en la carpeta del lote el número que la ruta no puede codificar."""
    try:
        os.makedirs(lote["ruta"], exist_ok=True)
        ruta = os.path.join(lote["ruta"], _ARCHIVO_META_LOTE)
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump({"numero_sec": lote.get("numero_sec", "")}, f, ensure_ascii=False)
    except Exception:
        pass


def leer_meta_lote(ruta):
    """Metadatos del lote guardados en su carpeta. {} si no hay o no se pudo."""
    try:
        with open(os.path.join(ruta, _ARCHIVO_META_LOTE), "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _numero_sec_en_disco(ruta):
    return (leer_meta_lote(ruta).get("numero_sec") or "").strip()


def etiqueta_lote(lote):
    """Rótulo corto para la UI: 'ACME · OT-553 · ENTREGAS'."""
    cfg = _MOV_CFG[lote["mov"]]
    return (
        f"{lote['cliente'].replace('_', ' ')} · "
        f"{carpeta_lote(lote['mov'], lote['numero'])} · {cfg['tipo_doc'].upper()}"
    )


# ─────────────────────────────────────────────────────────
#  Estado de un lote: abierto / cerrado
# ─────────────────────────────────────────────────────────
def esta_cerrado(ruta):
    return os.path.exists(os.path.join(ruta, _MARCADOR_CERRADO))


def leer_cerrado(ruta):
    """Metadatos del cierre (cliente, número, nº de equipos, fecha). {} si no."""
    try:
        with open(os.path.join(ruta, _MARCADOR_CERRADO), "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def marcar_cerrado(lote, equipos):
    """
    Escribe '.cerrado' con un resumen. Lleva metadatos a propósito: permite
    saber qué hay en una OT (cliente, número, cuántos equipos) sin abrir
    ninguno de sus JSON.
    """
    meta = {
        "cliente": lote["cliente"],
        "movimiento": lote["mov"],
        "numero": lote["numero"],
        "numero_sec": lote.get("numero_sec", ""),
        "equipos": equipos,
        "cerrado": datetime.now().isoformat(timespec="seconds"),
    }
    with open(
        os.path.join(lote["ruta"], _MARCADOR_CERRADO), "w", encoding="utf-8"
    ) as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def reabrir_lote(lote):
    """
    Vuelve a poner en curso una OT ya cerrada: borra '.cerrado' y también
    '.enviado', para que lo que se agregue se re-sincronice con el NAS.
    """
    for marcador in (_MARCADOR_CERRADO, _MARCADOR_ENVIADO):
        ruta = os.path.join(lote["ruta"], marcador)
        if os.path.exists(ruta):
            try:
                os.remove(ruta)
            except Exception:
                pass


def contar_equipos(ruta):
    """Equipos de un lote = sus JSON individuales (el FUSIONADO no cuenta)."""
    try:
        return sum(
            1
            for a in os.listdir(ruta)
            if a.endswith(".json") and not a.endswith("_FUSIONADO.json")
        )
    except Exception:
        return 0


def listar_lotes_abiertos():
    """
    Lotes SIN '.cerrado', ordenados por actividad reciente. De las OT cerradas
    —que son la mayoría— no abre ningún archivo: solo hace stat de la carpeta.
    De las abiertas lee además su '.lote', que trae el número secundario.
    """
    base = ruta_reportes()
    abiertos = []
    for mov, cfg in _MOV_CFG.items():
        raiz = os.path.join(base, cfg["grupo"])
        if not os.path.isdir(raiz):
            continue
        pref = cfg["prefijo"] + "-"
        try:
            with os.scandir(raiz) as clientes:
                for c in clientes:
                    if not c.is_dir():
                        continue
                    with os.scandir(c.path) as lotes:
                        for l in lotes:
                            if not l.is_dir() or esta_cerrado(l.path):
                                continue
                            numero = (
                                l.name[len(pref):] if l.name.startswith(pref) else l.name
                            )
                            lote = hacer_lote(
                                mov,
                                c.name,
                                numero,
                                ruta=l.path,
                                numero_sec=_numero_sec_en_disco(l.path),
                            )
                            lote["equipos"] = contar_equipos(l.path)
                            lote["modificado"] = os.path.getmtime(l.path)
                            abiertos.append(lote)
        except Exception:
            continue
    abiertos.sort(key=lambda x: x["modificado"], reverse=True)
    return abiertos


# ─────────────────────────────────────────────────────────
#  Lote activo (el que recibe lo que se escanea)
# ─────────────────────────────────────────────────────────
def guardar_lote_activo(lote):
    os.makedirs(ruta_reportes(), exist_ok=True)
    with open(
        os.path.join(ruta_reportes(), _ARCHIVO_LOTE_ACTIVO), "w", encoding="utf-8"
    ) as f:
        json.dump(
            {k: lote.get(k, "") for k in ("mov", "cliente", "numero", "numero_sec")},
            f,
            ensure_ascii=False,
        )


def cargar_lote_activo():
    """
    Lote activo, o None si no hay / desapareció la carpeta / ya fue cerrado
    (por ejemplo desde otro equipo o a mano en el explorador).
    """
    try:
        with open(
            os.path.join(ruta_reportes(), _ARCHIVO_LOTE_ACTIVO), "r", encoding="utf-8"
        ) as f:
            cfg = json.load(f)
        lote = hacer_lote(
            cfg["mov"],
            cfg["cliente"],
            cfg["numero"],
            numero_sec=cfg.get("numero_sec", ""),
        )
    except Exception:
        return None
    if not os.path.isdir(lote["ruta"]) or esta_cerrado(lote["ruta"]):
        return None
    # Punteros escritos por versiones anteriores no traen el número secundario:
    # se recupera del '.lote' de la carpeta en vez de darlo por perdido.
    if not lote["numero_sec"]:
        lote["numero_sec"] = _numero_sec_en_disco(lote["ruta"])
    return lote


def limpiar_lote_activo():
    ruta = os.path.join(ruta_reportes(), _ARCHIVO_LOTE_ACTIVO)
    try:
        if os.path.exists(ruta):
            os.remove(ruta)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────
#  Archivos del lote. El JSON individual es la ÚNICA fuente de verdad;
#  el HTML y el FUSIONADO se rehacen desde ellos en cada guardado.
# ─────────────────────────────────────────────────────────
def _nombre_base(lote):
    return (
        f"Reporte_{_MOV_CFG[lote['mov']]['tipo_doc']}_"
        f"{lote['cliente']}_{lote['numero']}"
    )


def ruta_json_equipo(lote, serial):
    # Sin fecha en el nombre a propósito: si el equipo se re-escanea otro día
    # debe PISAR su registro anterior, no crear un duplicado en el fusionado.
    serial_clean = re.sub(r"[^a-zA-Z0-9]", "", str(serial)) or "SINSERIE"
    return os.path.join(lote["ruta"], f"{_nombre_base(lote)}_{serial_clean}.json")


def ruta_fusionado(lote):
    return os.path.join(lote["ruta"], f"{_nombre_base(lote)}_FUSIONADO.json")


def ruta_html_lote(lote):
    cfg = _MOV_CFG[lote["mov"]]
    cliente_legible = lote["cliente"].replace("_", " ")
    return os.path.join(
        lote["ruta"],
        f"{cfg['tipo_doc']} - {cliente_legible} - "
        f"{cfg['etiqueta_id']} {lote['numero']}.html",
    )


def _clave_orden(jdata):
    """Timestamp ordenable del escaneo (los JSON viejos no lo traen)."""
    ts = jdata.get("ESCANEADO", "")
    if ts:
        return ts
    fecha = (jdata.get("DATA") or {}).get("fecha", "")
    try:
        return datetime.strptime(fecha, "%d/%m/%Y %H:%M").isoformat(timespec="seconds")
    except Exception:
        return ""


def leer_equipos_lote_con_ruta(lote):
    """
    (ruta, jdata) de cada equipo del lote, del más reciente al más antiguo.

    La ruta real importa cuando hay que REESCRIBIR el JSON: los registros
    viejos pueden tener un nombre que ruta_json_equipo() ya no genera, y
    recalcularlo crearía un archivo nuevo en vez de actualizar el que existe.
    """
    equipos = []
    try:
        archivos = os.listdir(lote["ruta"])
    except Exception:
        return equipos
    for archivo in archivos:
        if not archivo.endswith(".json") or archivo.endswith("_FUSIONADO.json"):
            continue
        ruta = os.path.join(lote["ruta"], archivo)
        try:
            with open(ruta, "r", encoding="utf-8") as f:
                jdata = json.load(f)
            if isinstance(jdata, dict) and "SERIAL" in jdata:
                equipos.append((ruta, jdata))
        except Exception:
            continue
    equipos.sort(key=lambda par: _clave_orden(par[1]), reverse=True)
    return equipos


def leer_equipos_lote(lote):
    """JSON individuales del lote, del más reciente al más antiguo."""
    return [jdata for _, jdata in leer_equipos_lote_con_ruta(lote)]


def regenerar_lote(lote):
    """
    Rehace el HTML y el JSON FUSIONADO del lote a partir de sus JSON.
    Al ser derivados y no incrementales, el fusionado siempre sale completo
    (aunque la OT se haya trabajado en varios días) y sin duplicados.
    Devuelve el total de equipos.

    NO toca el informe Word de fallas: ese lo escribe cerrar_lote() y solo al
    cerrar. Correr esto sobre una OT de retiro con equipos malos deja el .docx
    como estaba (ver sección 5).
    """
    equipos = leer_equipos_lote(lote)
    cfg = _MOV_CFG[lote["mov"]]
    cliente_legible = lote["cliente"].replace("_", " ")

    contenido = HTML_START.replace(
        "Inventario Maestro de Equipos",
        f"{cfg['tipo_doc']} - {cliente_legible} · "
        f"{cfg['etiqueta_id']} {lote['numero']}",
    ).replace(
        "<title>Inventario Maestro</title>",
        f"<title>{cfg['tipo_doc']} - {lote['cliente']}</title>",
    )

    entries = []
    for jdata in equipos:
        try:
            entries.append(_reconstruir_entry_desde_json(jdata))
        except Exception as e:
            print(f"No se pudo rearmar {jdata.get('SERIAL', '?')}: {e}")
    contenido += "\n".join(entries) + HTML_END
    contenido = fijar_total_count(contenido, len(entries))

    with open(ruta_html_lote(lote), "w", encoding="utf-8") as f:
        f.write(contenido)
    with open(ruta_fusionado(lote), "w", encoding="utf-8") as f:
        json.dump(
            [_resumen_fusionado(eq) for eq in equipos], f, ensure_ascii=False, indent=4
        )
    return len(equipos)


def marcar_pendiente_envio(lote):
    """Contenido nuevo en la OT ⇒ vuelve a estar pendiente de subir al NAS."""
    ruta = os.path.join(lote["ruta"], _MARCADOR_ENVIADO)
    if os.path.exists(ruta):
        try:
            os.remove(ruta)
        except Exception:
            pass


def guardar_registro_en_lote(lote, registro):
    """Escribe el JSON del equipo y regenera los derivados. Devuelve el total."""
    os.makedirs(lote["ruta"], exist_ok=True)
    guardar_meta_lote(lote)
    with open(ruta_json_equipo(lote, registro["SERIAL"]), "w", encoding="utf-8") as f:
        json.dump(registro, f, ensure_ascii=False, indent=4)
    marcar_pendiente_envio(lote)
    return regenerar_lote(lote)


def pendientes_legacy():
    """
    Equipos escaneados con el modelo anterior (data_*.json sueltos en la raíz
    de Reportes_Guardados, sin carpeta de OT). Devuelve {mov: [rutas]}.
    """
    base = ruta_reportes()
    encontrados = {"Entrega": [], "Retiro": []}
    if os.path.isdir(base):
        for archivo in sorted(os.listdir(base)):
            if archivo.startswith("data_") and archivo.endswith(".json"):
                mov = "Retiro" if archivo.startswith("data_Retiro_") else "Entrega"
                encontrados[mov].append(os.path.join(base, archivo))
    return encontrados


def migrar_legacy_a_lote(rutas, lote):
    """
    Adopta los data_*.json sueltos dentro de una OT: los reescribe con el
    cliente/número del lote y regenera HTML y FUSIONADO. Devuelve cuántos entraron.
    """
    os.makedirs(lote["ruta"], exist_ok=True)
    guardar_meta_lote(lote)
    num_ot, num_guia = numeros_lote(lote)
    migrados = 0
    for ruta in rutas:
        try:
            with open(ruta, "r", encoding="utf-8") as f:
                jdata = json.load(f)
            jdata["CLIENTE"] = lote["cliente"]
            jdata["TIPO_MOVIMIENTO"] = lote["mov"]
            jdata["GUIA_ID"] = num_guia
            jdata["OT_ID"] = num_ot
            destino = ruta_json_equipo(lote, jdata.get("SERIAL", ""))
            with open(destino, "w", encoding="utf-8") as f:
                json.dump(jdata, f, ensure_ascii=False, indent=4)
            os.remove(ruta)
            migrados += 1
        except Exception as e:
            print(f"No se pudo migrar {os.path.basename(ruta)}: {e}")

    # El HTML acumulado del modelo viejo ya no aplica: el de la OT se regenera
    # desde los JSON. Se borra solo el del movimiento migrado.
    legacy_html = os.path.join(
        ruta_reportes(), f"Reporte_{_MOV_CFG[lote['mov']]['tipo_doc']}.html"
    )
    if os.path.exists(legacy_html):
        try:
            os.remove(legacy_html)
        except Exception:
            pass

    marcar_pendiente_envio(lote)
    regenerar_lote(lote)
    return migrados


def buscar_serial_en_abiertos(serial, excluir_ruta=None):
    """
    Busca el serial en los demás lotes ABIERTOS. Con varias OT en curso el
    mismo equipo podría quedar registrado en dos, y eso hay que avisarlo.
    """
    for otro in listar_lotes_abiertos():
        if excluir_ruta and os.path.normcase(otro["ruta"]) == os.path.normcase(
            excluir_ruta
        ):
            continue
        if os.path.exists(ruta_json_equipo(otro, serial)):
            return otro
    return None


# ============================================================================
# PLANTILLA COMPARTIDA DE LA ENTRADA <details> DEL REPORTE HTML
# ----------------------------------------------------------------------------
# Constructor único, alimentado SIEMPRE desde el JSON del equipo vía
# _reconstruir_entry_desde_json: el HTML del lote se regenera entero en cada
# guardado, así que el JSON individual es la única fuente de verdad.
#
# Por eso la reconstrucción debe reponer todo lo que el registro contiene
# (batería y la Key de Office, que vive en SN_APP): si algo no se rearma aquí,
# desaparece del reporte en la siguiente regeneración.
# ============================================================================
def construir_entry_html(
    *,
    safe_id,
    model,
    serial,
    fecha,
    key,
    cpu,
    mov,
    guia_final,
    num_ot,
    obs_final,
    ram_html,
    disk_html,
    net_rows,
    comps_json,
    fusion_json="",
    tag_falla_html="",
    falla_html="",
    btn_imprimir="",
    office_html="",
    bateria_html="",
    transf_html="",
):
    return f"""<details id="{safe_id}" data-comps='{comps_json}' data-fusion='{fusion_json}' open>
<summary>
  <span class="model-name">{model}</span>
  <span style="font-size:12px;opacity:.85">S/N: {serial} &nbsp;|&nbsp; 📅 {fecha} {tag_falla_html}</span>
</summary>
<div class="content">
  <div class="content-header">
    {btn_imprimir}
    <button class="btn btn-delete" onclick="eliminarEquipo('{safe_id}')">🗑️ Quitar del Reporte</button>
  </div>
  <table>
    <tr class="logistics-row"><th>Tipo Movimiento</th><td><b>{mov}</b></td></tr>
    <tr class="logistics-row"><th>N° de Guía</th><td><b>{guia_final}</b></td></tr>
    <tr class="logistics-row"><th>N° de OT</th><td><b>{num_ot}</b></td></tr>
    {falla_html}
    <tr class="obs-row"><th>⚠️ Obs. General</th><td>{obs_final}</td></tr>
    {office_html}
    <tr><th>Número de Serie</th><td class="serial-tag">{serial}</td></tr>
    <tr><th>Licencia Windows (OA3)</th><td class="win-key">{key}</td></tr>
    <tr><th>Procesador</th><td>{cpu}</td></tr>{bateria_html}
    {ram_html}{disk_html}{transf_html}{net_rows}
  </table>
</div></details>"""


def fijar_total_count(content, total):
    """Reescribe el contador <span id="total-count"> del reporte HTML."""
    return re.sub(
        r'<span id="total-count">\d+</span>',
        f'<span id="total-count">{total}</span>',
        content,
    )


def guardar_equipo_general(
    lote, data, obs, tiene_office, office_key, version_office, detalle_componentes
):
    """
    Registra un equipo en el lote (OT/guía) activo: escribe su JSON individual
    y regenera el HTML y el FUSIONADO de la OT. El movimiento y el número de
    documento los aporta el lote, no el equipo. Devuelve el total del lote.
    """
    tiene_fallas = any(c["estado"] != "OK" for c in detalle_componentes)
    num_ot, num_guia = numeros_lote(lote)
    sn_hdd = "_".join(data.get("disk_serials", [])) or "PENDIENTE"
    sn_app = (
        office_key.strip()
        if (tiene_office and office_key and office_key.strip())
        else "PENDIENTE"
    )

    registro = {
        "CLIENTE": lote["cliente"],
        "SERIAL": data["serial"],
        "SN_Win": data["key"],
        "SN_HDD": sn_hdd,
        "SN_Transf": "PENDIENTE",
        "SN_APP": sn_app,
        "MODELO": data["model"],
        "TIPO_MOVIMIENTO": lote["mov"],
        "GUIA_ID": num_guia,
        "OT_ID": num_ot,
        "OBS": obs or "Sin observaciones",
        "DETALLE_COMPONENTES": detalle_componentes,
        "TIENE_FALLAS": tiene_fallas,
        "DATA": data,
        "OFFICE": f"Office {version_office}" if tiene_office else "No",
        # Timestamp ordenable: el HTML se rearma con el más reciente arriba.
        "ESCANEADO": datetime.now().isoformat(timespec="seconds"),
    }
    return guardar_registro_en_lote(lote, registro)


def accion_agregar_lote(
    ventana, lote, data, obs, tiene_office, office_key, version_office, comps_data
):
    detalle_componentes = []
    for c in comps_data:
        est = c["estado"].get()
        ob = c["obs"].get().strip()
        detalle_componentes.append({"nombre": c["nombre"], "estado": est, "obs": ob})

    # Con varias OT en curso el mismo equipo puede acabar registrado en dos.
    # No se bloquea (a veces es legítimo re-escanear), pero se avisa.
    otro = buscar_serial_en_abiertos(data["serial"], excluir_ruta=lote["ruta"])
    if otro and not messagebox.askyesno(
        "Equipo en otra OT",
        f"El serial {data['serial']} ya está registrado en:\n"
        f"{etiqueta_lote(otro)}\n\n"
        f"¿Agregarlo igual a {etiqueta_lote(lote)}?",
    ):
        return

    total = guardar_equipo_general(
        lote, data, obs, tiene_office, office_key, version_office, detalle_componentes
    )

    messagebox.showinfo(
        "✅ Guardado",
        f"Equipo agregado a {etiqueta_lote(lote)}.\n\n"
        f"La OT lleva {total} equipo(s) y su JSON FUSIONADO ya está actualizado.",
    )
    mostrar_menu_principal(ventana)


def _ruta_no_pisar(ruta):
    """
    Devuelve una ruta que NO exista aún: si el archivo ya está, agrega
    ' (2)', ' (3)'… antes de la extensión. Evita sobrescribir un HTML,
    FUSIONADO o Word de un lote anterior del mismo cliente/día.
    """
    if not os.path.exists(ruta):
        return ruta
    raiz, ext = os.path.splitext(ruta)
    n = 2
    while os.path.exists(f"{raiz} ({n}){ext}"):
        n += 1
    return f"{raiz} ({n}){ext}"


def crear_backup(motivo=""):
    """
    Comprime ENTREGA/ y RETIROS/ (con todas sus OT) en un ZIP con fecha y hora,
    dentro de Reportes_Guardados/_BACKUPS/.

    El ZIP va FUERA de ENTREGA/RETIROS a propósito: esas carpetas se recorren
    para listar lotes, clientes y pendientes de envío, así que un backup adentro
    aparecería como un cliente/OT falso y se subiría al NAS.

    Devuelve la ruta del ZIP creado, o None si no había nada que respaldar.
    Nunca lanza: un fallo de backup no debe cortar el cierre ni el envío.
    """
    try:
        base = ruta_reportes()
        origenes = [
            (grupo, os.path.join(base, grupo))
            for grupo in ("ENTREGA", "RETIROS")
            if os.path.isdir(os.path.join(base, grupo))
        ]
        if not origenes:
            return None

        ruta_backups = os.path.join(base, "_BACKUPS")
        os.makedirs(ruta_backups, exist_ok=True)

        sello = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        sufijo = f"_{motivo}" if motivo else ""
        destino = _ruta_no_pisar(
            os.path.join(ruta_backups, f"Backup_{sello}{sufijo}.zip")
        )

        archivos = 0
        with zipfile.ZipFile(destino, "w", zipfile.ZIP_DEFLATED) as zf:
            for grupo, ruta_grupo in origenes:
                for carpeta, _, nombres in os.walk(ruta_grupo):
                    for nombre in nombres:
                        completo = os.path.join(carpeta, nombre)
                        # Dentro del ZIP se respeta ENTREGA/<CLIENTE>/OT-553/...
                        interno = os.path.join(
                            grupo, os.path.relpath(completo, ruta_grupo)
                        )
                        zf.write(completo, interno)
                        archivos += 1

        if archivos == 0:  # nada dentro: no dejamos un ZIP vacío
            os.remove(destino)
            return None
        return destino
    except Exception as e:
        print(f"Error generando backup: {e}")
        return None


def _resumen_fusionado(eq):
    """Resumen logístico de un equipo para el JSON FUSIONADO del lote."""
    sn_app = eq.get("SN_APP", "PENDIENTE")
    obs = eq.get("OBS", "")
    return {
        "SERIAL": eq.get("SERIAL", ""),
        "SN_Win": eq.get("SN_Win", ""),
        "SN_HDD": eq.get("SN_HDD", ""),
        "SN_Transf": eq.get("SN_Transf", "PENDIENTE"),
        "SN_APP": None if sn_app in ("PENDIENTE", "", None) else sn_app,
        "OBS": None if obs in ("Sin observaciones", "", None) else obs,
    }


def cerrar_lote(lote):
    """
    Finaliza una OT/guía: regenera el fusionado con TODO lo escaneado (aunque
    se haya trabajado en varios días), emite el informe Word si hay equipos con
    fallas y escribe el marcador '.cerrado'.

    Este es el ÚNICO punto del flujo normal que produce el .docx, y solo para
    los RETIROS: en una entrega equipos_malos queda vacío a propósito, porque
    el formulario ni siquiera muestra el panel de componentes. Cerrar dos veces
    (reabrir, sumar equipos y volver a cerrar) no pisa el informe anterior:
    _ruta_no_pisar le agrega ' (2)', así que el archivo con el nombre original
    conserva los datos del primer cierre.

    No mueve ni renombra nada: los equipos viven en la carpeta de la OT desde
    el primer escaneo. Devuelve (total_equipos, equipos_malos).
    """
    # Hay que preguntarlo ANTES de marcar: cargar_lote_activo() descarta los
    # lotes cerrados, así que después del marcador siempre daría "no es activo"
    # y el puntero quedaría apuntando a una OT ya cerrada.
    era_activo = _es_activo(lote)

    total = regenerar_lote(lote)
    equipos = leer_equipos_lote(lote)

    equipos_malos = (
        [e for e in equipos if e.get("TIENE_FALLAS")]
        if lote["mov"] == "Retiro"
        else []
    )
    if equipos_malos:
        generar_informe_word(
            lote["cliente"], equipos_malos, lote["ruta"], lote["numero"]
        )

    marcar_cerrado(lote, total)
    if era_activo:
        limpiar_lote_activo()
    return total, equipos_malos


def _es_activo(lote):
    activo = cargar_lote_activo()
    return bool(activo) and os.path.normcase(activo["ruta"]) == os.path.normcase(
        lote["ruta"]
    )


def _clientes_existentes():
    """Clientes que ya tienen alguna OT/guía (carpetas dentro de ENTREGA/ y RETIROS/)."""
    ruta_base = os.path.join(os.getcwd(), "Reportes_Guardados")
    nombres = set()
    for grupo in ("ENTREGA", "RETIROS"):
        ruta_grupo = os.path.join(ruta_base, grupo)
        if os.path.isdir(ruta_grupo):
            for cli in os.listdir(ruta_grupo):
                if os.path.isdir(os.path.join(ruta_grupo, cli)):
                    nombres.add(cli)
    return sorted(nombres)


def _normalizar_cliente(valor):
    """Normaliza el nombre de cliente a MAYÚSCULAS y sin espacios (para carpeta)."""
    return (
        re.sub(r"[^a-zA-Z0-9\s_\-]", "", valor).strip().replace(" ", "_").upper()
        or "GENERAL"
    )


def dialogo_nuevo_lote(ventana, mov_sugerido="Entrega"):
    """
    Crea (o reabre) un lote: cliente + movimiento + sus dos documentos, la OT
    y la guia. El que nombra la carpeta es obligatorio y depende del
    movimiento; el otro es opcional y se puede completar despues.
    Devuelve el lote listo para recibir equipos, o None si se cancela.
    """
    existentes = _clientes_existentes()
    win = ventana_modal(ventana, "Nueva OT / Guia", "480x430")
    resultado = {"lote": None}

    titulo_ui(
        win, "🗂️ ¿A qué OT pertenece lo que vas a escanear?", size=12, pady=(16, 2)
    )
    tk.Label(
        win,
        text="El movimiento y el número los define la OT, no cada equipo.",
        font=fuente(9),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(0, 10))

    frm = tk.Frame(win, bg=COLORS["fondo"])
    frm.pack(pady=4)

    tk.Label(frm, text="Cliente:", font=fuente(10, True), bg=COLORS["fondo"]).grid(
        row=0, column=0, sticky="e", padx=6, pady=5
    )
    combo_cli = ttk.Combobox(frm, values=existentes, font=fuente(11), width=26)
    combo_cli.grid(row=0, column=1, padx=6, pady=5)
    combo_cli.focus()

    tk.Label(frm, text="Movimiento:", font=fuente(10, True), bg=COLORS["fondo"]).grid(
        row=1, column=0, sticky="e", padx=6, pady=5
    )
    combo_mov = ttk.Combobox(
        frm, values=["Entrega", "Retiro"], state="readonly", font=fuente(11), width=24
    )
    combo_mov.set(mov_sugerido)
    combo_mov.grid(row=1, column=1, padx=6, pady=5)

    lbl_num = tk.Label(frm, text="N° de OT:", font=fuente(10, True), bg=COLORS["fondo"])
    lbl_num.grid(row=2, column=0, sticky="e", padx=6, pady=5)
    entry_num = tk.Entry(frm, font=fuente(11), width=28)
    entry_num.grid(row=2, column=1, padx=6, pady=5)

    # El segundo documento del lote. Es opcional: muchas veces se conoce la OT
    # antes que la guía (o al revés) y no puede frenar la creación del lote.
    lbl_sec = tk.Label(frm, text="N° de Guía:", font=fuente(10, True), bg=COLORS["fondo"])
    lbl_sec.grid(row=3, column=0, sticky="e", padx=6, pady=5)
    entry_sec = tk.Entry(frm, font=fuente(11), width=28)
    entry_sec.grid(row=3, column=1, padx=6, pady=5)
    tk.Label(
        frm,
        text="(opcional, se puede completar después)",
        font=fuente(8),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).grid(row=4, column=1, sticky="w", padx=6)

    lbl_prev = tk.Label(
        win, text="", font=fuente(9, True), bg=COLORS["fondo"], fg=COLORS["azul"]
    )
    lbl_prev.pack(pady=(8, 0))

    def refrescar(event=None):
        mov = combo_mov.get()
        # El número que nombra la carpeta cambia con el movimiento: la entrega
        # se archiva por OT y el retiro por guía. El otro campo es el opuesto.
        es_entrega = mov == "Entrega"
        lbl_num.config(text="N° de OT:" if es_entrega else "N° de Guía:")
        lbl_sec.config(text="N° de Guía:" if es_entrega else "N° de OT:")
        numero = normalizar_numero(entry_num.get(), mov)
        cliente = _normalizar_cliente(combo_cli.get()) if combo_cli.get().strip() else ""
        if numero and cliente:
            lbl_prev.config(
                text=f"📁 {_MOV_CFG[mov]['grupo']} / {cliente} / "
                f"{carpeta_lote(mov, numero)}"
            )
        else:
            lbl_prev.config(text="")

    combo_mov.bind("<<ComboboxSelected>>", refrescar)
    entry_num.bind("<KeyRelease>", refrescar)
    combo_cli.bind("<KeyRelease>", refrescar)
    refrescar()

    def confirmar(event=None):
        if not combo_cli.get().strip():
            messagebox.showwarning(
                "Falta el cliente", "Elige o escribe un cliente.", parent=win
            )
            return
        mov = combo_mov.get()
        numero = normalizar_numero(entry_num.get(), mov)
        if not numero:
            messagebox.showwarning(
                "Número inválido",
                "Escribe el número de OT / guía.\n\n"
                "Solo letras, números, '-' y '_'.",
                parent=win,
            )
            return

        lote = hacer_lote(
            mov,
            _normalizar_cliente(combo_cli.get()),
            numero,
            numero_sec=normalizar_numero(entry_sec.get(), mov),
        )

        # Reabrir una OT ya cerrada es posible, pero nunca en silencio.
        if esta_cerrado(lote["ruta"]):
            meta = leer_cerrado(lote["ruta"])
            if not messagebox.askyesno(
                "OT ya cerrada",
                f"La {carpeta_lote(mov, numero)} de {lote['cliente']} ya está "
                f"cerrada ({meta.get('equipos', '?')} equipo(s), "
                f"{str(meta.get('cerrado', ''))[:10]}).\n\n"
                "¿Reabrirla para agregar más equipos?",
                parent=win,
            ):
                return
            reabrir_lote(lote)

        os.makedirs(lote["ruta"], exist_ok=True)
        # Si el lote ya existía y no se escribió el segundo número, se conserva
        # el que tenga la carpeta en vez de borrarlo por dejar el campo vacío.
        if not lote["numero_sec"]:
            lote["numero_sec"] = _numero_sec_en_disco(lote["ruta"])
        guardar_meta_lote(lote)
        resultado["lote"] = lote
        win.destroy()

    entry_num.bind("<Return>", confirmar)
    entry_sec.bind("<Return>", confirmar)

    frame_bts = tk.Frame(win, bg=COLORS["fondo"])
    frame_bts.pack(pady=16)
    boton_accion(frame_bts, "✅ Usar esta OT", confirmar).pack(side="left", padx=6)
    tk.Button(
        frame_bts,
        text="Cancelar",
        command=win.destroy,
        bg=COLORS["gris"],
        fg="white",
        font=fuente(10, True),
        padx=14,
        pady=7,
        cursor="hand2",
        relief="flat",
    ).pack(side="left", padx=6)

    win.wait_window()
    return resultado["lote"]


def dialogo_elegir_lote(ventana, titulo, texto_boton):
    """
    Lista las OT abiertas para elegir una. Devuelve el lote o None.
    Solo mira carpetas sin '.cerrado': no abre archivos de las OT finalizadas.
    """
    abiertos = listar_lotes_abiertos()
    if not abiertos:
        messagebox.showinfo(
            "Sin OT abiertas",
            "No hay ninguna OT en curso.\n\nCrea una con '➕ NUEVA OT'.",
            parent=ventana,
        )
        return None

    win = ventana_modal(ventana, titulo, "580x430")
    resultado = {"lote": None}

    titulo_ui(win, titulo, size=12, pady=(16, 6))

    cols = ("cliente", "lote", "equipos", "modificado")
    tree = ttk.Treeview(win, columns=cols, show="headings", height=10)
    for c, txt, w in (
        ("cliente", "Cliente", 180),
        ("lote", "OT / Guía", 130),
        ("equipos", "Equipos", 70),
        ("modificado", "Última actividad", 140),
    ):
        tree.heading(c, text=txt)
        tree.column(c, width=w, anchor="center" if c == "equipos" else "w")
    tree.pack(fill="both", expand=True, padx=16, pady=6)

    activo = cargar_lote_activo()
    for i, lote in enumerate(abiertos):
        tree.insert(
            "",
            "end",
            iid=str(i),
            values=(
                lote["cliente"].replace("_", " "),
                carpeta_lote(lote["mov"], lote["numero"]),
                lote["equipos"],
                datetime.fromtimestamp(lote["modificado"]).strftime("%d-%m %H:%M"),
            ),
        )
        if activo and os.path.normcase(activo["ruta"]) == os.path.normcase(lote["ruta"]):
            tree.selection_set(str(i))
    if not tree.selection():
        tree.selection_set("0")

    def confirmar(event=None):
        sel = tree.selection()
        if not sel:
            return
        resultado["lote"] = abiertos[int(sel[0])]
        win.destroy()

    tree.bind("<Double-1>", confirmar)

    frame_bts = tk.Frame(win, bg=COLORS["fondo"])
    frame_bts.pack(pady=12)
    boton_accion(frame_bts, texto_boton, confirmar).pack(side="left", padx=6)
    tk.Button(
        frame_bts,
        text="Cancelar",
        command=win.destroy,
        bg=COLORS["gris"],
        fg="white",
        font=fuente(10, True),
        padx=14,
        pady=7,
        cursor="hand2",
        relief="flat",
    ).pack(side="left", padx=6)

    win.wait_window()
    return resultado["lote"]


def accion_migrar_legacy(ventana):
    """Adopta los equipos sueltos del modelo anterior dentro de una OT."""
    legacy = pendientes_legacy()
    total = sum(len(v) for v in legacy.values())
    if not total:
        return

    for mov, rutas in legacy.items():
        if not rutas:
            continue
        messagebox.showinfo(
            "Migrar equipos sueltos",
            f"Hay {len(rutas)} equipo(s) de {_MOV_CFG[mov]['tipo_doc'].upper()} "
            "escaneados con el formato anterior.\n\n"
            "Elige a qué OT pertenecen.",
        )
        lote = dialogo_nuevo_lote(ventana, mov_sugerido=mov)
        if not lote:
            continue
        if lote["mov"] != mov:
            messagebox.showwarning(
                "Movimiento distinto",
                f"Esos equipos son de {mov} y la OT elegida es de "
                f"{lote['mov']}.\n\nSe migran igual con el movimiento de la OT.",
            )
        migrados = migrar_legacy_a_lote(rutas, lote)
        messagebox.showinfo(
            "✅ Migrados",
            f"{migrados} equipo(s) quedaron en {etiqueta_lote(lote)}.",
        )
        guardar_lote_activo(lote)

    mostrar_menu_principal(ventana)


def accion_nueva_ot(ventana):
    lote = dialogo_nuevo_lote(ventana)
    if not lote:
        return
    guardar_lote_activo(lote)
    mostrar_menu_principal(ventana)


def accion_cambiar_lote(ventana):
    lote = dialogo_elegir_lote(ventana, "Cambiar de OT", "✅ Activar")
    if not lote:
        return
    guardar_lote_activo(lote)
    mostrar_menu_principal(ventana)


def accion_cerrar_lote(ventana):
    lote = dialogo_elegir_lote(ventana, "Cerrar OT", "📦 Cerrar esta OT")
    if not lote:
        return

    if lote["equipos"] == 0 and not messagebox.askyesno(
        "OT vacía",
        f"{etiqueta_lote(lote)} no tiene equipos.\n\n¿Cerrarla igual?",
    ):
        return

    if not messagebox.askyesno(
        "Confirmar cierre",
        f"¿Cerrar {etiqueta_lote(lote)}?\n\n"
        f"{lote['equipos']} equipo(s). Se regenera el JSON FUSIONADO con todo "
        "lo escaneado y la carpeta queda marcada como finalizada.",
    ):
        return

    total, equipos_malos = cerrar_lote(lote)
    ruta_bk = crear_backup("cierre")

    extra = ""
    if equipos_malos:
        extra = (
            f"\n\n📄 Informe técnico generado con {len(equipos_malos)} "
            "equipo(s) defectuoso(s)."
        )
    if ruta_bk:
        extra += f"\n\n🗄️ Backup: {os.path.basename(ruta_bk)}"
    messagebox.showinfo(
        "📦 OT Cerrada",
        f"{etiqueta_lote(lote)}\n\n"
        f"{total} equipo(s) en el JSON FUSIONADO.\n"
        f"Carpeta: {lote['ruta']}{extra}",
    )
    mostrar_menu_principal(ventana)


# ─────────────────────────────────────────────────────────
#  ENVÍO POR SCP (SSH)
#
#  Sube al servidor solo lo que es fuente de verdad o entregable —JSON
#  individuales, JSON FUSIONADO, informe Word y el marcador '.cerrado'—
#  replicando el árbol <GRUPO>/<CLIENTE>/<OT>. El HTML se deja fuera a
#  propósito: es un derivado pesado que allá se rearma desde los JSON con
#  "RECONSTRUIR DESDE JSONs", idéntico y con todas sus funcionalidades.
# ─────────────────────────────────────────────────────────

# Extensiones/archivos que SÍ viajan. Todo lo demás (HTML, .enviado, .lote_activo,
# temporales de Office) se queda en la máquina.
_SCP_EXTENSIONES = (".json", ".docx")


def archivos_a_enviar(ruta_ot):
    """
    Archivos de una OT que deben subir al servidor, ya ordenados.

    Se envían los JSON (individuales + FUSIONADO), el Word de fallas y los
    marcadores de estado '.cerrado' y '.lote' (este último trae el número de
    documento que la ruta no codifica y sin él la OT se reconstruye sin su
    guía/OT secundaria). Quedan fuera el HTML —derivado, se rearma allá— y
    '.enviado', que es control local de sincronización.
    """
    archivos = []
    try:
        nombres = sorted(os.listdir(ruta_ot))
    except Exception:
        return archivos
    for nombre in nombres:
        ruta = os.path.join(ruta_ot, nombre)
        if not os.path.isfile(ruta):
            continue
        if nombre == _MARCADOR_ENVIADO:
            continue
        if nombre in (_MARCADOR_CERRADO, _ARCHIVO_META_LOTE):
            archivos.append(ruta)
            continue
        if nombre.startswith("~$"):  # temporales de Word
            continue
        if nombre.lower().endswith(_SCP_EXTENSIONES):
            archivos.append(ruta)
    return archivos


def _candidatas_clave_ssh():
    """
    Dónde se busca la clave privada, en orden de prioridad.

    El ejecutable se copia a equipos recién clonados por FOG, donde el perfil
    de usuario es otro y C:\\Users\\<alguien>\\.ssh no existe. Por eso la clave
    que viaja AL LADO DEL EXE manda sobre la ruta fija: es la única que existe
    seguro en esa máquina. La ruta de SCP_CLAVE queda como respaldo para el
    equipo de escritorio donde se trabaja a diario.
    """
    nombre = os.path.basename(SCP_CLAVE) if SCP_CLAVE else "id_ed25519"
    candidatas = [
        os.path.join(ruta_base_proyecto, nombre),
        os.path.join(ruta_base_proyecto, "clave_scp"),
    ]
    if SCP_CLAVE:
        candidatas.append(SCP_CLAVE)
    candidatas.append(os.path.join(os.path.expanduser("~"), ".ssh", nombre))
    # Sin duplicados y respetando el orden.
    vistas, unicas = set(), []
    for c in candidatas:
        clave = os.path.normcase(os.path.abspath(c))
        if clave not in vistas:
            vistas.add(clave)
            unicas.append(c)
    return unicas


def resolver_clave_ssh():
    """Primera clave privada que exista de verdad, o None si no hay ninguna."""
    for ruta in _candidatas_clave_ssh():
        if os.path.isfile(ruta):
            return ruta
    return None


def _ruta_remota(*partes):
    """
    Une un path remoto estilo POSIX. No se usa os.path.join: acá corre Windows
    y devolvería '\\', que el servidor Linux trata como parte del nombre.
    """
    absoluta = str(partes[0]).startswith("/")
    limpias = [str(p).replace("\\", "/").strip("/") for p in partes if p]
    unida = "/".join(limpias)
    return f"/{unida}" if absoluta else unida


def abrir_conexion_scp():
    """
    Cliente SSH conectado y listo. Lanza una excepción con mensaje legible si
    la clave no existe, el host no responde o las credenciales fallan.
    """
    if not SCP_AVAILABLE:
        raise RuntimeError(
            "Falta la librería 'paramiko'/'scp'.\n\n"
            "Instalar con:  pip install paramiko scp"
        )

    ruta_clave = resolver_clave_ssh()
    clave = None
    if ruta_clave:
        # El tipo de clave (RSA/Ed25519/ECDSA) no se conoce de antemano: se
        # prueban todos y se usa el primero que la lea.
        ultimo = None
        for tipo in (
            paramiko.Ed25519Key,
            paramiko.RSAKey,
            paramiko.ECDSAKey,
        ):
            try:
                clave = tipo.from_private_key_file(
                    ruta_clave, password=SCP_CLAVE_PASSPHRASE or None
                )
                break
            except Exception as e:
                ultimo = e
        if clave is None:
            raise RuntimeError(
                f"No se pudo leer la clave SSH:\n{ruta_clave}\n\n{ultimo}\n\n"
                "Si la clave tiene passphrase, cárgala en SCP_CLAVE_PASSPHRASE."
            )
    elif SCP_CLAVE:
        raise RuntimeError(
            "No se encontró la clave SSH en ninguna de estas ubicaciones:\n\n"
            + "\n".join(f"  • {r}" for r in _candidatas_clave_ssh())
            + "\n\nEn un equipo clonado la clave tiene que viajar junto al "
            "ejecutable, en la misma carpeta."
        )

    cliente = paramiko.SSHClient()
    cliente.load_system_host_keys()
    # El servidor es de la empresa y su huella puede cambiar al reinstalarlo:
    # aceptarla evita que el envío se caiga por un host key desconocido.
    cliente.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        cliente.connect(
            hostname=SCP_HOST,
            port=SCP_PUERTO,
            username=SCP_USUARIO,
            pkey=clave,
            timeout=SCP_TIMEOUT,
            banner_timeout=SCP_TIMEOUT,
            auth_timeout=SCP_TIMEOUT,
            look_for_keys=clave is None,
            allow_agent=clave is None,
        )
    except paramiko.AuthenticationException:
        cliente.close()
        raise RuntimeError(
            f"El servidor rechazó las credenciales de '{SCP_USUARIO}'.\n\n"
            "Verifica el usuario y que la clave pública esté en el "
            "authorized_keys del servidor."
        )
    except Exception as e:
        cliente.close()
        raise RuntimeError(
            f"No se pudo conectar a {SCP_HOST}:{SCP_PUERTO}\n\n{e}\n\n"
            "Verifica la VPN/Wi-Fi y que el servicio SSH esté arriba."
        )
    return cliente


def _correr_remoto(cliente, comando):
    """Ejecuta un comando en el servidor. Devuelve (código, stdout, stderr)."""
    _stdin, stdout, stderr = cliente.exec_command(comando, timeout=SCP_TIMEOUT)
    salida = stdout.read().decode("utf-8", "replace").strip()
    error = stderr.read().decode("utf-8", "replace").strip()
    return stdout.channel.recv_exit_status(), salida, error


def verificar_destino_remoto(cliente):
    """
    Comprueba que SCP_DESTINO exista y sea escribible ANTES de subir nada.

    La carpeta base debe existir de antes (la crea el NAS al publicar el
    recurso compartido); acá solo se crean las subcarpetas de cliente y OT. Si
    la ruta estuviera mal escrita, un 'mkdir -p' la crearía igual —en el disco
    de sistema y fuera del recurso compartido— y los reportes acabarían en un
    árbol fantasma que nadie ve por la red. Mejor fallar y decirlo.
    """
    codigo, _salida, _error = _correr_remoto(cliente, f"test -d '{SCP_DESTINO}'")
    if codigo != 0:
        raise RuntimeError(
            f"La carpeta base remota no existe:\n{SCP_DESTINO}\n\n"
            "Tiene que ser la ruta ABSOLUTA de la carpeta compartida en el "
            "servidor. En OpenMediaVault no es el nombre del recurso: mirá la "
            "ruta real en Almacenamiento → Carpetas compartidas (suele ser "
            "/srv/dev-disk-by-uuid-…/<carpeta>).\n\n"
            "Corregí la constante SCP_DESTINO."
        )
    codigo, _salida, _error = _correr_remoto(cliente, f"test -w '{SCP_DESTINO}'")
    if codigo != 0:
        raise RuntimeError(
            f"El usuario '{SCP_USUARIO}' no puede escribir en:\n{SCP_DESTINO}\n\n"
            "Dale permiso de lectura/escritura sobre la carpeta compartida "
            "(en OMV: Carpetas compartidas → ACL)."
        )


def _mkdir_remoto(cliente, ruta):
    """Crea la carpeta remota (y sus padres). Falla ruidosamente si no puede."""
    codigo, _salida, error = _correr_remoto(cliente, f"mkdir -p '{ruta}'")
    if codigo != 0:
        raise RuntimeError(f"No se pudo crear la carpeta remota {ruta}: {error}")


def enviar_ot_por_scp(cliente, grupo, cliente_folder, nombre_ot, ruta_origen):
    """
    Sube los archivos de UNA OT. Devuelve la cantidad de archivos subidos.
    Sobrescribe lo que ya esté allá: una OT abierta se re-envía varias veces y
    cada envío debe dejar el servidor al día.
    """
    archivos = archivos_a_enviar(ruta_origen)
    if not archivos:
        return 0
    destino = _ruta_remota(SCP_DESTINO, grupo, cliente_folder, nombre_ot)
    _mkdir_remoto(cliente, destino)
    with SCPClient(cliente.get_transport(), socket_timeout=SCP_TIMEOUT) as scp:
        # De a uno y no en bloque: si un archivo falla se sabe cuál fue.
        for ruta in archivos:
            scp.put(ruta, remote_path=_ruta_remota(destino, os.path.basename(ruta)))
    return len(archivos)


# ============================================================================
#  CONTROL CENTRAL POR SSH — OT activa compartida y seguimiento del lote
# ----------------------------------------------------------------------------
#  El servidor hace de pizarra común: no corre ningún servicio nuevo, solo
#  guarda los archivos de control en <SCP_DESTINO>/_control. El equipo de
#  escritorio publica la OT activa y los equipos recién clonados la leen para
#  saber a qué OT mandar lo que escanean.
#
#      _control/lote_activo.json     OT que recibe lo que se escanea
#      <GRUPO>/<CLIENTE>/<OT>/       equipos que fueron llegando
#      <GRUPO>/<CLIENTE>/<OT>/.esperados   seriales que se esperan del lote
# ============================================================================
_CONTROL_REMOTO = "_control"
_ARCHIVO_ESPERADOS = ".esperados"


def _ruta_control_remota(*partes):
    return _ruta_remota(SCP_DESTINO, _CONTROL_REMOTO, *partes)


def ruta_remota_lote(lote):
    """Carpeta de la OT en el servidor, con el mismo árbol que en local."""
    cfg = _MOV_CFG[lote["mov"]]
    return _ruta_remota(
        SCP_DESTINO,
        cfg["grupo"],
        lote["cliente"],
        carpeta_lote(lote["mov"], lote["numero"]),
    )


def _subir_texto(cliente, texto, ruta_destino_remota):
    """
    Escribe un archivo de texto en el servidor pasando por un temporal local.

    No se hace con 'echo' por SSH a propósito: los nombres de cliente y las
    observaciones traen tildes, comillas y espacios, y armar ese comando a mano
    es una fuente segura de archivos corruptos.
    """
    _mkdir_remoto(cliente, os.path.dirname(ruta_destino_remota).replace("\\", "/"))
    tmp = tempfile.NamedTemporaryFile(
        mode="w", suffix=".tmp", delete=False, encoding="utf-8"
    )
    try:
        tmp.write(texto)
        tmp.close()
        with SCPClient(cliente.get_transport(), socket_timeout=SCP_TIMEOUT) as scp:
            scp.put(tmp.name, remote_path=ruta_destino_remota)
    finally:
        try:
            os.unlink(tmp.name)
        except Exception:
            pass


def _leer_texto_remoto(cliente, ruta):
    """Contenido de un archivo del servidor, o None si no existe."""
    codigo, salida, _error = _correr_remoto(cliente, f"cat '{ruta}' 2>/dev/null")
    return salida if codigo == 0 else None


def publicar_lote_activo(cliente, lote):
    """
    Deja en el servidor cuál es la OT que recibe lo escaneado.

    Es el mismo contenido que el '.lote_activo.json' local, pero compartido:
    un equipo recién clonado no tiene historia ni archivos previos, así que la
    única forma de que sepa a qué OT pertenece es preguntárselo al servidor.
    """
    datos = {k: lote.get(k, "") for k in ("mov", "cliente", "numero", "numero_sec")}
    datos["publicado"] = datetime.now().isoformat(timespec="seconds")
    datos["publicado_por"] = socket.gethostname()
    _subir_texto(
        cliente,
        json.dumps(datos, ensure_ascii=False, indent=2),
        _ruta_control_remota(_ARCHIVO_LOTE_ACTIVO),
    )


def leer_lote_activo_remoto(cliente):
    """OT activa publicada en el servidor, o None si no hay ninguna."""
    crudo = _leer_texto_remoto(cliente, _ruta_control_remota(_ARCHIVO_LOTE_ACTIVO))
    if not crudo:
        return None
    try:
        cfg = json.loads(crudo)
        lote = hacer_lote(
            cfg["mov"], cfg["cliente"], cfg["numero"], numero_sec=cfg.get("numero_sec", "")
        )
    except Exception:
        return None
    # Una OT ya cerrada no debe seguir recibiendo equipos aunque el puntero
    # haya quedado publicado.
    codigo, _s, _e = _correr_remoto(
        cliente, f"test -e '{_ruta_remota(ruta_remota_lote(lote), _MARCADOR_CERRADO)}'"
    )
    return None if codigo == 0 else lote


def despublicar_lote_activo(cliente):
    _correr_remoto(cliente, f"rm -f '{_ruta_control_remota(_ARCHIVO_LOTE_ACTIVO)}'")


def enviar_equipo_por_scp(cliente, lote, ruta_json):
    """
    Sube el JSON de UN equipo a la OT del servidor.

    Es lo único que manda un equipo clonado, y es deliberado: su carpeta local
    contiene un solo equipo, así que su FUSIONADO y su HTML describen ese
    equipo nada más. Subirlos pisaría los del servidor —que acumulan todo el
    lote— y el reporte de la OT quedaría con un solo equipo. Los derivados se
    rehacen en el panel central, que sí ve el lote completo.
    """
    destino = ruta_remota_lote(lote)
    _mkdir_remoto(cliente, destino)
    with SCPClient(cliente.get_transport(), socket_timeout=SCP_TIMEOUT) as scp:
        scp.put(ruta_json, remote_path=_ruta_remota(destino, os.path.basename(ruta_json)))
    return _ruta_remota(destino, os.path.basename(ruta_json))


_SEPARADOR_REMOTO = "===ARCHIVO:"


def equipos_remotos(cliente, lote):
    """
    Equipos que ya llegaron a la OT en el servidor: [{SERIAL, MODELO, ...}].

    Se traen todos los JSON en UNA sola vuelta (un 'cat' encadenado) en vez de
    un comando por archivo: con un lote de decenas de equipos la diferencia
    entre una conexión y decenas es la diferencia entre un panel usable y uno
    que tarda medio minuto en refrescar.
    """
    carpeta = ruta_remota_lote(lote)
    comando = (
        f"for f in '{carpeta}'/*.json; do "
        f'[ -e "$f" ] || continue; '
        f'case "$f" in *_FUSIONADO.json) continue;; esac; '
        f'echo "{_SEPARADOR_REMOTO}$(basename "$f")"; cat "$f"; '
        f"done"
    )
    codigo, salida, _error = _correr_remoto(cliente, comando)
    if codigo != 0 or not salida:
        return []

    equipos = []
    for bloque in salida.split(_SEPARADOR_REMOTO):
        if not bloque.strip():
            continue
        nombre, _, cuerpo = bloque.partition("\n")
        try:
            jdata = json.loads(cuerpo)
        except Exception:
            continue
        if isinstance(jdata, dict) and jdata.get("SERIAL"):
            jdata["_archivo"] = nombre.strip()
            equipos.append(jdata)
    equipos.sort(key=_clave_orden, reverse=True)
    return equipos


def leer_esperados(cliente, lote):
    """Seriales que se esperan en el lote. Lista vacía si no se cargó ninguna."""
    crudo = _leer_texto_remoto(
        cliente, _ruta_remota(ruta_remota_lote(lote), _ARCHIVO_ESPERADOS)
    )
    if not crudo:
        return []
    try:
        datos = json.loads(crudo)
        return [str(s).strip().upper() for s in datos if str(s).strip()]
    except Exception:
        return []


def guardar_esperados(cliente, lote, seriales):
    """Fija la lista de seriales que deben llegar antes de dar el lote por cerrado."""
    limpios = []
    for s in seriales:
        s = str(s).strip().upper()
        if s and s not in limpios:
            limpios.append(s)
    _subir_texto(
        cliente,
        json.dumps(limpios, ensure_ascii=False, indent=2),
        _ruta_remota(ruta_remota_lote(lote), _ARCHIVO_ESPERADOS),
    )
    return limpios


def descargar_equipos_remotos(cliente, lote):
    """
    Baja a la carpeta local de la OT los JSON que mandaron los equipos y rehace
    el HTML y el FUSIONADO. Devuelve (bajados, total_equipos).

    Es el paso que cierra el circuito: cada equipo clonado sube su JSON y nada
    más, así que el reporte completo del lote no existe en ningún lado hasta
    que alguien junta las partes. Acá se juntan.
    """
    carpeta = ruta_remota_lote(lote)
    codigo, salida, _error = _correr_remoto(
        cliente,
        f"ls -1 '{carpeta}'/*.json 2>/dev/null | grep -v '_FUSIONADO.json$' || true",
    )
    remotos = [l.strip() for l in salida.splitlines() if l.strip()] if codigo == 0 else []
    if not remotos:
        return 0, contar_equipos(lote["ruta"])

    os.makedirs(lote["ruta"], exist_ok=True)
    bajados = 0
    with SCPClient(cliente.get_transport(), socket_timeout=SCP_TIMEOUT) as scp:
        for remoto in remotos:
            local = os.path.join(lote["ruta"], os.path.basename(remoto))
            try:
                scp.get(remoto, local_path=local)
                bajados += 1
            except Exception as e:
                print(f"No se pudo bajar {remoto}: {e}")
    guardar_meta_lote(lote)
    return bajados, regenerar_lote(lote)


def cerrar_lote_remoto(cliente, lote, total):
    """Marca la OT como cerrada en el servidor y baja el puntero de OT activa."""
    marca = json.dumps(
        {
            "cliente": lote["cliente"],
            "numero": lote["numero"],
            "equipos": total,
            "cerrado": datetime.now().isoformat(timespec="seconds"),
            "cerrado_por": socket.gethostname(),
        },
        ensure_ascii=False,
        indent=2,
    )
    _subir_texto(
        cliente, marca, _ruta_remota(ruta_remota_lote(lote), _MARCADOR_CERRADO)
    )
    activo = leer_lote_activo_remoto(cliente)
    if activo and os.path.normcase(activo["ruta"]) == os.path.normcase(lote["ruta"]):
        despublicar_lote_activo(cliente)


def accion_enviar_red(ventana):
    carpeta_local = ruta_reportes()

    if not os.path.exists(carpeta_local):
        messagebox.showwarning("Aviso", "No hay reportes locales para enviar.")
        return

    if not SCP_AVAILABLE:
        messagebox.showerror(
            "Falta dependencia",
            "El envío por SCP necesita las librerías 'paramiko' y 'scp'.\n\n"
            "Instálalas con:  pip install paramiko scp\n"
            "y vuelve a compilar el ejecutable.",
            parent=ventana,
        )
        return

    # Pendientes = carpetas de OT SIN el marcador '.enviado'. El marcador va por
    # OT (no por cliente): agregar un equipo a una sola OT no obliga a re-subir
    # las demás carpetas de ese cliente, que pueden ser miles.
    # En el servidor se replica el árbol COMPLETO ENTREGA/ACME/OT-553: con el
    # grupo delante, lo que se baje de allá se reconstruye como OT canónica.
    pendientes = []  # (grupo, cliente, nombre_ot, ruta_origen, abierta)
    hay_carpetas = False
    for grupo in ("ENTREGA", "RETIROS"):
        ruta_grupo = os.path.join(carpeta_local, grupo)
        if not os.path.isdir(ruta_grupo):
            continue
        for cli in os.listdir(ruta_grupo):
            ruta_cli = os.path.join(ruta_grupo, cli)
            if not os.path.isdir(ruta_cli):
                continue
            for ot in os.listdir(ruta_cli):
                ruta_ot = os.path.join(ruta_cli, ot)
                if not os.path.isdir(ruta_ot):
                    continue
                hay_carpetas = True
                if not os.path.exists(os.path.join(ruta_ot, _MARCADOR_ENVIADO)):
                    pendientes.append(
                        (grupo, cli, ot, ruta_ot, not esta_cerrado(ruta_ot))
                    )

    if not pendientes:
        if hay_carpetas:
            messagebox.showinfo(
                "Nada nuevo por enviar",
                "Todas las OT ya fueron enviadas al servidor.",
            )
        else:
            messagebox.showwarning(
                "Aviso",
                "No hay ninguna OT para enviar.\n\nEscanea equipos primero.",
            )
        return

    total_archivos = sum(len(archivos_a_enviar(p[3])) for p in pendientes)
    abiertas = sum(1 for p in pendientes if p[4])
    aviso_abiertas = ""
    if abiertas:
        aviso_abiertas = (
            f"\n\n⚠️ {abiertas} de ellas siguen ABIERTAS: se enviarán tal como "
            "van y se volverán a enviar cuando les agregues equipos."
        )

    if not messagebox.askyesno(
        "Confirmar Envío",
        f"¿Enviar {len(pendientes)} carpeta(s) de OT al servidor por SCP?\n\n"
        f"Destino: {SCP_USUARIO}@{SCP_HOST}:{SCP_DESTINO}\n"
        f"Archivos: {total_archivos} (JSON individuales, FUSIONADO, Word y "
        f"marcadores).\n\n"
        f"El HTML no se sube: se reconstruye desde los JSON en destino."
        f"{aviso_abiertas}",
    ):
        return

    # --- Ventana de Carga Visual ---
    win_carga = tk.Toplevel(ventana)
    win_carga.title("Enviando...")
    win_carga.geometry("420x175")
    win_carga.configure(bg=COLORS["fondo"])
    win_carga.grab_set()  # Bloquea la ventana principal durante el envío

    tk.Label(win_carga, text="🚀", font=("Segoe UI", 24), bg=COLORS["fondo"]).pack(
        pady=(15, 5)
    )
    tk.Label(
        win_carga,
        text=f"Subiendo por SCP a {SCP_HOST}...",
        font=("Segoe UI", 10, "bold"),
        bg=COLORS["fondo"],
        fg=COLORS["azul"],
    ).pack()
    lbl_progreso = tk.Label(
        win_carga,
        text="Conectando…",
        font=("Segoe UI", 9),
        bg=COLORS["fondo"],
        fg=COLORS["gris_oscuro"],
    )
    lbl_progreso.pack(pady=(6, 0))
    tk.Label(
        win_carga,
        text="Por favor, no cierres el programa.",
        font=("Segoe UI", 9),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(6, 0))

    def progreso(texto):
        """Refresca el rótulo de avance desde el hilo de trabajo."""
        def _set():
            if win_carga.winfo_exists():
                lbl_progreso.config(text=texto)
        ventana.after(0, _set)

    # --- Función que hace el trabajo pesado en segundo plano ---
    def tarea_envio():
        # Backup previo al envío. Va en el hilo para no congelar la UI y nunca
        # lanza: si falla, el envío continúa igual.
        ruta_bk = crear_backup("envio")
        errores = []
        enviados = 0
        subidos = 0

        # Conexión y destino se validan juntos y ANTES del primer archivo: si
        # algo está mal, ninguna OT queda marcada como enviada.
        try:
            cliente = abrir_conexion_scp()
        except Exception as e:
            ventana.after(0, finalizar_envio, 0, 0, [str(e)], ruta_bk, True)
            return

        try:
            verificar_destino_remoto(cliente)
        except Exception as e:
            cliente.close()
            ventana.after(0, finalizar_envio, 0, 0, [str(e)], ruta_bk, True)
            return

        try:
            for i, (grupo, cli, nombre_ot, ruta_origen, _abierta) in enumerate(
                pendientes, 1
            ):
                progreso(f"[{i}/{len(pendientes)}] {cli} · {nombre_ot}")
                try:
                    n = enviar_ot_por_scp(cliente, grupo, cli, nombre_ot, ruta_origen)
                    subidos += n
                    # Marcar DESPUÉS de subir: si el envío falla, la OT sigue
                    # pendiente y se reintenta en el próximo envío.
                    with open(
                        os.path.join(ruta_origen, _MARCADOR_ENVIADO),
                        "w",
                        encoding="utf-8",
                    ) as mf:
                        mf.write(datetime.now().isoformat())
                    enviados += 1
                except Exception as e:
                    errores.append(f"{cli}/{nombre_ot}: {e}")
                    print(f"Error enviando {cli}/{nombre_ot}: {e}")
        finally:
            try:
                cliente.close()
            except Exception:
                pass

        ventana.after(0, finalizar_envio, enviados, subidos, errores, ruta_bk, False)

    # --- Función que cierra la carga y muestra el resultado ---
    def finalizar_envio(enviados, subidos, errores, ruta_bk=None, fatal=False):
        if win_carga.winfo_exists():
            win_carga.destroy()

        aviso_bk = (
            f"\n\n🗄️ Backup previo: {os.path.basename(ruta_bk)}" if ruta_bk else ""
        )
        if fatal:
            messagebox.showerror(
                "Error de conexión",
                f"{errores[0]}{aviso_bk}",
                parent=ventana,
            )
            return
        if errores:
            messagebox.showwarning(
                "Proceso con Observaciones",
                f"Se enviaron {enviados} carpeta(s) ({subidos} archivo(s)), "
                f"pero {len(errores)} fallaron:\n\n"
                + "\n".join(f"  • {e}" for e in errores[:8])
                + aviso_bk,
                parent=ventana,
            )
        else:
            messagebox.showinfo(
                "✅ Envío Exitoso",
                f"Se enviaron {enviados} carpeta(s) de OT "
                f"({subidos} archivo(s)) a {SCP_HOST} correctamente."
                f"{aviso_bk}",
                parent=ventana,
            )

    # Lanzamos el envío en un hilo (Thread) para que la pantalla no se congele
    threading.Thread(target=tarea_envio, daemon=True).start()

# ============================================================================
# 8. RECONSTRUCCIÓN / MERGE — rearmar el HTML a partir de los JSON guardados
# ============================================================================
def _es_html(valor):
    """True si el valor ya viene renderizado como filas de tabla."""
    return isinstance(valor, str) and valor.lstrip().startswith("<")


def _fecha_legible(jdata):
    """Fecha del encabezado del equipo cuando DATA no la trae, desde ESCANEADO."""
    ts = jdata.get("ESCANEADO", "")
    try:
        return datetime.fromisoformat(ts).strftime("%d/%m/%Y %H:%M")
    except Exception:
        return ""


def filas_hardware_desde_json(jdata):
    """
    (ram_html, disk_html, net_html) de un equipo, con cadena de respaldo.

    El HTML del reporte NO es decorativo: el JS lee de estas filas la RAM, los
    discos y sus seriales para "Copiar Filas" y el CSV de Valida. Un JSON al
    que le falte el HTML pre-renderizado —viejo, editado a mano o bajado del
    servidor— daría un reporte que se ve bien pero exporta "RAM Desconocida" y
    sin discos. Por eso se intenta, en orden:

        1. el HTML ya guardado en DATA (lo normal),
        2. los datos crudos (ram_raw / disk_raw / net_raw),
        3. lo que quede: los seriales sueltos o el SN_HDD del nivel superior.
    """
    data = jdata.get("DATA") or {}

    # --- RAM ---
    ram_html = data.get("ram_html")
    if not _es_html(ram_html):
        ram_html = render_filas_ram(data.get("ram_raw") or data.get("ram_rows"))

    # --- Discos ---
    disk_html = data.get("disk_html")
    if not _es_html(disk_html):
        discos = data.get("disk_raw") or data.get("disks_data")
        if not discos:
            # Sin descripción no se puede inventar el modelo del disco, pero el
            # serial sí se conserva: es el dato que viaja a Valida.
            seriales = data.get("disk_serials") or [
                s for s in str(jdata.get("SN_HDD", "")).split("_") if s
            ]
            seriales = [s for s in seriales if s and s != "PENDIENTE"]
            discos = [{"desc": "Disco interno", "serial": s} for s in seriales]
        disk_html, _ = render_filas_discos(discos)

    # --- Red ---
    net_html = data.get("net_rows")
    if not _es_html(net_html):
        if _es_html(data.get("net_html")):
            net_html = data["net_html"]
        else:
            net_html = render_filas_red(data.get("net_raw") or data.get("net_rows"))

    return ram_html, disk_html, net_html


def _reconstruir_entry_desde_json(jdata):
    data = jdata.get("DATA") or {}
    obs_final = jdata.get("OBS", "Sin observaciones")
    guia_final = jdata.get("GUIA_ID", "Sin Guía")
    ot = jdata.get("OT_ID", "Sin Orden")
    mov = jdata.get("TIPO_MOVIMIENTO", "Entrega")
    detalle_componentes = jdata.get("DETALLE_COMPONENTES", [])
    office_str = jdata.get("OFFICE", "No")

    # La Key de Office vive en SN_APP: hay que reponerla en la fila porque el
    # export a Valida (filasValida) la saca leyendo el "(Key: ...)" de este
    # texto. Si no se repone, la fila igual se ve en el reporte pero el export
    # la saltea: sin key no hay nada que cargar en Valida.
    office_html = ""
    if office_str and office_str != "No":
        sn_app = jdata.get("SN_APP", "")
        key_part = (
            f" (Key: {sn_app})" if sn_app and sn_app not in ("PENDIENTE", None) else ""
        )
        office_html = (
            f'<tr class="office-row"><th>🔑 Office Instalado</th>'
            f"<td><b>{office_str}</b>{key_part}</td></tr>"
        )

    # El nivel superior del JSON repite serial, modelo y licencia; se usan como
    # respaldo para que un registro con DATA incompleta —editado a mano o de una
    # versión vieja— igual salga con su identidad correcta en el reporte.
    serial = data.get("serial") or jdata.get("SERIAL") or "?"
    model = data.get("model") or jdata.get("MODELO") or "Desconocido"
    win_key = data.get("key") or jdata.get("SN_Win") or "N/A"

    safe_id = f"dev-{re.sub(r'[^a-zA-Z0-9]', '', str(serial)) or 'DESCONOCIDO'}"

    tiene_fallas = jdata.get("TIENE_FALLAS", False)
    fallas_str = ", ".join(
        [
            f"{c['nombre']} ({c['estado']})"
            for c in detalle_componentes
            if c.get("estado", "OK") != "OK"
        ]
    )

    falla_html = ""
    tag_falla_html = ""
    btn_imprimir = ""

    if tiene_fallas:
        tag_falla_html = '<span class="tag-falla">CON FALLAS</span>'
        falla_html = f'<tr class="falla-row"><th>Fallas Detectadas</th><td>{fallas_str}</td></tr>'
        btn_imprimir = (
            f'<button class="btn" style="background:#475569; padding:5px 12px; font-size:11px;" '
            f"onclick=\"imprimirEtiqueta('{safe_id}')\">🖨️ Etiqueta Individual</button>"
        )

    comps_json = html.escape(json.dumps(detalle_componentes))
    fusion_json = html.escape(json.dumps(_resumen_fusionado(jdata)))

    ram_html, disk_html, net_rows = filas_hardware_desde_json(jdata)
    bateria_html = (
        f'\n    <tr><th>🔋 Salud Batería</th>'
        f'<td>{data.get("Battery", "No detectada")}</td></tr>'
    )

    # El transformador se carga después del escaneo (módulo aparte), así que la
    # fila solo aparece si ya tiene serial. Va al HTML porque el export a
    # Valida (copiarFilas / CSV) lee del DOM, no de los JSON.
    sn_transf = jdata.get("SN_Transf", _TRANSF_PENDIENTE)
    transf_html = ""
    if sn_transf and sn_transf != _TRANSF_PENDIENTE:
        transf_html = (
            f'\n    <tr><th>🔌 Serie Transformador</th>'
            f'<td class="transf-tag">{sn_transf}</td></tr>'
        )

    return construir_entry_html(
        safe_id=safe_id,
        model=model,
        serial=serial,
        fecha=data.get("fecha") or _fecha_legible(jdata),
        key=win_key,
        cpu=data.get("cpu", "Desconocido"),
        mov=mov,
        num_ot=ot,
        guia_final=guia_final,
        obs_final=obs_final,
        ram_html=ram_html,
        disk_html=disk_html,
        net_rows=net_rows,
        comps_json=comps_json,
        fusion_json=fusion_json,
        tag_falla_html=tag_falla_html,
        falla_html=falla_html,
        btn_imprimir=btn_imprimir,
        office_html=office_html,
        bateria_html=bateria_html,
        transf_html=transf_html,
    )


def _lote_desde_ruta(ruta):
    """
    Deduce el lote de una carpeta con forma de OT: <CLIENTE>/<PREFIJO-NUMERO>.
    Devuelve None si es una carpeta cualquiera con JSON sueltos.

    No se exige el nivel de grupo (ENTREGA/RETIROS): una OT bajada del servidor
    puede llegar sin sus carpetas padre, y el prefijo de la carpeta —OT- o
    GUIA-— ya identifica el movimiento sin ambigüedad. Sin esto una OT
    recuperada se reconstruía como "carpeta suelta": HTML aparte, sin FUSIONADO
    y sin quedar integrada al árbol.
    """
    ruta = os.path.normpath(os.path.abspath(ruta))
    nombre = os.path.basename(ruta)
    cliente = os.path.basename(os.path.dirname(ruta))
    if not cliente:
        return None

    for mov, cfg in _MOV_CFG.items():
        pref = cfg["prefijo"] + "-"
        if nombre.startswith(pref) and nombre[len(pref):]:
            return hacer_lote(
                mov,
                cliente,
                nombre[len(pref):],
                ruta=ruta,
                numero_sec=_numero_sec_en_disco(ruta),
            )
    return None


def _carpetas_con_equipos(raiz):
    """Carpetas que contienen JSON de equipos, recursivo (incluida la raíz)."""
    encontradas = []
    for actual, _dirs, archivos in os.walk(raiz):
        if any(
            a.endswith(".json")
            and not a.endswith("_FUSIONADO.json")
            and not a.startswith(".")
            for a in archivos
        ):
            encontradas.append(actual)
    return sorted(encontradas)


def _cliente_dominante(equipos):
    """
    Cliente que aparece en más JSON del grupo. El JS del reporte saca de ahí el
    nombre de los archivos que descarga (nombreBaseReporte lee el <title>), así
    que usar el cliente real y no el nombre de la carpeta hace que el CSV y el
    fusionado exportados salgan con el mismo nombre que si los hubiera generado
    el escaneo.
    """
    conteo = {}
    for eq in equipos:
        cli = (eq.get("CLIENTE") or "").strip()
        if cli:
            conteo[cli] = conteo.get(cli, 0) + 1
    return max(conteo, key=conteo.get) if conteo else ""


def _reconstruir_carpeta_suelta(carpeta):
    """
    Rearma el HTML de una carpeta que NO pertenece al árbol de OT (por ejemplo
    JSON bajados del servidor a un directorio cualquiera). Agrupa por movimiento
    y escribe archivos '_RECONSTRUIDO.html' sin tocar nada más.
    Devuelve (generados, equipos, errores).

    El HTML sale con las mismas funcionalidades que el original: mismos
    HTML_START/HTML_END (copiar filas, CSV, fusionado, etiquetas), cada equipo
    con su data-comps y data-fusion, y un <title> con el cliente real para que
    las descargas salgan con el nombre correcto.

    A diferencia de regenerar_lote(), acá SÍ se escribe el informe Word cuando
    el grupo es de retiros y hay equipos con fallas: la carpeta suelta no pasó
    nunca por cerrar_lote(), así que el .docx no existe y hay que producirlo.
    """
    equipos_por_tipo = {"Entregas": [], "Retiros": []}
    errores = []

    for archivo in sorted(os.listdir(carpeta)):
        if not archivo.endswith(".json") or archivo.startswith("."):
            continue
        try:
            with open(os.path.join(carpeta, archivo), "r", encoding="utf-8") as f:
                jdata = json.load(f)
            # El fusionado (lista de equipos) ya está cubierto por los JSON
            # individuales; se ignora para no duplicar.
            if isinstance(jdata, list):
                continue
            # Basta el SERIAL: un registro sin DATA todavía sirve —el resto de
            # las filas se rearman desde el nivel superior del JSON— y
            # descartarlo dejaría al equipo fuera del reporte, que es justo lo
            # que hay que evitar al reconstruir.
            if not isinstance(jdata, dict) or not jdata.get("SERIAL"):
                errores.append(f"{archivo}: sin SERIAL, no es un JSON de equipo")
                continue
            if not jdata.get("DATA"):
                errores.append(
                    f"{archivo}: sin DATA — reconstruido desde el nivel superior "
                    "(puede faltar detalle de hardware)"
                )
            mov = jdata.get("TIPO_MOVIMIENTO", "Entrega")
            equipos_por_tipo["Entregas" if mov == "Entrega" else "Retiros"].append(jdata)
        except Exception as e:
            errores.append(f"{archivo}: {e}")

    nombre_carpeta = os.path.basename(os.path.normpath(carpeta))
    fecha_hoy = datetime.now().strftime("%Y%m%d")
    generados = []
    total = 0

    for tipo, equipos in equipos_por_tipo.items():
        if not equipos:
            continue
        equipos.sort(key=_clave_orden, reverse=True)
        etiqueta = (_cliente_dominante(equipos) or nombre_carpeta).replace("_", " ")
        titulo = f"{tipo} - {etiqueta}"
        contenido = HTML_START.replace("Inventario Maestro de Equipos", titulo).replace(
            "<title>Inventario Maestro</title>", f"<title>{titulo}</title>"
        )

        entries = []
        for jdata in equipos:
            try:
                entries.append(_reconstruir_entry_desde_json(jdata))
            except Exception as e:
                errores.append(f"S/N {jdata.get('SERIAL', '?')}: {e}")

        contenido += "\n".join(entries) + HTML_END
        contenido = fijar_total_count(contenido, len(entries))

        base = f"Reporte_{tipo}_{nombre_carpeta}_{fecha_hoy}_RECONSTRUIDO"
        with open(os.path.join(carpeta, base + ".html"), "w", encoding="utf-8") as f:
            f.write(contenido)
        generados.append(base + ".html")

        # El fusionado también se deja en disco: es lo que se sube a los
        # sistemas y no debería depender de que alguien abra el HTML y aprete
        # el botón de descarga.
        with open(
            os.path.join(carpeta, base + "_FUSIONADO.json"), "w", encoding="utf-8"
        ) as f:
            json.dump(
                [_resumen_fusionado(eq) for eq in equipos],
                f,
                ensure_ascii=False,
                indent=4,
            )
        generados.append(base + "_FUSIONADO.json")

        # Retiros con fallas: el Word es un entregable, no un derivado visual.
        malos = [e for e in equipos if e.get("TIENE_FALLAS")]
        if tipo == "Retiros" and malos and DOCX_AVAILABLE:
            try:
                generar_informe_word(
                    _cliente_dominante(equipos) or nombre_carpeta,
                    malos,
                    carpeta,
                    nombre_carpeta,
                )
            except Exception as e:
                errores.append(f"Informe Word: {e}")

        total += len(entries)

    return generados, total, errores


def accion_unir_Json(ventana):
    """
    Rearma los HTML desde los JSON guardados. Funciona recursivamente: se puede
    apuntar a una OT, a un cliente (con todas sus OT) o a ENTREGA/ completo.

    Si la carpeta pertenece al árbol de reportes se regenera el lote CANÓNICO
    (mismo HTML y mismo FUSIONADO que produce el escaneo, sin archivos sueltos);
    si es una carpeta cualquiera, se escribe un '_RECONSTRUIDO.html' aparte.

    Las dos ramas difieren en el informe Word: la de carpeta suelta lo genera,
    la canónica no (el .docx de una OT es cosa de cerrar_lote, sección 5).
    """
    carpeta = filedialog.askdirectory(
        title="Selecciona la carpeta a reconstruir (OT, cliente o grupo)",
        parent=ventana,
    )
    if not carpeta:
        return

    carpetas = _carpetas_con_equipos(carpeta)
    if not carpetas:
        messagebox.showwarning(
            "Sin JSONs",
            f"No se encontraron JSON de equipos en:\n{carpeta}\n\n"
            "(Se busca también dentro de las subcarpetas.)",
            parent=ventana,
        )
        return

    # Aquí sí se leen carpetas cerradas: el usuario las eligió a propósito.
    # El marcador '.cerrado' gobierna los recorridos automáticos, no esto.
    if len(carpetas) > 1 and not messagebox.askyesno(
        "Reconstrucción múltiple",
        f"Se encontraron {len(carpetas)} carpetas con equipos dentro de:\n"
        f"{carpeta}\n\n¿Reconstruir todas?",
        parent=ventana,
    ):
        return

    regenerados, sueltos, errores, total = [], [], [], 0

    for sub in carpetas:
        lote = _lote_desde_ruta(sub)
        try:
            if lote:
                n = regenerar_lote(lote)
                total += n
                regenerados.append(f"{etiqueta_lote(lote)} — {n} equipo(s)")
            else:
                generados, n, errs = _reconstruir_carpeta_suelta(sub)
                total += n
                errores.extend(errs)
                sueltos.extend(generados)
        except Exception as e:
            errores.append(f"{os.path.basename(sub)}: {e}")

    if not total and not regenerados:
        msg = "No se pudo leer ningún JSON válido."
        if errores:
            msg += "\n\nErrores:\n" + "\n".join(errores[:15])
        messagebox.showerror("Error", msg, parent=ventana)
        return

    resumen = f"✅ Reconstrucción completa: {total} equipo(s).\n"
    if regenerados:
        resumen += (
            f"\n\n📁 OT regeneradas ({len(regenerados)}) — HTML + JSON FUSIONADO:\n"
            + "\n".join(f"  • {r}" for r in regenerados[:12])
        )
        if len(regenerados) > 12:
            resumen += f"\n  … y {len(regenerados) - 12} más"
    if sueltos:
        resumen += f"\n\n📄 Carpetas sueltas ({len(sueltos)}):\n" + "\n".join(
            f"  • {a}" for a in sueltos[:12]
        )
    if errores:
        resumen += f"\n\n⚠️ Advertencias ({len(errores)}):\n" + "\n".join(
            errores[:10]
        )

    messagebox.showinfo("🔧 Reconstrucción Completa", resumen, parent=ventana)


# ============================================================================
# 9. MÓDULOS (ventanas Toplevel) — etiquetas manuales, transformadores,
#    panel de clonación y cambio de nombre del equipo.
# ============================================================================
def abrir_modulo_etiquetas_manual(ventana_padre):
    win = ventana_modal(
        ventana_padre, "Generador Manual de Etiquetas (Equipos Defectuosos)", "750x700"
    )
    titulo_ui(win, "🏷️ Generador Manual de Etiquetas", pady=(15, 5))
    tk.Label(
        win,
        text="Crea e imprime una etiqueta de diagnóstico sin alterar los reportes de inventario.",
        font=("Segoe UI", 9),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(0, 15))

    # --- DATOS DEL EQUIPO ---
    frame_datos = ttk.LabelFrame(win, text=" Datos del Equipo ")
    frame_datos.pack(fill="x", padx=25, pady=5)

    tk.Label(frame_datos, text="N° de Serie:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=0, column=0, padx=10, pady=8, sticky="e")
    ent_serial = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_serial.grid(row=0, column=1, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="Modelo:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=0, column=2, padx=10, pady=8, sticky="e")
    ent_modelo = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_modelo.grid(row=0, column=3, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="Procesador (CPU):", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=1, column=0, padx=10, pady=8, sticky="e")
    ent_cpu = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_cpu.grid(row=1, column=1, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="N° de Guía:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=1, column=2, padx=10, pady=8, sticky="e")
    ent_guia = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_guia.grid(row=1, column=3, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="N° de OT:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=2, column=0, padx=10, pady=8, sticky="e")
    ent_ot = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_ot.grid(row=2, column=1, padx=5, pady=8, sticky="w")

    # --- REVISIÓN DE COMPONENTES ---
    lf_fallas = ttk.LabelFrame(win, text=" Revisión de Componentes ")
    lf_fallas.pack(fill="both", expand=True, padx=25, pady=10)

    componentes_etiqueta = [
        "Carcaza", "Pantalla", "Teclado", "Touchpad", "Puertos USB",
        "Puertos Video", "Ethernet/Wi-Fi", "Placa base", "Memoria", "Disco duro", "Batería"
    ]

    comps_data = []

    frame_grilla = tk.Frame(lf_fallas, bg=COLORS["fondo"])
    frame_grilla.pack(padx=10, pady=10)

    tk.Label(frame_grilla, text="Componente", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=0, sticky="w", padx=5)
    tk.Label(frame_grilla, text="Estado", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=1, padx=5)
    tk.Label(frame_grilla, text="Observación (Si es OBS/MALO)", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=2, sticky="w", padx=5)

    for i, comp_name in enumerate(componentes_etiqueta, start=1):
        tk.Label(frame_grilla, text=comp_name, bg=COLORS["fondo"], font=("Segoe UI", 9)).grid(row=i, column=0, sticky="w", padx=5, pady=2)

        cmb_estado = ttk.Combobox(frame_grilla, values=["OK", "OBS", "MALO"], width=6, state="readonly", font=("Segoe UI", 9))
        cmb_estado.set("OK")
        cmb_estado.grid(row=i, column=1, padx=5, pady=2)

        ent_obs = tk.Entry(frame_grilla, width=45, font=("Segoe UI", 9), state=tk.DISABLED, bg="#f1f5f9")
        ent_obs.grid(row=i, column=2, padx=5, pady=2, sticky="w")

        def toggle_entry(event, e=ent_obs, c=cmb_estado):
            if c.get() != "OK":
                e.config(state=tk.NORMAL, bg="#ffffff")
                e.focus()
            else:
                e.delete(0, tk.END)
                e.config(state=tk.DISABLED, bg="#f1f5f9")

        cmb_estado.bind("<<ComboboxSelected>>", toggle_entry)
        comps_data.append({"nombre": comp_name, "estado": cmb_estado, "obs": ent_obs})

    def generar_y_imprimir():
        serial = ent_serial.get().strip() or "S/N"
        modelo = ent_modelo.get().strip() or "Desconocido"
        cpu = ent_cpu.get().strip() or "Desconocido"
        guia = ent_guia.get().strip() or "S/N"
        ot = ent_ot.get().strip() or "S/N"
        fecha = datetime.now().strftime('%d/%m/%Y')

        trs = ""
        for c in comps_data:
            nombre = c["nombre"]
            estado = c["estado"].get()
            obs = c["obs"].get().strip()

            iconOK = "☑" if estado == "OK" else "□"
            iconOBS = "☑" if estado == "OBS" else "□"
            iconMalo = "☑" if estado == "MALO" else "□"

            trs += f"""<tr>
              <td>{nombre}</td>
              <td style="text-align: center;">{iconOK} OK &nbsp;&nbsp;&nbsp; {iconOBS} OBS &nbsp;&nbsp;&nbsp; {iconMalo} MALO</td>
              <td>{obs}</td>
            </tr>"""

        html_etiqueta = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Etiqueta_{serial}</title>
            <style>
              @media print {{ @page {{ margin: 15mm; size: auto; }} body {{ margin: 0; }} }}
              body {{ font-family: 'Arial', sans-serif; font-size: 11px; color: #000; max-width: 700px; margin: auto; padding: 20px; line-height: 1.4; }}
              .header-title {{ display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 20px; font-size: 13px; }}
              .row {{ margin-bottom: 10px; font-size: 11px;}}
              .row span {{ margin-right: 25px; }}
              table {{ width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 20px; }}
              th, td {{ border: 1px solid #000; padding: 5px 8px; text-align: left; }}
              .checkboxes {{ font-family: 'Segoe UI Symbol', sans-serif; }}
              .bold {{ font-weight: bold; }}
            </style>
        </head>
        <body onload="setTimeout(() => {{ window.print(); }}, 500);">
            <div class="header-title">
              <span>ETIQUETAS NOTEBOOK</span>
              <span>FECHA: {fecha}</span>
            </div>
            <div class="row">
              <span><span class="bold">MODELO:</span> {modelo}</span>
              <span><span class="bold">OT:</span> {ot}</span>
              <span><span class="bold">PROCESADOR:</span> {cpu}</span>
              <span><span class="bold">GUIA:</span> {guia}</span>
              <span><span class="bold">SERIAL:</span> {serial}</span>
            </div>
            <div class="row checkboxes" style="margin-bottom:20px;">
              <span class="bold">ESTADO GLOBAL:</span> &nbsp;&nbsp;&nbsp;&nbsp;
              □ BUENA &nbsp;&nbsp;&nbsp; ☑ POR REPARAR &nbsp;&nbsp;&nbsp; □ MALA &nbsp;&nbsp;&nbsp; □ REPUESTO
            </div>
            <table>
              <thead>
                <tr>
                  <th width="20%">REVISIÓN DE COMPONENTES</th>
                  <th width="35%" style="text-align: center;">ESTADO</th>
                  <th width="45%">OBSERVACIÓN</th>
                </tr>
              </thead>
              <tbody class="checkboxes">
                {trs}
              </tbody>
            </table>
            <div class="row checkboxes bold" style="margin-top:20px;">
              CHECKLIST: &nbsp;&nbsp;&nbsp;&nbsp; □ PRUEBA S/O &nbsp;&nbsp;&nbsp;&nbsp; □ HYDRA &nbsp;&nbsp;&nbsp;&nbsp; □ HARDWARE DEFAULT
            </div>
        </body>
        </html>
        """

        temp_path = os.path.join(tempfile.gettempdir(), "etiqueta_manual_temp.html")
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write(html_etiqueta)

        # as_uri() y no f"file://{ruta}": en Windows la ruta empieza con "C:\"
        # y concatenarla deja "file://C:\..." , donde el navegador toma "C:"
        # como nombre de host y la etiqueta no abre.
        webbrowser.open(Path(temp_path).as_uri())

    tk.Button(
        win,
        text="🖨️ GENERAR E IMPRIMIR ETIQUETA",
        command=generar_y_imprimir,
        bg="#0078d4",
        fg="white",
        font=("Segoe UI", 10, "bold"),
        padx=20,
        pady=8,
        cursor="hand2",
        relief="flat"
    ).pack(pady=10)


# ============================================================================
#  Carga masiva de transformadores (cargadores) sobre los equipos de la OT
# ============================================================================
# El transformador se etiqueta aparte del equipo y casi nunca aparece en el
# mismo momento: llega después y en bloque, como una columna de Excel o como
# una ráfaga de la pistola. De ahí que sea su propio módulo y no un campo del
# formulario de escaneo.
#
# La lista va en ORDEN DE ESCANEO (el más viejo arriba, al revés que el resto
# de la app): es el orden en que se etiquetaron los equipos, así que es el que
# trae la columna pegada y el pegado cae 1-a-1 sobre las filas.
_TRANSF_PENDIENTE = "PENDIENTE"


def _normalizar_sn_transf(valor):
    """Serial de transformador utilizable, o '' si la celda quedó vacía."""
    limpio = " ".join(str(valor or "").split())
    return "" if limpio.upper() in ("", _TRANSF_PENDIENTE) else limpio


def _seriales_del_portapapeles(texto):
    """
    Un serial por línea. Si la línea trae tabs (columna copiada de Excel con
    celdas vecinas) se usa la primera celda no vacía: pegar una columna nunca
    debe convertirse en varios transformadores para el mismo equipo.
    """
    seriales = []
    for linea in texto.replace("\r", "\n").split("\n"):
        celdas = (" ".join(c.split()) for c in linea.split("\t"))
        celda = next((c for c in celdas if c), "")
        if celda:
            seriales.append(celda)
    return seriales


def abrir_modulo_agregar_transformador(ventana_padre):
    # Los chequeos van ANTES de crear la ventana: abrir un modal para cerrarlo
    # de inmediato hace parpadear la pantalla y roba el foco por un instante.
    lote = cargar_lote_activo()
    if not lote:
        messagebox.showwarning(
            "Sin OT activa",
            "Los transformadores se asignan a los equipos de una OT.\n\n"
            "Abrí o creá una OT antes de cargarlos.",
            parent=ventana_padre,
        )
        return

    # Orden de escaneo explícito (el más viejo primero). No alcanza con dar
    # vuelta la lista: ESCANEADO tiene resolución de segundos, así que dos
    # equipos del mismo segundo empatan y el reverso los deja invertidos entre
    # sí. Ordenar ascendente respeta el empate en vez de darlo vuelta.
    registros = sorted(
        leer_equipos_lote_con_ruta(lote), key=lambda par: _clave_orden(par[1])
    )
    if not registros:
        messagebox.showwarning(
            "OT sin equipos",
            f"{etiqueta_lote(lote)} todavía no tiene equipos escaneados.",
            parent=ventana_padre,
        )
        return

    win = ventana_modal(ventana_padre, "Asignar Transformadores", "780x640")
    titulo_ui(win, "🔌 Asignar Transformadores", size=14, pady=(18, 2))
    tk.Label(
        win,
        text=f"{etiqueta_lote(lote)}  ·  {len(registros)} equipo(s)",
        font=fuente(9, True),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack()
    tk.Label(
        win,
        text="Los equipos están en orden de escaneo. Pegá la columna de seriales "
        "o escaneá con la pistola (Enter salta a la fila siguiente).",
        font=fuente(8),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
        wraplength=720,
    ).pack(pady=(2, 8))

    # --- Grilla scrolleable: una OT puede traer 50+ equipos ---
    marco, grilla, canvas = marco_scrolleable(win)
    marco.pack(fill="both", expand=True, padx=18)

    for col, (txt, ancho) in enumerate(
        (("#", 3), ("MODELO", 26), ("SERIAL EQUIPO", 20), ("SERIE TRANSFORMADOR", 24))
    ):
        tk.Label(
            grilla,
            text=txt,
            width=ancho,
            anchor="w",
            font=fuente(8, True),
            bg=COLORS["fondo"],
            fg=COLORS["gris_oscuro"],
        ).grid(row=0, column=col, padx=4, pady=(0, 4), sticky="w")

    filas = []
    for i, (ruta, jdata) in enumerate(registros, start=1):
        data = jdata.get("DATA", {})
        actual = jdata.get("SN_Transf", _TRANSF_PENDIENTE) or _TRANSF_PENDIENTE
        fondo = COLORS["fondo"] if i % 2 else "#e3e9f2"

        for col, txt in (
            (0, str(i)),
            (1, data.get("model", jdata.get("MODELO", "Desconocido"))),
            (2, jdata.get("SERIAL", "?")),
        ):
            tk.Label(
                grilla,
                text=txt,
                anchor="w",
                font=fuente(9),
                bg=fondo,
                fg=COLORS["gris_oscuro"],
            ).grid(row=i, column=col, padx=4, pady=1, sticky="we")

        ent = tk.Entry(grilla, font=fuente(9), width=24)
        ent.grid(row=i, column=3, padx=4, pady=1, sticky="w")
        # Lo ya cargado se muestra: este módulo se usa varias veces sobre la
        # misma OT, a medida que van llegando los cargadores.
        if actual != _TRANSF_PENDIENTE:
            ent.insert(0, actual)
        filas.append({"ruta": ruta, "jdata": jdata, "entry": ent})

    grilla.grid_columnconfigure(1, weight=1)

    lbl_estado = tk.Label(
        win, text="", font=fuente(9, True), bg=COLORS["fondo"], fg=COLORS["azul"]
    )
    lbl_estado.pack(pady=(8, 0))

    def refrescar_estado():
        cargados = sum(1 for f in filas if _normalizar_sn_transf(f["entry"].get()))
        lbl_estado.config(
            text=f"{cargados} de {len(filas)} equipo(s) con transformador"
        )

    # Enter = siguiente fila: la pistola manda Enter al final de cada lectura,
    # así se escanea la OT completa sin tocar el mouse.
    for idx, f in enumerate(filas):
        siguiente = filas[idx + 1]["entry"] if idx + 1 < len(filas) else None

        # sig se captura por default: sin eso, el closure se quedaría con la
        # última fila del bucle y todos los Enter saltarían al mismo Entry.
        def salto(event, sig=siguiente):
            refrescar_estado()
            if sig:
                sig.focus_set()
                canvas.yview_scroll(1, "units")
            return "break"

        f["entry"].bind("<Return>", salto)
        f["entry"].bind("<FocusOut>", lambda e: refrescar_estado())

    def pegar_portapapeles():
        try:
            crudo = win.clipboard_get()
        except Exception:
            messagebox.showwarning(
                "Portapapeles vacío",
                "No hay texto para pegar. Copiá primero la columna de seriales.",
                parent=win,
            )
            return

        seriales = _seriales_del_portapapeles(crudo)
        if not seriales:
            messagebox.showwarning(
                "Nada que pegar",
                "El portapapeles no trae ningún serial reconocible.",
                parent=win,
            )
            return

        # El pegado es posicional y sobrescribe desde la primera fila: si ya
        # había datos cargados a mano, se avisa antes de taparlos.
        if any(_normalizar_sn_transf(f["entry"].get()) for f in filas):
            if not messagebox.askyesno(
                "Sobrescribir",
                "Ya hay transformadores cargados en la lista.\n\n"
                f"Pegar reemplaza las primeras {min(len(seriales), len(filas))} "
                "fila(s) en orden. ¿Continuar?",
                parent=win,
            ):
                return

        # zip corta por el más corto: si sobran seriales o sobran filas, no
        # revienta ni desalinea, simplemente llega hasta donde alcanza.
        for f, sn in zip(filas, seriales):
            f["entry"].delete(0, tk.END)
            f["entry"].insert(0, sn)
        refrescar_estado()

        if len(seriales) > len(filas):
            messagebox.showwarning(
                "Sobran seriales",
                f"Pegaste {len(seriales)} seriales pero la OT tiene "
                f"{len(filas)} equipos.\n\n"
                f"Se usaron los primeros {len(filas)} y se descartó el resto.",
                parent=win,
            )

    def limpiar():
        if not messagebox.askyesno(
            "Limpiar",
            "¿Vaciar la columna de transformadores?\n\n"
            "Todavía no se guarda nada: los JSON quedan como están hasta que "
            "aprietes GUARDAR.",
            parent=win,
        ):
            return
        for f in filas:
            f["entry"].delete(0, tk.END)
        refrescar_estado()

    def guardar():
        cambios, vistos, duplicados = [], {}, []
        for f in filas:
            serial_eq = f["jdata"].get("SERIAL", "?")
            nuevo = _normalizar_sn_transf(f["entry"].get()) or _TRANSF_PENDIENTE
            if nuevo != _TRANSF_PENDIENTE:
                clave = nuevo.upper()
                if clave in vistos:
                    duplicados.append(f"{nuevo} → {vistos[clave]} y {serial_eq}")
                vistos[clave] = serial_eq
            actual = f["jdata"].get("SN_Transf", _TRANSF_PENDIENTE) or _TRANSF_PENDIENTE
            # Vaciar una celda es un cambio válido: devuelve el equipo a
            # PENDIENTE cuando el transformador se cargó por error.
            if nuevo != actual:
                cambios.append((f, nuevo))

        # Un transformador no puede estar en dos equipos: casi siempre es una
        # columna pegada corrida una fila, así que conviene frenar y mirar.
        if duplicados and not messagebox.askyesno(
            "Transformadores repetidos",
            "Estos seriales quedaron en más de un equipo:\n\n"
            + "\n".join(duplicados[:8])
            + "\n\n¿Guardar igual?",
            parent=win,
        ):
            return

        if not cambios:
            messagebox.showinfo(
                "Sin cambios", "No hay nada nuevo para guardar.", parent=win
            )
            return

        errores = []
        for f, nuevo in cambios:
            # Se relee del disco por si el equipo se re-escaneó mientras esta
            # ventana estaba abierta: así se pisa solo SN_Transf y no se
            # revierte el resto del registro a como estaba al abrir.
            try:
                with open(f["ruta"], "r", encoding="utf-8") as fh:
                    jdata = json.load(fh)
                if not isinstance(jdata, dict):
                    jdata = f["jdata"]
            except Exception:
                jdata = f["jdata"]
            jdata["SN_Transf"] = nuevo
            try:
                with open(f["ruta"], "w", encoding="utf-8") as fh:
                    json.dump(jdata, fh, ensure_ascii=False, indent=4)
                f["jdata"] = jdata
            except Exception as e:
                errores.append(f"{jdata.get('SERIAL', '?')}: {e}")

        marcar_pendiente_envio(lote)
        regenerar_lote(lote)

        if errores:
            messagebox.showerror(
                "Guardado parcial",
                f"Se guardaron {len(cambios) - len(errores)} de {len(cambios)} "
                "equipo(s).\n\nFallaron:\n" + "\n".join(errores[:8]),
                parent=win,
            )
            return

        messagebox.showinfo(
            "✅ Transformadores guardados",
            f"{len(cambios)} equipo(s) actualizado(s) en {etiqueta_lote(lote)}.\n\n"
            "El HTML y el JSON FUSIONADO ya se regeneraron: los seriales salen "
            "en 'Copiar Filas' y en el CSV de Valida.",
            parent=win,
        )
        win.destroy()

    barra_bts = tk.Frame(win, bg=COLORS["fondo"])
    barra_bts.pack(pady=12)
    for txt, cmd, color in (
        ("📋 Pegar del portapapeles", pegar_portapapeles, COLORS["azul_btn"]),
        ("🧹 Limpiar", limpiar, COLORS["pizarra"]),
        ("💾 GUARDAR", guardar, COLORS["verde"]),
        ("Cancelar", win.destroy, COLORS["gris"]),
    ):
        tk.Button(
            barra_bts,
            text=txt,
            command=cmd,
            bg=color,
            fg="white",
            font=fuente(9, True),
            padx=14,
            pady=7,
            cursor="hand2",
            relief="flat",
        ).pack(side="left", padx=5)

    refrescar_estado()
    filas[0]["entry"].focus_set()
    win.wait_window()



def abrir_panel_clonacion(ventana_padre):
    """
    Panel central del flujo de clonación: qué OT están tomando los equipos
    clonados, cuáles ya llegaron, cuáles faltan, y el cierre del lote.

    Todo el estado vive en el servidor, no acá: esta ventana lo lee y lo
    escribe, pero cualquier equipo puede abrirla y ver lo mismo.
    """
    win = ventana_modal(ventana_padre, "Panel de clonación — OT activa", "980x680")
    titulo_ui(win, "🛰️ Panel de clonación")

    estado = {"lote": None, "llegados": [], "esperados": []}

    # --- Encabezado: qué OT está publicada ---
    lbl_ot = tk.Label(
        win, text="Consultando el servidor…", font=fuente(11, True),
        bg="#fef3c7", fg="#92400e", pady=10,
    )
    lbl_ot.pack(fill="x", padx=20, pady=(0, 8))

    lbl_resumen = tk.Label(
        win, text="", font=fuente(10, True), bg=COLORS["fondo"], fg=COLORS["gris_oscuro"]
    )
    lbl_resumen.pack(pady=(0, 6))

    tree = treeview_auditoria(
        win,
        [
            ("estado", 90, "Estado"),
            ("serial", 170, "N° de Serie"),
            ("modelo", 260, "Modelo"),
            ("fecha", 140, "Escaneado"),
            ("obs", 260, "Observaciones"),
        ],
        height=15,
    )

    def _con_cliente(tarea, titulo_error="Error"):
        """Abre la conexión, corre la tarea y la cierra siempre."""
        try:
            cliente = abrir_conexion_scp()
        except Exception as e:
            messagebox.showerror(titulo_error, str(e), parent=win)
            return None
        try:
            return tarea(cliente)
        except Exception as e:
            messagebox.showerror(titulo_error, str(e), parent=win)
            return None
        finally:
            try:
                cliente.close()
            except Exception:
                pass

    def refrescar():
        def tarea(cliente):
            lote = leer_lote_activo_remoto(cliente)
            if not lote:
                return None, [], []
            return lote, equipos_remotos(cliente, lote), leer_esperados(cliente, lote)

        datos = _con_cliente(tarea, "No se pudo consultar el servidor")
        if datos is None:
            return
        lote, llegados, esperados = datos
        estado.update(lote=lote, llegados=llegados, esperados=esperados)

        for fila in tree.get_children():
            tree.delete(fila)

        if not lote:
            lbl_ot.config(
                text="⚠️  No hay ninguna OT publicada — los equipos clonados no "
                     "tienen dónde mandar sus datos",
                bg="#fef3c7", fg="#92400e",
            )
            lbl_resumen.config(text="")
            return

        lbl_ot.config(
            text=f"📍 OT publicada: {etiqueta_lote(lote)}",
            bg="#dcfce7", fg="#166534",
        )

        recibidos = {}
        for eq in llegados:
            recibidos[str(eq.get("SERIAL", "")).strip().upper()] = eq

        for sn, eq in recibidos.items():
            marca = "esperado" if not esperados or sn in [
                s.upper() for s in esperados
            ] else "sobra"
            tree.insert(
                "", "end",
                values=(
                    "✓ llegó" if marca == "esperado" else "⚠ no esperado",
                    eq.get("SERIAL", ""),
                    eq.get("MODELO", "") or (eq.get("DATA") or {}).get("model", ""),
                    (eq.get("DATA") or {}).get("fecha", "") or eq.get("ESCANEADO", ""),
                    eq.get("OBS", ""),
                ),
                tags=("ok" if marca == "esperado" else "sobra",),
            )

        faltan = [s for s in esperados if s.upper() not in recibidos]
        for sn in faltan:
            tree.insert(
                "", "end",
                values=("✗ falta", sn, "—", "—", "todavía no envió"),
                tags=("falta",),
            )

        if esperados:
            lbl_resumen.config(
                text=f"Llegaron {len(recibidos)} de {len(esperados)} esperados  ·  "
                     f"faltan {len(faltan)}"
            )
        else:
            lbl_resumen.config(
                text=f"{len(recibidos)} equipo(s) recibidos  ·  "
                     "sin lista esperada cargada"
            )

    def publicar():
        lote = cargar_lote_activo()
        if not lote:
            messagebox.showwarning(
                "Sin OT activa",
                "No hay una OT activa en este equipo.\n\n"
                "Creá o cambiá de OT desde el menú y volvé a publicar.",
                parent=win,
            )
            return
        if not messagebox.askyesno(
            "Publicar OT",
            f"¿Publicar {etiqueta_lote(lote)} como OT activa?\n\n"
            "Todos los equipos que se clonen a partir de ahora van a mandar "
            "sus datos a esta OT.",
            parent=win,
        ):
            return
        if _con_cliente(
            lambda c: publicar_lote_activo(c, lote) or True, "No se pudo publicar"
        ):
            messagebox.showinfo("Publicada", f"{etiqueta_lote(lote)}", parent=win)
            refrescar()

    def cargar_esperados():
        lote = estado["lote"]
        if not lote:
            messagebox.showwarning("Sin OT", "Primero publicá una OT.", parent=win)
            return
        dlg = ventana_modal(win, "Seriales esperados", "460x520")
        titulo_ui(dlg, "Seriales que deben llegar", size=12)
        tk.Label(
            dlg, text="Uno por línea. Pegá la lista de equipos a clonar.",
            font=fuente(9), bg=COLORS["fondo"], fg=COLORS["gris"],
        ).pack(pady=(0, 6))
        caja = tk.Text(dlg, font=fuente(9), relief="solid", bd=1)
        caja.pack(fill="both", expand=True, padx=16)
        caja.insert("1.0", "\n".join(estado["esperados"]))

        def guardar():
            seriales = [l.strip() for l in caja.get("1.0", "end").splitlines()]
            seriales = [s for s in seriales if s]
            guardados = _con_cliente(
                lambda c: guardar_esperados(c, lote, seriales), "No se pudo guardar"
            )
            if guardados is not None:
                dlg.destroy()
                messagebox.showinfo(
                    "Lista guardada", f"{len(guardados)} serial(es).", parent=win
                )
                refrescar()

        boton_accion(dlg, "💾 Guardar lista", guardar, COLORS["verde"]).pack(pady=10)

    def traer_y_reconstruir():
        lote = estado["lote"]
        if not lote:
            messagebox.showwarning("Sin OT", "No hay OT publicada.", parent=win)
            return
        res = _con_cliente(
            lambda c: descargar_equipos_remotos(c, lote), "No se pudo traer"
        )
        if res is None:
            return
        bajados, total = res
        messagebox.showinfo(
            "Reconstrucción lista",
            f"Se trajeron {bajados} JSON del servidor.\n\n"
            f"La OT quedó con {total} equipo(s) y su HTML y JSON FUSIONADO "
            f"ya se regeneraron en:\n{lote['ruta']}",
            parent=win,
        )

    def cerrar_remoto():
        lote = estado["lote"]
        if not lote:
            messagebox.showwarning("Sin OT", "No hay OT publicada.", parent=win)
            return
        faltan = [
            s for s in estado["esperados"]
            if s.upper() not in {
                str(e.get("SERIAL", "")).strip().upper() for e in estado["llegados"]
            }
        ]
        aviso = ""
        if faltan:
            aviso = (
                f"\n\n⚠️ Todavía faltan {len(faltan)} equipo(s):\n"
                + "\n".join(f"  • {s}" for s in faltan[:10])
                + ("\n  …" if len(faltan) > 10 else "")
                + "\n\nSi cerrás ahora, esos equipos ya no van a poder enviar."
            )
        if not messagebox.askyesno(
            "Cerrar lote",
            f"¿Cerrar {etiqueta_lote(lote)} en el servidor?\n\n"
            f"Recibidos: {len(estado['llegados'])} equipo(s).{aviso}",
            parent=win,
        ):
            return
        if _con_cliente(
            lambda c: cerrar_lote_remoto(c, lote, len(estado["llegados"])) or True,
            "No se pudo cerrar",
        ):
            messagebox.showinfo(
                "Lote cerrado",
                "La OT quedó cerrada en el servidor y ya no es la OT activa.",
                parent=win,
            )
            refrescar()

    frame_btns = tk.Frame(win, bg=COLORS["fondo"])
    frame_btns.pack(pady=12)
    for texto, comando, color in (
        ("🔄 Refrescar", refrescar, COLORS["celeste"]),
        ("📤 Publicar OT activa", publicar, COLORS["azul"]),
        ("📋 Seriales esperados", cargar_esperados, COLORS["pizarra"]),
        ("⬇️ Traer y reconstruir", traer_y_reconstruir, COLORS["morado"]),
        ("✅ Cerrar lote", cerrar_remoto, COLORS["verde"]),
    ):
        boton_accion(frame_btns, texto, comando, color).pack(side="left", padx=4)

    refrescar()


def abrir_modulo_cambiar_nombre(ventana_padre):
    # 1. Obtener el número de serie de la BIOS
    try:
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = 0

        cmd = ["powershell", "-NoProfile", "-Command", "(Get-CimInstance Win32_Bios).SerialNumber"]
        serial = subprocess.check_output(cmd, text=True, startupinfo=si).strip()
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo obtener el número de serie de la BIOS.\n{e}", parent=ventana_padre)
        return

    if not serial or serial.lower() in ("to be filled by o.e.m.", "default string"):
        messagebox.showwarning("Aviso", "No se detectó un número de serie válido en la BIOS de este equipo.", parent=ventana_padre)
        return

    # 2. Limpiar y formatear el nombre (Máximo 15 caracteres para NetBIOS en Windows)
    serial_limpio = re.sub(r"[^a-zA-Z0-9\-]", "", serial)
    nuevo_nombre = f"ARR-{serial_limpio}"

    if len(nuevo_nombre) > 15:
        nuevo_nombre = nuevo_nombre[:15] # Truncar a 15 caracteres si es muy largo

    respuesta = messagebox.askyesno(
        "Confirmar Cambio de Nombre",
        f"El número de serie detectado es: {serial}\n\n"
        f"¿Estás seguro de cambiar el nombre de este equipo a:\n\n{nuevo_nombre}?\n\n"
        "(Debes aceptar la ventana de permisos de Administrador que aparecerá a continuación).",
        parent=ventana_padre
    )

    if not respuesta:
        return

    # 3. Aplicar el cambio de nombre
    try:
        ps_command = f'Rename-Computer -NewName "{nuevo_nombre}"'
        cmd = [
            "powershell",
            "-NoProfile",
            "-Command",
            f"Start-Process powershell -ArgumentList '-NoProfile -WindowStyle Hidden -Command {ps_command}' -Verb RunAs"
        ]

        subprocess.run(cmd, check=True)

        messagebox.showinfo(
            "✅ Solicitud Enviada",
            f"Se ha ordenado el cambio de nombre a '{nuevo_nombre}'.\n\nPor favor, REINICIA EL EQUIPO manualmente para que los cambios surtan efecto.",
            parent=ventana_padre
        )

    except subprocess.CalledProcessError:
        messagebox.showerror(
            "Error",
            "No se pudo cambiar el nombre del equipo.\n\nAsegúrate de aceptar la ventana de permisos de Administrador.",
            parent=ventana_padre
        )
    except Exception as e:
        messagebox.showerror(
            "Error Inesperado",
            f"Ocurrió un error al intentar cambiar el nombre: {e}",
            parent=ventana_padre
        )








# ─────────────────────────────────────────────────────────
#  Navegación e Interfaz Principal
# ─────────────────────────────────────────────────────────
# ============================================================================
# 10. UI PRINCIPAL — navegación en la ventana raíz: menú, pantalla de escaneo
#     y formulario de registro.
# ============================================================================
def limpiar_ventana(ventana):
    for widget in ventana.winfo_children():
        widget.destroy()


def _seccion_menu(padre, titulo):
    """Encabezado de grupo del menú: línea divisoria + rótulo chico."""
    ttk.Separator(padre, orient="horizontal").pack(fill="x", pady=(14, 5))
    tk.Label(
        padre,
        text=titulo.upper(),
        font=fuente(8, True),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(anchor="w", pady=(0, 2))


def mostrar_menu_principal(ventana):
    limpiar_ventana(ventana)
    ventana.geometry("700x800")

    titulo_ui(ventana, "🖥️ Sistema de Control de Inventario", size=16, pady=(24, 8))

    # Los equipos escaneados van SIEMPRE al lote activo. Mostrarlo grande y
    # permanente evita el error caro: escanear contra la OT equivocada.
    lote = cargar_lote_activo()
    abiertos = listar_lotes_abiertos()

    if lote:
        equipos = contar_equipos(lote["ruta"])
        tk.Label(
            ventana,
            text=f"📍 OT ACTIVA — {etiqueta_lote(lote)}\n{equipos} equipo(s) escaneado(s)",
            font=fuente(10, True),
            bg="#dcfce7",
            fg="#166534",
            justify="center",
            padx=12,
            pady=8,
        ).pack(fill="x", padx=40, pady=(0, 4))
    else:
        tk.Label(
            ventana,
            text=(
                "⚠️  Sin OT activa\n"
                "Crea una OT o activa una en curso antes de escanear."
            ),
            font=fuente(10, True),
            bg="#fef3c7",
            fg="#92400e",
            justify="center",
            padx=12,
            pady=8,
        ).pack(fill="x", padx=40, pady=(0, 4))

    otras = [
        l
        for l in abiertos
        if not lote
        or os.path.normcase(l["ruta"]) != os.path.normcase(lote["ruta"])
    ]
    if otras:
        resumen = ", ".join(
            f"{o['cliente'].replace('_', ' ')} {carpeta_lote(o['mov'], o['numero'])}"
            f" ({o['equipos']})"
            for o in otras[:3]
        )
        if len(otras) > 3:
            resumen += f" y {len(otras) - 3} más"
        tk.Label(
            ventana,
            text=f"También abiertas: {resumen}",
            font=fuente(8),
            bg=COLORS["fondo"],
            fg=COLORS["gris"],
            wraplength=430,
            justify="center",
        ).pack(pady=(0, 8))
    else:
        tk.Label(ventana, text="", bg=COLORS["fondo"]).pack(pady=(0, 4))

    # El menú va con scroll: la lista de acciones creció hasta rozar el alto de
    # la ventana y sin esto el próximo botón que se agregue queda cortado, sin
    # forma de llegar a él.
    contenedor, frame_btns, _ = marco_scrolleable(ventana)
    contenedor.pack(fill="both", expand=True, padx=40)

    # El color agrupa, no decora: celeste = manejo de la OT, verde = cerrarla,
    # naranja = sale de esta máquina, morado = reparación, pizarra = utilitario
    # suelto. Dos botones del mismo color hacen cosas del mismo tipo.
    #
    # Equipos del modelo anterior (sueltos, sin carpeta de OT): hay que
    # adoptarlos en una OT para que entren al fusionado.
    legacy = sum(len(v) for v in pendientes_legacy().values())
    if legacy:
        boton_menu(
            frame_btns,
            f"⚠️ MIGRAR {legacy} EQUIPO(S) DEL FORMATO ANTERIOR",
            lambda: accion_migrar_legacy(ventana),
            "#d97706",
            size=10,
        )

    boton_menu(
        frame_btns,
        "🔍 ESCANEAR ESTE EQUIPO\n(Crear registro)",
        lambda: iniciar_escaneo(ventana),
        COLORS["azul_btn"],
        size=12,
    )

    _seccion_menu(frame_btns, "Orden de trabajo")

    frame_ot = tk.Frame(frame_btns, bg=COLORS["fondo"])
    frame_ot.pack(fill="x", pady=2)
    for texto, comando, lado in (
        ("➕ NUEVA OT", lambda: accion_nueva_ot(ventana), (0, 3)),
        (
            f"🔄 CAMBIAR OT ({len(abiertos)})",
            lambda: accion_cambiar_lote(ventana),
            (3, 0),
        ),
    ):
        tk.Button(
            frame_ot,
            text=texto,
            command=comando,
            bg=COLORS["celeste"],
            fg="white",
            font=fuente(10, True),
            pady=8,
            cursor="hand2",
            relief="flat",
        ).pack(side="left", fill="x", expand=True, padx=lado)

    boton_menu(
        frame_btns,
        "✅ CERRAR OT",
        lambda: accion_cerrar_lote(ventana),
        COLORS["verde"],
    )

    _seccion_menu(frame_btns, "Archivo y respaldo")
    boton_menu(
        frame_btns,
        f"🚀 ENVIAR OT AL SERVIDOR (SCP · {SCP_HOST})",
        lambda: accion_enviar_red(ventana),
        COLORS["naranja"],
        size=10,
    )
    boton_menu(
        frame_btns,
        "🔧 RECONSTRUIR / REPARAR OT DESDE JSONs",
        lambda: accion_unir_Json(ventana),
        COLORS["morado"],
        size=10,
    )
    boton_menu(
        frame_btns,
        "🛰️ PANEL DE CLONACIÓN (OT activa y seguimiento)",
        lambda: abrir_panel_clonacion(ventana),
        COLORS["verde_esmeralda"],
        size=10,
    )

    _seccion_menu(frame_btns, "Herramientas")
    for texto, comando in (
        ("🔌 AGREGAR TRANSFORMADORES", lambda: abrir_modulo_agregar_transformador(ventana)),
        ("🏷️ GENERAR ETIQUETA MANUAL (Equipos Malos)", lambda: abrir_modulo_etiquetas_manual(ventana)),
        ("💻 CAMBIAR NOMBRE DEL EQUIPO", lambda: abrir_modulo_cambiar_nombre(ventana)),
    ):
        boton_menu(frame_btns, texto, comando, COLORS["pizarra"], size=9)


def iniciar_escaneo(ventana):
    # Sin OT activa no se escanea: el equipo no tendría dónde caer. Se ofrece
    # crearla en el momento para no cortar el flujo.
    lote = cargar_lote_activo()
    if not lote:
        if not messagebox.askyesno(
            "Sin OT activa",
            "No hay ninguna OT activa y el equipo escaneado necesita una.\n\n"
            "¿Crear o elegir una ahora?",
        ):
            return
        lote = dialogo_nuevo_lote(ventana)
        if not lote:
            return
        guardar_lote_activo(lote)

    limpiar_ventana(ventana)
    ventana.geometry("750x850")

    frame_carga = tk.Frame(ventana, bg=COLORS["fondo"])
    frame_carga.place(relx=0.5, rely=0.5, anchor="center")

    tk.Label(frame_carga, text="⏳", font=fuente(48), bg=COLORS["fondo"]).pack()
    lbl_scan = tk.Label(
        frame_carga,
        text="Escaneando hardware...",
        font=fuente(15, True),
        bg=COLORS["fondo"],
        fg=COLORS["azul"],
    )
    lbl_scan.pack(pady=6)
    tk.Label(
        frame_carga,
        text="Esto tomará unos segundos.",
        font=fuente(10),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack()

    dots = ["", ".", "..", "..."]
    dot_idx = [0]

    def animar():
        if frame_carga.winfo_exists():
            lbl_scan.config(text=f"Escaneando hardware{dots[dot_idx[0] % 4]}")
            dot_idx[0] += 1
            ventana.after(450, animar)

    animar()

    resultado = [None]

    def escanear():
        resultado[0] = get_inventory_fast()
        ventana.after(0, lambda: construir_ui_formulario(ventana, resultado[0], lote))

    threading.Thread(target=escanear, daemon=True).start()


def construir_ui_formulario(ventana, data, lote):
    limpiar_ventana(ventana)

    if not data:
        messagebox.showerror("Error", "No se pudo obtener información del hardware.")
        mostrar_menu_principal(ventana)
        return

    top_frame = tk.Frame(ventana, bg=COLORS["fondo"])
    top_frame.pack(fill="x", pady=(14, 4), padx=30)

    tk.Button(
        top_frame,
        text="⬅ Volver",
        command=lambda: mostrar_menu_principal(ventana),
        bg="#cbd5e1",
        fg=COLORS["gris_oscuro"],
        font=("Segoe UI", 9, "bold"),
        cursor="hand2",
        relief="flat",
        padx=10,
    ).pack(side="left")
    tk.Label(
        top_frame,
        text="💻 Asistente de Entrega / Retiro",
        font=("Segoe UI", 15, "bold"),
        bg=COLORS["fondo"],
        fg=COLORS["azul"],
    ).pack(side="left", padx=50)

    info_frame = tk.Frame(ventana, bg="#dbeafe", bd=1, relief="solid")
    info_frame.pack(padx=30, fill="x")
    tk.Label(
        info_frame,
        text=f"  {data['model']}   |   S/N: {data['serial']}  ",
        font=("Consolas", 10),
        bg="#dbeafe",
        fg="#1d4ed8",
        pady=6,
    ).pack()

    marco = tk.Frame(ventana, bg=COLORS["fondo"])
    marco.pack(pady=10, fill="x", padx=30)

    # El movimiento y el número ya no se eligen por equipo: los define la OT
    # activa. Se muestran fijos para que siempre sea evidente dónde va a caer
    # lo que se está escaneando.
    cfg_mov = _MOV_CFG[lote["mov"]]
    lf_log = ttk.LabelFrame(marco, text=" 🚚 Destino de este equipo ")
    lf_log.pack(fill="x", pady=(0, 10))
    tk.Label(
        lf_log,
        text=f"📍 {etiqueta_lote(lote)}",
        bg=COLORS["fondo"],
        fg="#166534",
        font=("Segoe UI", 11, "bold"),
    ).pack(pady=(6, 0))
    tk.Label(
        lf_log,
        text=f"{cfg_mov['etiqueta_campo']}: {lote['numero']}\n"
        f"{cfg_mov['etiqueta_campo_sec']}: {lote.get('numero_sec') or '—'}",
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
        font=("Segoe UI", 9),
    ).pack(pady=(0, 6))

    lf_fallas = ttk.LabelFrame(marco, text=" ⚠️ Revisión de Componentes (Solo Retiros) ")

    componentes_etiqueta = [
        "Carcaza",
        "Pantalla",
        "Teclado",
        "Touchpad",
        "Puertos USB",
        "Puertos Video",
        "Ethernet/Wi-Fi",
        "Placa base",
        "Memoria",
        "Disco duro",
        "Batería",
    ]

    comps_data = []

    frame_grilla = tk.Frame(lf_fallas, bg=COLORS["fondo"])
    frame_grilla.pack(padx=10, pady=10)

    tk.Label(
        frame_grilla, text="Componente", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]
    ).grid(row=0, column=0, sticky="w", padx=5)
    tk.Label(
        frame_grilla, text="Estado", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]
    ).grid(row=0, column=1, padx=5)
    tk.Label(
        frame_grilla,
        text="Observación (Obligatorio si es OBS/MALO)",
        font=("Segoe UI", 9, "bold"),
        bg=COLORS["fondo"],
    ).grid(row=0, column=2, sticky="w", padx=5)

    for i, comp_name in enumerate(componentes_etiqueta, start=1):
        tk.Label(frame_grilla, text=comp_name, bg=COLORS["fondo"], font=("Segoe UI", 9)).grid(
            row=i, column=0, sticky="w", padx=5, pady=2
        )

        cmb_estado = ttk.Combobox(
            frame_grilla,
            values=["OK", "OBS", "MALO"],
            width=6,
            state="readonly",
            font=("Segoe UI", 9),
        )
        cmb_estado.set("OK")
        cmb_estado.grid(row=i, column=1, padx=5, pady=2)

        ent_obs = tk.Entry(
            frame_grilla,
            width=40,
            font=("Segoe UI", 9),
            state=tk.DISABLED,
            bg="#f1f5f9",
        )
        ent_obs.grid(row=i, column=2, padx=5, pady=2, sticky="w")

        def toggle_entry(event, e=ent_obs, c=cmb_estado):
            if c.get() != "OK":
                e.config(state=tk.NORMAL, bg="#ffffff")
                e.focus()
            else:
                e.delete(0, tk.END)
                e.config(state=tk.DISABLED, bg="#f1f5f9")

        cmb_estado.bind("<<ComboboxSelected>>", toggle_entry)

        comps_data.append({"nombre": comp_name, "estado": cmb_estado, "obs": ent_obs})

    # La revisión de componentes solo aplica a retiros, y el movimiento ya no
    # cambia dentro del formulario: lo fija la OT activa.
    if lote["mov"] == "Retiro":
        lf_fallas.pack(fill="x", pady=(0, 10))
    else:
        lf_fallas.pack_forget()
        for c in comps_data:
            c["estado"].set("OK")
            c["obs"].config(state=tk.NORMAL)
            c["obs"].delete(0, tk.END)
            c["obs"].config(state=tk.DISABLED, bg="#f1f5f9")

    tk.Label(
        marco,
        text="Observación General para el Acta:",
        bg=COLORS["fondo"],
        font=("Segoe UI", 10, "bold"),
        fg=COLORS["gris_oscuro"],
    ).pack(anchor="w", pady=(4, 2))
    entry_obs_gen = tk.Entry(marco, width=50, font=("Segoe UI", 10))
    entry_obs_gen.pack(fill="x", ipady=4, pady=(0, 8))

    frame_off = tk.Frame(marco, bg=COLORS["fondo"])
    frame_off.pack(fill="x", pady=(4, 0))
    var_office = tk.IntVar()

    def toggle_office():
        state = tk.NORMAL if var_office.get() else tk.DISABLED
        entry_office.config(state=state)
        combo_ver.config(state="readonly" if var_office.get() else tk.DISABLED)
        if not var_office.get():
            entry_office.delete(0, tk.END)

    tk.Checkbutton(
        frame_off,
        text="📦 Registrar Office",
        variable=var_office,
        command=toggle_office,
        bg=COLORS["fondo"],
        font=("Segoe UI", 10, "bold"),
        fg="#1e40af",
        activebackground="#eef2f7",
        cursor="hand2",
    ).pack(side="left")
    combo_ver = ttk.Combobox(
        frame_off,
        values=["2013", "2016", "2019", "2021", "2024", "365"],
        width=8,
        state=tk.DISABLED,
        font=("Segoe UI", 10),
    )
    combo_ver.set("2016")
    combo_ver.pack(side="left", padx=6)

    entry_office = tk.Entry(
        marco, font=("Consolas", 10), justify="center", state=tk.DISABLED, fg="#555"
    )
    entry_office.pack(fill="x", ipady=4, pady=6)

    frame_bts = tk.Frame(ventana, bg=COLORS["fondo"])
    frame_bts.pack(pady=10)

    tk.Button(
        frame_bts,
        text=f"💾 GUARDAR EN {carpeta_lote(lote['mov'], lote['numero'])}",
        command=lambda: accion_agregar_lote(
            ventana,
            lote,
            data,
            entry_obs_gen.get().strip(),
            var_office.get() == 1,
            entry_office.get().strip(),
            combo_ver.get(),
            comps_data,
        ),
        bg="#0078d4",
        fg="white",
        font=("Segoe UI", 12, "bold"),
        padx=20,
        pady=10,
        cursor="hand2",
        relief="flat",
    ).pack()


# ============================================================================
# 11. MODO CLONACIÓN (sin menú) — para el post-script de FOG
# ----------------------------------------------------------------------------
# Se invoca como:  REPORTE3.exe --clonacion [--espera SEGUNDOS]
#
# Corre solo en el equipo recién clonado: pregunta al servidor a qué OT
# pertenece, escanea, abre UNA ventana para las observaciones y manda el JSON
# por SCP. No hay menú ni OT que elegir: el operador no decide nada que pueda
# equivocarse, y si se distrae, el equipo se manda igual al vencer la espera.
# ============================================================================
_ESPERA_CLONACION = 180  # segundos antes del envío automático

# Con --sin-ui no se abre NINGUNA ventana: ni la de observaciones ni los avisos
# finales. Es el modo que necesita el post-script de FOG, donde no hay nadie
# para cerrar un messagebox y el proceso se quedaría colgado sin devolver el
# código de salida.
#
# Que no haya ventanas no quiere decir que no haya operador: si el programa
# corre en un cmd, las observaciones se piden por consola con la misma cuenta
# regresiva, y recién después se hace el SCP. Sin consola (post-script de FOG
# de verdad) se manda directo.
_SIN_UI = False


def _log_clon(texto):
    """Traza para el post-script de FOG, que se queda con la salida estándar."""
    print(f"[REPORTE3] {texto}", flush=True)


def _hay_consola():
    """
    ¿Hay alguien del otro lado del teclado? Con el .exe lanzado desde un cmd
    sys.stdin es la consola; lanzado por el post-script de FOG o con la entrada
    redirigida, no hay a quién preguntarle nada.
    """
    try:
        return sys.stdin is not None and sys.stdin.isatty()
    except Exception:
        return False


def _leer_linea_con_espera(prompt, espera):
    """
    Lee una línea de la consola con cuenta regresiva. Devuelve (texto, respondió).

    Apenas el operador toca una tecla se cancela el auto-envío: si está
    escribiendo la observación no se le puede cortar la frase por la mitad.
    Con espera <= 0 se espera indefinidamente, igual que en la ventana.
    """
    print(prompt, end="", flush=True)
    try:
        import msvcrt
    except ImportError:  # fuera de Windows no hay cuenta regresiva posible
        return input().strip(), True

    escrito = []
    limite = time.monotonic() + espera if espera > 0 else None
    ultimo_aviso = None

    while True:
        if msvcrt.kbhit():
            tecla = msvcrt.getwch()
            if tecla in ("\x00", "\xe0"):  # flechas y F1..F12 llegan de a dos
                msvcrt.getwch()
                continue
            if tecla == "\x03":            # Ctrl+C
                raise KeyboardInterrupt
            if tecla in ("\r", "\n"):
                print(flush=True)
                return "".join(escrito).strip(), True
            if tecla == "\b":
                if escrito:
                    escrito.pop()
                    print("\b \b", end="", flush=True)
                continue
            escrito.append(tecla)
            print(tecla, end="", flush=True)
            limite = None                  # ya está escribiendo: no se le corta
            continue

        if limite is not None:
            restante = limite - time.monotonic()
            if restante <= 0:
                print("\n[REPORTE3] Espera agotada: se envía automáticamente.",
                      flush=True)
                return "".join(escrito).strip(), False
            # Hacia arriba: con int() a secas la espera se acortaba un segundo
            # y --espera 1 no llegaba a mostrarse nunca.
            segundos = int(restante) + 1
            if segundos != ultimo_aviso and segundos % 30 == 0:
                ultimo_aviso = segundos
                print(f"\n(envío automático en {segundos}s) {prompt}"
                      f"{''.join(escrito)}", end="", flush=True)
        time.sleep(0.05)


def _observaciones_consola(data, lote, espera):
    """
    La ventana de observaciones, pero en la consola. Mismo contrato que
    _ventana_observaciones: devuelve (observaciones, tiene_fallas, enviar), y
    todo esto pasa ANTES del SCP.
    """
    print("")
    print("=" * 64)
    print(f"  EQUIPO ESCANEADO - {etiqueta_lote(lote)}")
    print("=" * 64)
    for etiqueta, valor in (
        ("Modelo", data.get("model", "?")),
        ("N de Serie", data.get("serial", "?")),
        ("Procesador", data.get("cpu", "?")),
        ("Licencia Windows", data.get("key", "N/A")),
    ):
        print(f"  {etiqueta:<18}{valor}")
    print("-" * 64)
    if espera > 0:
        print(f"  Si no escribís nada, se envía solo en {espera}s.")
    print("  Ctrl+C cancela el registro.")
    print("")

    try:
        obs, respondio = _leer_linea_con_espera(
            "Observaciones (Enter = ninguna): ", espera
        )
        if not respondio:
            return obs, False, True
        # Ya está frente al teclado: para la segunda pregunta alcanza con un
        # margen corto, no hace falta repetir la espera larga.
        resp, _ = _leer_linea_con_espera("¿El equipo tiene fallas? [s/N]: ", 60)
        return obs, resp.strip().lower().startswith("s"), True
    except KeyboardInterrupt:
        print("\n[REPORTE3] Cancelado por el operador.", flush=True)
        return "", False, False


def _ventana_observaciones(data, lote, espera):
    """
    Ventana única del modo clonación: muestra lo escaneado y toma las
    observaciones. Devuelve (observaciones, tiene_fallas, enviar).

    'enviar' sale en False solo si el operador cancela a propósito; si la
    ventana se cierra o se agota la espera se manda igual, porque el costo de
    perder el registro de un equipo ya clonado es mucho mayor que el de
    mandarlo sin observaciones.
    """
    resultado = {"obs": "", "fallas": False, "enviar": True}

    win = tk.Tk()
    win.title(f"Registro de equipo — {etiqueta_lote(lote)}")
    win.configure(bg=COLORS["fondo"])
    win.attributes("-topmost", True)
    win.geometry("560x460")

    titulo_ui(win, "🖥️ Equipo escaneado", size=14, pady=(16, 4))
    tk.Label(
        win,
        text=etiqueta_lote(lote),
        font=fuente(10, True),
        bg="#dcfce7",
        fg="#166534",
        pady=6,
    ).pack(fill="x", padx=24, pady=(0, 10))

    marco = tk.Frame(win, bg="white", relief="flat", bd=1)
    marco.pack(fill="x", padx=24)
    for etiqueta, valor in (
        ("Modelo", data.get("model", "?")),
        ("N° de Serie", data.get("serial", "?")),
        ("Procesador", data.get("cpu", "?")),
        ("Licencia Windows", data.get("key", "N/A")),
    ):
        fila = tk.Frame(marco, bg="white")
        fila.pack(fill="x", padx=10, pady=3)
        tk.Label(
            fila, text=f"{etiqueta}:", font=fuente(9, True), bg="white",
            fg=COLORS["gris_oscuro"], width=16, anchor="w",
        ).pack(side="left")
        tk.Label(
            fila, text=valor, font=fuente(9), bg="white", fg=COLORS["azul"],
            anchor="w", wraplength=330, justify="left",
        ).pack(side="left", fill="x", expand=True)

    tk.Label(
        win, text="Observaciones (opcional)", font=fuente(9, True),
        bg=COLORS["fondo"], fg=COLORS["gris_oscuro"],
    ).pack(anchor="w", padx=24, pady=(14, 2))
    txt_obs = tk.Text(win, height=4, font=fuente(9), relief="solid", bd=1)
    txt_obs.pack(fill="x", padx=24)
    txt_obs.focus_set()

    var_fallas = tk.BooleanVar(value=False)
    tk.Checkbutton(
        win, text="El equipo tiene fallas", variable=var_fallas,
        font=fuente(9), bg=COLORS["fondo"], fg="#991b1b",
        activebackground=COLORS["fondo"], selectcolor="white",
    ).pack(anchor="w", padx=22, pady=(6, 0))

    lbl_cuenta = tk.Label(
        win, text="", font=fuente(8), bg=COLORS["fondo"], fg=COLORS["gris"]
    )
    lbl_cuenta.pack(pady=(6, 0))

    def terminar(enviar=True):
        resultado["obs"] = txt_obs.get("1.0", "end").strip()
        resultado["fallas"] = var_fallas.get()
        resultado["enviar"] = enviar
        try:
            win.destroy()
        except Exception:
            pass

    frame_btn = tk.Frame(win, bg=COLORS["fondo"])
    frame_btn.pack(pady=10)
    boton_accion(frame_btn, "✅ ENVIAR AL SERVIDOR", lambda: terminar(True),
                 COLORS["verde"]).pack(side="left", padx=6)
    boton_accion(frame_btn, "Cancelar", lambda: terminar(False),
                 COLORS["gris"]).pack(side="left", padx=6)

    # Cerrar con la X equivale a enviar: es lo que hace el operador apurado.
    win.protocol("WM_DELETE_WINDOW", lambda: terminar(True))

    restante = {"seg": max(0, int(espera))}

    def tic():
        if restante["seg"] <= 0:
            _log_clon("Espera agotada: se envía automáticamente.")
            terminar(True)
            return
        minutos, seg = divmod(restante["seg"], 60)
        lbl_cuenta.config(text=f"Envío automático en {minutos}:{seg:02d}")
        restante["seg"] -= 1
        win.after(1000, tic)

    if espera > 0:
        tic()
    win.mainloop()
    return resultado["obs"], resultado["fallas"], resultado["enviar"]


def modo_clonacion(espera=_ESPERA_CLONACION, sin_ui=False):
    """
    Flujo completo del equipo clonado. Devuelve el código de salida del proceso
    (0 = enviado), pensado para que el post-script de FOG sepa si funcionó.

    Con sin_ui=True no se abre ninguna ventana: si corre en un cmd las
    observaciones se piden por consola antes del SCP, y si no hay consola el
    equipo se manda sin observaciones. Los errores quedan solo en el log.
    """
    global _SIN_UI
    _SIN_UI = sin_ui
    _log_clon(f"Modo clonación en {socket.gethostname()}")

    try:
        cliente = abrir_conexion_scp()
    except Exception as e:
        _log_clon(f"ERROR de conexión: {e}")
        _avisar_error_clonacion("No se pudo conectar al servidor", str(e))
        return 2

    try:
        lote = leer_lote_activo_remoto(cliente)
        if not lote:
            msg = (
                "No hay ninguna OT activa publicada en el servidor.\n\n"
                "Activá una OT desde el panel central antes de clonar."
            )
            _log_clon("ERROR: sin OT activa publicada")
            _avisar_error_clonacion("Sin OT activa", msg)
            return 3
        _log_clon(f"OT activa: {etiqueta_lote(lote)}")

        _log_clon("Escaneando hardware…")
        data = get_inventory_fast()
        if not data:
            _log_clon("ERROR: el escaneo falló")
            _avisar_error_clonacion(
                "No se pudo escanear",
                "El escaneo de hardware falló. Revisá el registro y reintentá.",
            )
            return 4
        _log_clon(f"Equipo: {data.get('model')} / {data.get('serial')}")

        if sin_ui and _hay_consola():
            obs, tiene_fallas, enviar = _observaciones_consola(data, lote, espera)
        elif sin_ui:
            _log_clon("Sin UI y sin consola: se envía sin observaciones.")
            obs, tiene_fallas, enviar = "", False, True
        else:
            obs, tiene_fallas, enviar = _ventana_observaciones(data, lote, espera)
        if not enviar:
            _log_clon("Cancelado por el operador.")
            return 1

        # Se guarda con la misma función que el flujo normal: el JSON queda
        # idéntico al de un escaneo manual, así el panel y la reconstrucción
        # no tienen que distinguir de dónde vino.
        comps = [{"nombre": "Estado general", "estado": "MALO" if tiene_fallas else "OK",
                  "obs": obs}] if tiene_fallas else []
        guardar_equipo_general(lote, data, obs, False, "", "", comps)
        ruta_json = ruta_json_equipo(lote, data["serial"])
        _log_clon(f"JSON local: {ruta_json}")

        destino = enviar_equipo_por_scp(cliente, lote, ruta_json)
        _log_clon(f"ENVIADO -> {destino}")
        _avisar_ok_clonacion(data, lote)
        return 0
    except Exception as e:
        _log_clon(f"ERROR inesperado: {e}")
        _avisar_error_clonacion("Falló el envío", str(e))
        return 5
    finally:
        try:
            cliente.close()
        except Exception:
            pass


def diagnostico_scp():
    """
    Chequeo completo del envío, corriendo DESDE ESTE equipo. Código 0 = todo OK.

    Va dentro del ejecutable a propósito: el equipo que importa verificar es el
    clonado, y ahí no hay Python ni código fuente, solo el .exe. Verificar desde
    el servidor de FOG no sirve —FOG nunca habla con el NAS, solo copia archivos
    al disco de Windows—, y verificar desde la PC de escritorio tampoco prueba
    que el equipo clonado tenga la clave donde corresponde.
    """
    ok, mal = "  [OK]   ", "  [FALLA] "
    print("\n=== Configuración ===")
    print(f"  Ejecutable : {ruta_base_proyecto}")
    print(f"  Servidor   : {SCP_USUARIO}@{SCP_HOST}:{SCP_PUERTO}")
    print(f"  Destino    : {SCP_DESTINO}")

    if not SCP_AVAILABLE:
        print(f"\n{mal}Falta paramiko/scp en el ejecutable.")
        print("       Recompilá con el .spec (trae los hiddenimports).")
        return 1

    print("\n=== 1. Clave privada ===")
    ruta = resolver_clave_ssh()
    for candidata in _candidatas_clave_ssh():
        marca = "  <-- se usa esta" if candidata == ruta else ""
        print(f"    {candidata}{marca}")
    if not ruta:
        print(f"{mal}No hay clave en ninguna de esas rutas.")
        print("       En un equipo clonado la clave va JUNTO al .exe.")
        return 1
    print(f"{ok}Clave encontrada")

    print("\n=== 2. Conexión ===")
    try:
        cliente = abrir_conexion_scp()
    except Exception as e:
        print(f"{mal}{e}")
        return 1
    print(f"{ok}Conectado y autenticado")

    try:
        print("\n=== 3. Permisos en el servidor ===")
        codigo, salida, _ = _correr_remoto(cliente, "id")
        if codigo != 0 or not salida:
            print(f"{mal}El usuario no puede ejecutar comandos (¿shell nologin?)")
            return 1
        print(f"{ok}{salida}")
        try:
            verificar_destino_remoto(cliente)
        except Exception as e:
            print(f"{mal}{e}")
            return 1
        print(f"{ok}Carpeta de destino escribible")

        print("\n=== 4. OT activa publicada ===")
        lote = leer_lote_activo_remoto(cliente)
        if lote:
            print(f"{ok}{etiqueta_lote(lote)}")
        else:
            # No es un fallo del equipo: es que en la central todavía no
            # publicaron la OT. Se distingue para no mandar a revisar la red.
            print("  [AVISO] No hay OT publicada.")
            print("       En la PC central: PANEL DE CLONACIÓN → Publicar OT activa.")

        print("\n=== 5. Subida real de prueba ===")
        local = os.path.join(tempfile.mkdtemp(), "prueba_scp.txt")
        with open(local, "w", encoding="utf-8") as f:
            f.write(f"prueba desde {socket.gethostname()}\n")
        carpeta = _ruta_remota(SCP_DESTINO, "_prueba_scp")
        remoto = _ruta_remota(carpeta, "prueba_scp.txt")
        try:
            _mkdir_remoto(cliente, carpeta)
            with SCPClient(cliente.get_transport(), socket_timeout=SCP_TIMEOUT) as scp:
                scp.put(local, remote_path=remoto)
            codigo, salida, _ = _correr_remoto(cliente, f"cat '{remoto}'")
            if codigo != 0 or socket.gethostname() not in salida:
                print(f"{mal}El archivo no llegó o no se pudo leer de vuelta")
                return 1
            print(f"{ok}Archivo subido y leído de vuelta")
        finally:
            _correr_remoto(cliente, f"rm -rf '{carpeta}'")

        print("\n" + "=" * 56)
        print("  TODO LISTO — este equipo puede enviar.")
        print("=" * 56 + "\n")
        return 0
    finally:
        try:
            cliente.close()
        except Exception:
            pass


def _avisar_error_clonacion(titulo, detalle):
    """Error visible en pantalla: en el equipo clonado nadie mira la consola."""
    if _SIN_UI:
        return
    try:
        raiz = tk.Tk()
        raiz.withdraw()
        raiz.attributes("-topmost", True)
        messagebox.showerror(f"❌ {titulo}", detalle)
        raiz.destroy()
    except Exception:
        pass


def _avisar_ok_clonacion(data, lote):
    if _SIN_UI:
        return
    try:
        raiz = tk.Tk()
        raiz.withdraw()
        raiz.attributes("-topmost", True)
        messagebox.showinfo(
            "✅ Equipo registrado",
            f"{data.get('model')}\nS/N: {data.get('serial')}\n\n"
            f"Enviado a {etiqueta_lote(lote)}.",
        )
        raiz.destroy()
    except Exception:
        pass


# ============================================================================
# 12. ENTRYPOINT
# ============================================================================
def _ocultar_consola():
    """
    El .exe se compila con console=True para que el post-script de FOG pueda
    quedarse con la traza de --clonacion. En el uso diario esa consola negra
    detrás del menú no aporta nada, así que se esconde apenas arranca la GUI.
    """
    try:
        import ctypes

        ventana_consola = ctypes.windll.kernel32.GetConsoleWindow()
        if ventana_consola:
            ctypes.windll.user32.ShowWindow(ventana_consola, 0)  # SW_HIDE
    except Exception:
        pass


def iniciar_interfaz_principal():
    _ocultar_consola()
    ventana = tk.Tk()
    ventana.title("Asistente de Entrega — Arrienda.cl")
    ventana.resizable(False, False)
    ventana.attributes("-topmost", True)
    ventana.configure(bg=COLORS["fondo"])

    style = ttk.Style()
    style.theme_use("clam")
    style.configure(
        "TCombobox", fieldbackground="#fff", background="#fff", font=("Segoe UI", 9)
    )
    style.configure("TEntry", fieldbackground="#fff", font=("Segoe UI", 9))
    style.configure("TLabelframe", background="#eef2f7")
    style.configure(
        "TLabelframe.Label",
        background="#eef2f7",
        font=("Segoe UI", 10, "bold"),
        foreground="#334155",
    )

    mostrar_menu_principal(ventana)
    ventana.mainloop()


def main(argv=None):
    """
    Punto de entrada. Sin argumentos abre el menú de siempre; con --clonacion
    corre el flujo automático del equipo recién clonado.
    """
    argv = list(sys.argv[1:] if argv is None else argv)

    # El .bat redirige la salida a un log con la codificación de la consola
    # (cp1252/cp850). Sin esto, un acento que no entre revienta el print, el
    # except de modo_clonacion lo convierte en "error inesperado" y se pierde
    # el registro de un equipo ya clonado por un problema de tipografía.
    for flujo in (sys.stdout, sys.stderr):
        try:
            flujo.reconfigure(errors="replace")
        except Exception:
            pass

    if "--ayuda" in argv or "-h" in argv or "--help" in argv:
        print(
            "REPORTE3 — Sistema de Control de Inventario\n\n"
            "  REPORTE3.exe                       menú normal\n" "  REPORTE3.exe --verificar           chequea el envío desde ESTE equipo\n"
            "  REPORTE3.exe --clonacion           registro automático del equipo\n"
            "  REPORTE3.exe --clonacion --espera 120\n"
            "                                     segundos antes del envío\n"
            "                                     automático (0 = sin límite)\n"
            "  REPORTE3.exe --clonacion --sin-ui  sin ninguna ventana: en un cmd\n"
            "                                     pide las observaciones por\n"
            "                                     consola antes de enviar; sin\n"
            "                                     consola manda directo. Respeta\n"
            "                                     --espera y Ctrl+C cancela.\n\n"
            "Códigos de salida en modo clonación:\n"
            "  0 enviado   1 cancelado   2 sin conexión\n"
            "  3 sin OT activa   4 falló el escaneo   5 error inesperado\n"
        )
        return 0

    if "--verificar" in argv:
        return diagnostico_scp()

    if "--clonacion" in argv:
        espera = _ESPERA_CLONACION
        if "--espera" in argv:
            try:
                espera = int(argv[argv.index("--espera") + 1])
            except (IndexError, ValueError):
                _log_clon(
                    f"--espera inválido, se usa el valor por defecto ({espera}s)"
                )
        return modo_clonacion(espera, sin_ui="--sin-ui" in argv)

    iniciar_interfaz_principal()
    return 0


if __name__ == "__main__":
    sys.exit(main())
