"""
Sistema de Control de Inventario — Arrienda.cl
==============================================

Aplicación de escritorio (Tkinter) para escanear el hardware de un equipo,
registrar entregas/retiros y generar reportes HTML, JSON y Word.

ÍNDICE DEL ARCHIVO (buscá el banner "==== NOMBRE ====" para saltar):

  1. CONFIGURACIÓN            — rutas y constantes globales.
  2. TEMA Y HELPERS DE UI      — paleta (COLORS), fuentes y widgets reutilizables
                                 de Tkinter (titulo_ui, ventana_modal, boton_menu,
                                 boton_accion, fila_selector, treeview_auditoria).
  3. SCANNER (PowerShell)      — PS_SCRIPT: script embebido que lee el hardware.
  4. PLANTILLA HTML DEL REPORTE — HTML_START / HTML_END (CSS + JS del reporte).
  5. EXPORT WORD               — helpers docx + generar_informe_word.
  6. SCANNER (ejecución)       — get_inventory_fast: corre PS_SCRIPT y parsea.
  7. PERSISTENCIA              — config de guía, construir_entry_html,
                                 guardar_equipo_general, empacar_lote, enviar a red.
  8. RECONSTRUCCIÓN / MERGE    — rearmar HTML desde los JSON guardados.
  9. MÓDULOS (Toplevel)        — auditorías, comparación Excel, etiquetas, nombre.
 10. UI PRINCIPAL              — menú, pantalla de escaneo y formulario.
 11. ENTRYPOINT               — iniciar_interfaz_principal / __main__.

NOTA: la apariencia del REPORTE HTML vive en HTML_START/HTML_END y en
construir_entry_html; no se debe alterar sin querer cambiar el reporte.
"""

import html  # Para escapar strings en el HTML
import json
import os
import re
import shutil
import socket
import subprocess
import sys
import tempfile
import threading
import webbrowser
import time
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, messagebox, simpledialog, ttk

# --- SOLUCIÓN PYINSTALLER: Asegurar que el directorio de trabajo sea siempre el correcto ---
if getattr(sys, "frozen", False):
    # Si el programa está compilado (.exe) con PyInstaller
    ruta_base_proyecto = os.path.dirname(sys.executable)
else:
    # Si se está ejecutando como script de Python normal (.py)
    ruta_base_proyecto = os.path.dirname(os.path.abspath(__file__))

# Cambiamos el directorio de trabajo a la carpeta real del .exe
os.chdir(ruta_base_proyecto)
# -------------------------------------------------------------------------------------------

# Importar librería para pandas
try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Importar librería para crear Word (docx)
try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# 1. CONFIGURACIÓN
# ============================================================================
CARPETA_DESTINO = r"\\192.168.1.68\transformadores"


# ============================================================================
# 2. TEMA Y HELPERS DE UI (solo para la ventana Tkinter de escritorio)
# ----------------------------------------------------------------------------
# Paleta y helpers centralizados para la interfaz Tkinter. NO afectan al HTML
# del reporte (ese usa sus propios colores dentro de HTML_START/HTML_END y de
# las plantillas f-string). Cambiar el look de la app se hace desde aquí.
# ============================================================================

# Color de marca (azul corporativo) como fuente única para Tkinter y Word.
MARCA_RGB = (0x1A, 0x3C, 0x6E)
MARCA_HEX = "1A3C6E"

COLORS = {
    "fondo": "#eef2f7",          # fondo general de ventanas
    "azul": "#1a3c6e",           # azul de marca (títulos, botones de ruta)
    "azul_btn": "#0078d4",       # botón principal "escanear"
    "verde": "#16a34a",          # acción / OK
    "verde_esmeralda": "#059669",
    "naranja": "#d97706",        # advertencia / enviar red
    "morado": "#7c3aed",         # reconstruir
    "celeste": "#0ea5e9",        # comparar excel
    "pizarra": "#475569",        # etiquetas
    "gris": "#64748b",           # texto secundario
    "gris_oscuro": "#334155",
    "blanco": "white",
}


def fuente(size=9, bold=False):
    """Tupla de fuente estándar de la app (Segoe UI)."""
    return ("Segoe UI", size, "bold") if bold else ("Segoe UI", size)


def titulo_ui(parent, texto, size=14, pady=12, fg=None):
    """Label de título estándar (azul de marca sobre fondo claro)."""
    lbl = tk.Label(
        parent,
        text=texto,
        font=fuente(size, True),
        bg=COLORS["fondo"],
        fg=fg or COLORS["azul"],
    )
    lbl.pack(pady=pady)
    return lbl


def ventana_modal(padre, titulo_ventana, geometria):
    """Crea un Toplevel modal con el fondo y grab_set estándar."""
    win = tk.Toplevel(padre)
    win.title(titulo_ventana)
    win.geometry(geometria)
    win.configure(bg=COLORS["fondo"])
    win.grab_set()
    return win


def boton_menu(parent, texto, comando, color, size=11, bold="bold"):
    """Botón ancho del menú/pantallas (relleno horizontal)."""
    tk.Button(
        parent,
        text=texto,
        command=comando,
        bg=color,
        fg="white",
        font=("Segoe UI", size, bold),
        pady=10,
        cursor="hand2",
        relief="flat",
    ).pack(fill="x", pady=6)


def boton_accion(parent, texto, comando, color=None, padx=18, pady=7):
    """Botón de acción destacado (por defecto verde). El caller lo empaqueta."""
    return tk.Button(
        parent,
        text=texto,
        command=comando,
        bg=color or COLORS["verde"],
        fg="white",
        font=fuente(10, True),
        padx=padx,
        pady=pady,
        cursor="hand2",
        relief="flat",
    )


# Estilo compartido de los botones "elegir ruta/carpeta" de los módulos.
def _estilo_btn_ruta():
    return dict(
        bg=COLORS["azul"], fg="white", font=fuente(9), cursor="hand2", relief="flat"
    )


def fila_selector(frame, fila, texto_btn, comando, var):
    """Fila 'botón elegir ruta + entry de solo lectura' usada por las auditorías."""
    tk.Button(frame, text=texto_btn, command=comando, **_estilo_btn_ruta()).grid(
        row=fila, column=0, padx=5, pady=3, sticky="ew"
    )
    tk.Entry(
        frame, textvariable=var, width=55, state="readonly", font=fuente(9)
    ).grid(row=fila, column=1, padx=5)


def treeview_auditoria(win, cols_spec, height, con_alterado=False):
    """
    Crea el Treeview estándar de auditoría (estilo + columnas + tags de color +
    scrollbar dentro de su frame) y lo devuelve listo para insertar filas.
    cols_spec: lista de (col_id, ancho, encabezado).
    """
    style = ttk.Style()
    style.configure("Audit.Treeview", rowheight=24, font=fuente(9))
    style.configure("Audit.Treeview.Heading", font=fuente(9, True))

    cols = tuple(c[0] for c in cols_spec)
    tree = ttk.Treeview(
        win, columns=cols, show="headings", height=height, style="Audit.Treeview"
    )
    for col, w, txt in cols_spec:
        tree.heading(col, text=txt)
        tree.column(
            col, width=w, anchor="center" if col in ("estado", "serial") else "w"
        )

    tree.tag_configure("ok", foreground="#16a34a")
    tree.tag_configure("falta", foreground="#dc2626", background="#fee2e2")
    tree.tag_configure("sobra", foreground="#d97706", background="#fef3c7")
    if con_alterado:
        tree.tag_configure(
            "alterado",
            foreground="#b91c1c",
            background="#fecaca",
            font=fuente(9, True),
        )

    sb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=sb.set)
    frame_tree = tk.Frame(win, bg=COLORS["fondo"])
    frame_tree.pack(fill="both", expand=True, padx=20, pady=8)
    tree.pack(side="left", fill="both", expand=True, in_=frame_tree)
    sb.pack(side="right", fill="y", in_=frame_tree)
    return tree

# ============================================================================
# 3. SCANNER (PowerShell)
# ----------------------------------------------------------------------------
# Script embebido que recolecta el hardware vía CIM/WMI + powercfg y lo emite
# como JSON. Se ejecuta desde get_inventory_fast() (sección 6).
# ============================================================================
PS_SCRIPT = """
$ErrorActionPreference = 'SilentlyContinue'
$bios   = Get-CimInstance Win32_Bios
$sys    = Get-CimInstance Win32_ComputerSystem
$cpu    = Get-CimInstance Win32_Processor
$svc    = Get-CimInstance SoftwareLicensingService
$winKey = if ($svc.OA3xOriginalProductKey) { $svc.OA3xOriginalProductKey } else { "No encontrada en BIOS" }

# --- NUEVO: LÓGICA DE BATERÍA CON POWERCFG XML ---
$battInfo = "No detectada"
$xmlPath = "$env:TEMP\\batt_report_temp.xml"
if (Test-Path $xmlPath) { Remove-Item $xmlPath -Force -ErrorAction SilentlyContinue }

# Ejecutamos powercfg silenciosamente para que genere el XML
& powercfg /batteryreport /xml /output $xmlPath | Out-Null

if (Test-Path $xmlPath) {
    try {
        [xml]$battXml = Get-Content $xmlPath
        $baterias = $battXml.BatteryReport.Batteries.Battery
        if ($baterias) {
            # Tomar la primera batería por si el equipo tiene dos
            $bat = if ($baterias.Count -gt 1) { $baterias[0] } else { $baterias }

            $design = [int]$bat.DesignCapacity
            $full = [int]$bat.FullChargeCapacity

            if ($design -gt 0) {
                $salud = [math]::Round(($full / $design) * 100)
                if ($salud -gt 100) { $salud = 100 } # Tope visual al 100%
                $battInfo = "$salud% | Diseño: $design mWh | Actual: $full mWh"
            }
        }
    } catch {
        $battInfo = "Error al leer datos XML"
    }
    # Autodestruir el reporte temporal para no dejar basura
    Remove-Item $xmlPath -Force -ErrorAction SilentlyContinue
}
# ------------------------------------------------

$ramInfo = @(Get-CimInstance Win32_PhysicalMemory | ForEach-Object {
    $speed = if ($_.Speed) { "$($_.Speed)" } else { "Desconocida" }
    $Partnumber = if ($_.PartNumber) { $_.PartNumber.Trim() } else { "Desconocida" }
    $Manufacturer = if ($_.Manufacturer) { $_.Manufacturer.Trim() } else { "Desconocida" }
    $typeNum = if ($_.SMBIOSMemoryType) { $_.SMBIOSMemoryType } else { $_.MemoryType }
    $typeStr = switch ($typeNum) { 24 {"PC3"} 26 {"PC4"} 30 {"LPDDR4"} 34 {"PC5"} 35 {"LPDDR5"} default {"PC4"} }
    "{0}GB | {1} | {2} | {3} | {4}" -f ([math]::Round($_.Capacity/1GB)), $speed, $typeStr, $Manufacturer, $Partnumber
})

$disksData = @(Get-PhysicalDisk | Where-Object { $_.BusType -notin @('USB','File Backed Virtual') } | ForEach-Object {
    $size  = [math]::Round($_.Size / 1000000000)
    $media = if ([string]::IsNullOrWhiteSpace($_.MediaType) -or $_.MediaType -eq 'Unspecified') { 'SSD' } else { [string]$_.MediaType }
    $bus   = if ([string]$_.BusType -match 'NVMe') { 'M.2' } else { [string]$_.BusType }
    @{ desc = "{0} {1}GB {2} {3}" -f $_.FriendlyName.Trim(), $size, $media, $bus
       serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
})
if ($disksData.Count -eq 0) {
    $disksData = @(Get-CimInstance Win32_DiskDrive | Where-Object { $_.InterfaceType -ne 'USB' } | ForEach-Object {
        @{ desc = "{0} {1}GB SSD" -f $_.Model.Trim(), ([math]::Round($_.Size/1000000000))
           serial = if ($_.SerialNumber) { $_.SerialNumber.Trim() } else { "SIN-SERIE" } }
    })
}

$netInfo = @(Get-CimInstance Win32_NetworkAdapterConfiguration | Where-Object { $_.IPEnabled } | ForEach-Object {
    "{0} | IP: {1} | MAC: {2}" -f $_.Description, $_.IPAddress[0], $_.MACAddress
})

@{ serial=($bios.SerialNumber); model=($sys.Model); win_key=($winKey)
   cpu=($cpu.Name); ram_rows=$ramInfo; disks_data=$disksData; net_rows=$netInfo; battery=($battInfo)
} | ConvertTo-Json -Compress -Depth 3
"""

# ─────────────────────────────────────────────────────────
#  HTML / CSS / JS  (plantilla del reporte)
# ─────────────────────────────────────────────────────────
# ============================================================================
# 4. PLANTILLA HTML DEL REPORTE  (⚠️ apariencia del reporte — NO modificar)
# ----------------------------------------------------------------------------
# HTML_START trae toda la página (CSS + JS de estadísticas, etiquetas, export
# CSV, etc.) y termina en <div id="main-list">. HTML_END cierra las etiquetas.
# Las entradas de cada equipo se insertan entre ambos (ver construir_entry_html).
# ============================================================================
HTML_START = r"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Inventario Maestro</title>
<style>
  *{box-sizing:border-box;margin:0;padding:0}
  body{font-family:'Segoe UI',Tahoma,sans-serif;background:#eef2f7;color:#1e293b;padding:30px}
  .container{max-width:1050px;margin:auto}
  .header{text-align:center;margin-bottom:28px}
  h1{font-size:26px;color:#1a3c6e;margin-bottom:6px}
  .subtitle{color:#64748b;font-size:13px;margin-bottom:14px}
  .counter-box{background:#dbeafe;color:#1d4ed8;padding:5px 18px;border-radius:20px;
    font-weight:700;font-size:13px;display:inline-block;border:1px solid #bfdbfe;margin-bottom:12px}

  .stats-container { display:flex; justify-content:center; gap:10px; margin-bottom:22px; flex-wrap:wrap; }
  .stat-pill { padding:5px 16px; border-radius:20px; font-weight:700; font-size:12px; border:1px solid; }
  .stat-ram { background:#fce7f3; color:#be185d; border-color:#fbcfe8; }
  .stat-cpu { background:#fef3c7; color:#b45309; border-color:#fde68a; }
  .stat-disk { background:#dcfce7; color:#15803d; border-color:#bbf7d0; }

  .button-group{display:flex;justify-content:center;gap:12px;margin-bottom:32px;flex-wrap:wrap}
  .btn{color:#fff;border:none;padding:10px 22px;font-size:13px;font-weight:700;border-radius:8px;
    cursor:pointer;box-shadow:0 2px 6px rgba(0,0,0,.18);transition:transform .1s,box-shadow .1s}
  .btn:hover{transform:translateY(-1px);box-shadow:0 4px 12px rgba(0,0,0,.22)}
  .btn:active{transform:translateY(0)}
  .btn-copy{background:linear-gradient(135deg,#0078d4,#005a9e)}
  .btn-export{background:linear-gradient(135deg,#16a34a,#15803d)}
  .btn-delete{background:#ef4444;padding:5px 12px;font-size:11px;border-radius:6px}
  .btn-delete:hover{background:#dc2626}

  details{background:#fff;margin-bottom:12px;border-radius:12px;
    box-shadow:0 2px 8px rgba(0,0,0,.08);overflow:hidden;border:1px solid #e2e8f0}
  summary{padding:14px 18px;font-weight:700;cursor:pointer;
    background:linear-gradient(135deg,#1a3c6e,#0f2d5a);color:#fff;
    list-style:none;display:flex;justify-content:space-between;align-items:center;gap:10px}
  summary:hover{background:linear-gradient(135deg,#1e4a84,#142f60)}
  summary::-webkit-details-marker{display:none}

  .tag-entrega{background:#22c55e;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-retiro{background:#f97316;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700}
  .tag-falla{background:#ef4444;color:#fff;padding:3px 9px;border-radius:20px;font-size:11px;font-weight:700;margin-left:8px}
  .logistics-row{background:#f8fafc}

  .content{padding:20px 22px;border-top:1px solid #f1f5f9}
  .content-header{display:flex;justify-content:flex-end;margin-bottom:10px; gap: 8px;}
  table{border-collapse:collapse;width:100%}
  th,td{border:1px solid #e2e8f0;padding:9px 12px;text-align:left;font-size:13px}
  th{background:#f8fafc;width:28%;color:#475569;font-weight:600}
  tr:hover td,tr:hover th{background:#f0f9ff}
  .serial-tag{color:#dc2626;font-weight:700;font-family:Consolas,monospace}
  .net-tag{background:#dcfce7;color:#166534;padding:2px 7px;border-radius:4px;font-weight:600;font-size:12px}
  .win-key{color:#1d4ed8;font-family:Consolas,monospace;font-weight:700;font-size:12px}
  .obs-row{background:#fef9c3!important}
  .obs-row td,.obs-row th{color:#854d0e;font-weight:600;font-style:italic}
  .falla-row{background:#fee2e2!important}
  .falla-row td,.falla-row th{color:#991b1b;font-weight:600}
  .office-row{background:#eff6ff}
  .office-row td,.office-row th{color:#1e40af}
</style>
<script>
function actualizarEstadisticas() {
  let eqs = document.querySelectorAll('details');
  document.getElementById('total-count').innerText = eqs.length;

  let cpus = {}, rams = {}, discos = {};

  eqs.forEach(eq => {
    let d = extraerDatos(eq);

    let cpuU = d.cpu.toUpperCase();
    let cpuKey = "Otro";
    if(cpuU.includes("I3")) cpuKey = "i3";
    else if(cpuU.includes("I5")) cpuKey = "i5";
    else if(cpuU.includes("I7")) cpuKey = "i7";
    else if(cpuU.includes("I9")) cpuKey = "i9";
    else if(cpuU.includes("RYZEN 3")) cpuKey = "Ryzen 3";
    else if(cpuU.includes("RYZEN 5")) cpuKey = "Ryzen 5";
    else if(cpuU.includes("RYZEN 7")) cpuKey = "Ryzen 7";
    cpus[cpuKey] = (cpus[cpuKey] || 0) + 1;

    let ramMatch = d.specs.match(/^(\d+GB)/);
    let ramKey = ramMatch ? ramMatch[1] : "Otra";
    rams[ramKey] = (rams[ramKey] || 0) + 1;

    d.discos.forEach(dk => {
      let desc = dk.desc.toUpperCase();
      let tipo = "Desconocido";
      if(desc.includes("M.2") || desc.includes("NVME")) tipo = "M.2 NVMe";
      else if(desc.includes("SSD")) tipo = "SSD 2.5\"";
      else if(desc.includes("HDD")) tipo = "HDD";
      discos[tipo] = (discos[tipo] || 0) + 1;
    });
  });

  let mkHTML = (obj, clase, prefijo) => {
    let arr = Object.entries(obj).map(([k,v]) => `${k} (${v})`);
    return arr.length ? `<span class="stat-pill ${clase}">${prefijo}: ${arr.join(' | ')}</span>` : '';
  };

  let statsDiv = document.getElementById('stats-container');
  if(statsDiv) {
    statsDiv.innerHTML = mkHTML(rams, 'stat-ram', 'RAM') +
                         mkHTML(cpus, 'stat-cpu', 'CPU') +
                         mkHTML(discos, 'stat-disk', 'Discos');
  }
}

document.addEventListener("DOMContentLoaded", actualizarEstadisticas);

function eliminarEquipo(id){
  if(confirm("¿Quitar este equipo del reporte actual?\n(No afecta los archivos ya guardados)")){
    let n=document.getElementById(id);
    if(n){
      n.remove();
      actualizarEstadisticas();
    }
  }
}

function extraerDatos(eq){
  let model  = eq.querySelector('.model-name').innerText.trim();
  let serial = eq.querySelector('.serial-tag').innerText.trim();
  let tds=eq.querySelectorAll('th,td');
  let obs="",cpu="",office="",guia="",mov="",ramSpeed="",ramTech="PC4",fallas="",Manufacturer="",Partnumber="Desconocida";
  let totalRam=0,ramArr=[],discos=[],tempDesc="";

  for(let i=0;i<tds.length;i++){
    if(tds[i].tagName !== 'TH') continue;

    let t=tds[i].innerText.trim();
    if(t==='N° de Guía')          guia   =tds[i+1].innerText.trim();
    if(t==='Tipo Movimiento')     mov    =tds[i+1].innerText.trim();
    if(t.includes('Obs. General'))obs    =tds[i+1].innerText.trim();
    if(t==='Fallas Detectadas')   fallas =tds[i+1].innerText.trim();
    if(t==='Procesador')          cpu    =tds[i+1].innerText.trim();
    if(t.includes('Office'))      office =tds[i+1].innerText.trim();
    if(t==='Discos Internos')     tempDesc=tds[i+1].innerText.trim();
    if(t==='Serie Disco Duro')    discos.push({desc:tempDesc,serial:tds[i+1].innerText.trim()});
    if(t==='Módulo RAM'){
      let p=tds[i+1].innerText.split('|');
      let n=parseInt(p[0].replace(/\D/g,''));
      if(!isNaN(n)){totalRam+=n;ramArr.push(n)}
      if(p[1]&&ramSpeed==="")ramSpeed=p[1].trim();
      if(p[2]&&ramTech==="PC4")ramTech=p[2].trim();
      if(p[3]&&Manufacturer==="")Manufacturer=p[3].trim()
      if(p[4]&&Partnumber==="Desconocida")Partnumber=p[4].trim()
    }
  }

  let ramFinal="RAM Desconocida";
  if(totalRam>0){
    let det=ramArr.length>1?(ramArr.every(v=>v===ramArr[0])?`(${ramArr[0]}X${ramArr.length})`:`(${ramArr.join('+')})`):"";
    let ts=ramSpeed!=="Desconocida"?`${ramTech}-${ramSpeed}Mhz`:ramTech;
    ramFinal=`${totalRam}GB${det} (${ts})`;
  }
  let diskSimple=discos.map(d=>{let m=d.desc.match(/(\d+)\s*GB\s*(.*)/i);return m?m[1]+" GB "+m[2]:d.desc});
  let specs=ramFinal+(diskSimple.length>0?", "+diskSimple.join(" + "):"");
  if(obs&&obs!=="Sin observaciones")specs+=", OBS: "+obs;
  if(fallas&&fallas!=="Ninguna")specs+=", FALLAS: "+fallas;

  return {model,serial,cpu,obs,office,guia,mov,discos,specs,equipoFull:model+" "+cpu, fallas, Manufacturer, Partnumber};
}

function copiarFilas(){
  let txt="";
  document.querySelectorAll('details').forEach(eq=>{
    let d=extraerDatos(eq);
    txt+=`${d.equipoFull}\t${d.specs}\t${d.serial}\n`;
    d.discos.forEach(dk=>{
      let lbl=dk.desc.toUpperCase().includes("HDD")?"HDD":"SSD";
      txt+=`${lbl}\t${dk.desc}\t${dk.serial}\n`;
    });
    if(d.office&&d.office.includes('Key:')){
      let ver=(d.office.match(/Office\s+([A-Za-z0-9]+)/)||[])[1]||"2016";
      let key=(d.office.match(/Key:\s*([^)]+)/)||[])[1]||"";
      txt+=`OFFICE ${ver}\tHOME AND BUSINESS\t${key.trim()}\n`;
    }
  });

  if (navigator.clipboard && window.isSecureContext) {
      navigator.clipboard.writeText(txt).then(() => {
          alert("✅ ¡Filas copiadas!\n\nPega directamente en tu Excel / Valida.");
      }).catch(err => {
          alert("Error al copiar al portapapeles.");
      });
  } else {
      let ta=document.createElement("textarea");ta.value=txt;document.body.appendChild(ta);ta.select();
      try{document.execCommand('copy');alert("✅ ¡Filas copiadas!\n\nPega directamente en tu Excel / Valida.");}
      catch(e){alert("Error al copiar.")}
      document.body.removeChild(ta);
  }
}

function exportarValidaCSV(){
  let csv="\uFEFFTIPO;GUIA;NOMBRE DEL EQUIPO;DESCRIPCION;NUMERO DE SERIE;OBSERVACIONES\n";
  document.querySelectorAll('details').forEach(eq=>{
    let d=extraerDatos(eq);
    csv+=`"${d.mov}";"${d.guia}";"${d.equipoFull}";"${d.specs}";"${d.serial}";"${d.obs}"\n`;
    d.discos.forEach(dk=>{
      let lbl=dk.desc.toUpperCase().includes("HDD")?"HDD":"SSD";
      csv+=`"${d.mov}";"${d.guia}";"${lbl}";"${dk.desc}";"${dk.serial}";""\n`;
    });
    if(d.office&&d.office.includes('Key:')){
      let ver=(d.office.match(/Office\s+([A-Za-z0-9]+)/)||[])[1]||"2016";
      let key=(d.office.match(/Key:\s*([^)]+)/)||[])[1]||"";
      csv+=`"${d.mov}";"${d.guia}";"OFFICE ${ver}";"HOME AND BUSINESS";"${key.trim()}";""\n`;
    }
  });
  let blob=new Blob([csv],{type:'text/csv;charset=utf-8;'}),a=document.createElement("a");
  a.href=URL.createObjectURL(blob);
  let titulo=document.querySelector('title').innerText;
  let tipo=titulo.includes("Entregas")?"Entregas":(titulo.includes("Retiros")?"Retiros":"General");
  let cliente=titulo.replace(tipo+" - ","").replace("Inventario Maestro","General").replace(/ /g,"_");
  a.download=`Reporte_${tipo}_${cliente}_${new Date().toISOString().slice(0,10).replace(/-/g,"")}.csv`;
  a.style.visibility='hidden';document.body.appendChild(a);a.click();document.body.removeChild(a);
}

/* --- LOGICA DE IMPRESIÓN CENTRALIZADA --- */

function getPrintStyle() {
  return `
    <style>
      @media print {
        @page { margin: 15mm; size: auto; }
        body { margin: 0; }
        .no-print { display: none; }
        .page-break { page-break-after: always; }
      }
      body { font-family: 'Arial', sans-serif; font-size: 11px; color: #000; max-width: 700px; margin: auto; padding: 20px; line-height: 1.4; }
      .header-title { display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 20px; font-size: 13px; }
      .row { margin-bottom: 10px; font-size: 11px;}
      .row span { margin-right: 25px; }
      table { width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 20px; }
      th, td { border: 1px solid #000; padding: 5px 8px; text-align: left; }
      .center { text-align: center; }
      .checkboxes { font-family: 'Segoe UI Symbol', sans-serif; }
      .bold { font-weight: bold; }
    </style>
  `;
}

function generarHtmlEtiqueta(d, comps) {
  let fecha = new Date().toLocaleDateString('es-CL');
  let cpuLimpia = d.cpu.replace(/Intel\(R\)|Core\(TM\)/gi, "").trim();

  let trs = comps.map(c => {
     let iconOK   = c.estado === "OK" ? "☑" : "□";
     let iconOBS  = c.estado === "OBS" ? "☑" : "□";
     let iconMalo = c.estado === "MALO" ? "☑" : "□";
     return `<tr>
       <td>${c.nombre}</td>
       <td class="center">${iconOK} OK &nbsp;&nbsp;&nbsp; ${iconOBS} OBS &nbsp;&nbsp;&nbsp; ${iconMalo} MALO</td>
       <td>${c.obs}</td>
     </tr>`;
  }).join("");

  return `
    <div class="header-title">
      <span>ETIQUETAS NOTEBOOK</span>
      <span>FECHA: ${fecha}</span>
    </div>
    <div class="row">
      <span><span class="bold">MODELO:</span> ${d.model}</span>
      <span><span class="bold">PROCESADOR:</span> ${cpuLimpia}</span>
      <span><span class="bold">GUIA:</span> ${d.guia}</span>
      <span><span class="bold">SERIAL:</span> ${d.serial}</span>
    </div>
    <div class="row checkboxes" style="margin-bottom:20px;">
      <span class="bold">ESTADO GLOBAL:</span> &nbsp;&nbsp;&nbsp;&nbsp;
      □ BUENA &nbsp;&nbsp;&nbsp; ☑ POR REPARAR &nbsp;&nbsp;&nbsp; □ MALA &nbsp;&nbsp;&nbsp; □ REPUESTO
    </div>
    <table>
      <thead>
        <tr>
          <th width="20%">REVISIÓN DE COMPONENTES</th>
          <th width="35%" class="center">ESTADO</th>
          <th width="45%">OBSERVACIÓN</th>
        </tr>
      </thead>
      <tbody class="checkboxes">
        ${trs}
      </tbody>
    </table>
    <div class="row checkboxes bold" style="margin-top:20px;">
      CHECKLIST: &nbsp;&nbsp;&nbsp;&nbsp; □ PRUEBA S/O &nbsp;&nbsp;&nbsp;&nbsp; □ HYDRA &nbsp;&nbsp;&nbsp;&nbsp; □ HARDWARE DEFAULT
    </div>
  `;
}

function imprimirEtiqueta(id) {
  let eq = document.getElementById(id);
  if(!eq) return;
  let d = extraerDatos(eq);

  let rawComps = eq.getAttribute('data-comps');
  let comps = [];
  try { comps = JSON.parse(rawComps); } catch(e) {}

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Etiqueta ${d.serial}</title>${getPrintStyle()}</head><body>${generarHtmlEtiqueta(d, comps)}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 350);
}

function imprimirTodasEtiquetas() {
  let etiquetasHtml = [];

  document.querySelectorAll('details').forEach(eq => {
    if (eq.querySelector('.tag-falla')) {
      let d = extraerDatos(eq);
      let rawComps = eq.getAttribute('data-comps');
      let comps = [];
      try { comps = JSON.parse(rawComps); } catch(e) {}

      etiquetasHtml.push(generarHtmlEtiqueta(d, comps));
    }
  });

  if (etiquetasHtml.length === 0) {
    alert("No hay equipos con fallas en este reporte para imprimir.");
    return;
  }

  let contenido = etiquetasHtml.join('<div class="page-break"></div>');

  let html = `<!DOCTYPE html><html><head><meta charset="utf-8"><title>Lote de Etiquetas</title>${getPrintStyle()}</head><body>${contenido}</body></html>`;

  let printWin = window.open('', '', 'width=800,height=600');
  printWin.document.open();
  printWin.document.write(html);
  printWin.document.close();
  printWin.focus();
  setTimeout(() => { printWin.print(); printWin.close(); }, 500);
}
</script>
</head>
<body>
<div class="container">
<div class="header">
  <h1>📋 Inventario Maestro de Equipos</h1>
  <p class="subtitle">Sistema de Control de Entregas y Retiros</p>
  <span class="counter-box">Total Equipos: <span id="total-count">0</span></span>

  <div id="stats-container" class="stats-container"></div>

  <div class="button-group">
    <button class="btn btn-copy"   onclick="copiarFilas()">📋 Copiar Filas (Pegar en Valida)</button>
    <button class="btn btn-export" onclick="exportarValidaCSV()">📥 Descargar CSV</button>
    <button class="btn" style="background:#334155;" onclick="imprimirTodasEtiquetas()">🖨️ Imprimir Lote de Etiquetas</button>
  </div>
</div>
<div id="main-list">"""

HTML_END = "\n</div></div></body></html>"

# ─────────────────────────────────────────────────────────
#  Helpers de formato para python-docx
# ─────────────────────────────────────────────────────────


# ============================================================================
# 5. EXPORT WORD
# ----------------------------------------------------------------------------
# Helpers de formato docx + generar_informe_word: informe técnico de equipos
# defectuosos. Requiere python-docx (DOCX_AVAILABLE).
# ============================================================================
def _set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(space_before * 20)))
    spacing.set(qn("w:after"), str(int(space_after * 20)))
    if line_spacing:
        spacing.set(qn("w:line"), str(int(line_spacing * 240)))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MARCA_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_paragraph_spacing(p, space_before=0, space_after=4)
    return p


def _add_campo(
    doc, etiqueta, valor, etiqueta_bold=True, valor_bold=False, font_size=11
):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=0, space_after=2)
    run_lbl = p.add_run(etiqueta)
    run_lbl.bold = etiqueta_bold
    run_lbl.font.size = Pt(font_size)
    run_lbl.font.name = "Arial"
    run_val = p.add_run(valor)
    run_val.bold = valor_bold
    run_val.font.size = Pt(font_size)
    run_val.font.name = "Arial"
    return p


def _add_titulo_seccion(
    doc, texto, font_size=11, uppercase=True, space_before=10, space_after=4
):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=space_before, space_after=space_after)
    run = p.add_run(texto.upper() if uppercase else texto)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(*MARCA_RGB)
    return p


def _add_falla_item(doc, texto, font_size=11):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, space_before=0, space_after=1)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)
    run = p.add_run(f"- {texto.upper()}")
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    return p


# ─────────────────────────────────────────────────────────
#  Generación de Word (Docx)
# ─────────────────────────────────────────────────────────


def generar_informe_word(cliente, equipos_malos, ruta_cliente):
    if not DOCX_AVAILABLE:
        messagebox.showwarning(
            "Falta Librería", "No se pudo generar el Word porque falta 'python-docx'."
        )
        return

    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        logo_path = os.path.join(os.getcwd(), "lg1.png")

        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            r_logo = p_logo.add_run()
            r_logo.add_picture(logo_path, width=Cm(6.5))
            _set_paragraph_spacing(p_logo, space_before=0, space_after=4)

        p_titulo = doc.add_paragraph()
        _set_paragraph_spacing(p_titulo, space_before=0, space_after=8)
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run("INFORME TÉCNICO DE RETIRO - RESUMEN DE LOTE")
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(*MARCA_RGB)

        _add_horizontal_rule(doc)

        _add_campo(doc, "CLIENTE:             ", cliente.upper())
        _add_campo(doc, "FECHA DE REVISIÓN:   ", fecha_hoy)
        _add_campo(doc, "EQUIPOS DEFECTUOSOS: ", str(len(equipos_malos)))

        _add_horizontal_rule(doc)
        _add_titulo_seccion(
            doc, "DETALLE DE EQUIPOS Y FALLAS", space_before=8, space_after=8
        )

        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Table Grid"

        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "GUÍA"
        hdr_cells[1].text = "MODELO"
        hdr_cells[2].text = "N° DE SERIE"
        hdr_cells[3].text = "DETALLE DE FALLAS"

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for equipo in equipos_malos:
            guia = equipo.get("GUIA_ID", "S/N")
            modelo = equipo.get("MODELO", "Desconocido")
            serie = equipo.get("SERIAL", "Desconocido")
            obs = equipo.get("OBS", "")

            detalle_comps = equipo.get("DETALLE_COMPONENTES", [])
            fallas_list = []
            for c in detalle_comps:
                if c["estado"] != "OK":
                    texto = f"- {c['nombre']}: {c['estado']}"
                    if c["obs"]:
                        texto += f" ({c['obs']})"
                    fallas_list.append(texto)

            texto_fallas = "\n".join(fallas_list)
            if obs and obs not in ("Sin observaciones", ""):
                texto_fallas += f"\nOBS GENERAL: {obs}"

            row_cells = tabla.add_row().cells
            row_cells[0].text = str(guia)
            row_cells[1].text = modelo
            row_cells[2].text = serie
            row_cells[3].text = texto_fallas

        _add_titulo_seccion(doc, "CONCLUSIÓN:", space_before=15, space_after=4)
        p_conclu = doc.add_paragraph()
        run_c = p_conclu.add_run(
            "Los equipos listados en la tabla superior presentan fallas físicas o de hardware. "
            "Se requiere revisión detallada, cambio de componentes afectados o envío a bodega "
            "a la espera de repuestos o reparación."
        )

        for _ in range(3):
            doc.add_paragraph()

        _add_horizontal_rule(doc)
        p_firma = doc.add_paragraph()
        run_f1 = p_firma.add_run("TÉCNICO ENCARGADO: ")
        run_f1.bold = True
        run_f2 = p_firma.add_run("TOMAS GAC\n")
        run_f2.bold = True
        run_r1 = p_firma.add_run("RUT: ")
        run_r1.bold = True
        p_firma.add_run("21.790.634-2")

        fecha_str = datetime.now().strftime("%Y%m%d")
        nombre_archivo = f"Informe_Tecnico_{cliente}_{fecha_str}.docx"
        ruta_guardado = os.path.join(ruta_cliente, nombre_archivo)
        doc.save(ruta_guardado)

    except Exception as e:
        print(f"Error generando Word: {e}")
        messagebox.showerror("Error al generar Word", str(e))


# ============================================================================
# 6. SCANNER (ejecución y parseo)
# ----------------------------------------------------------------------------
# Corre PS_SCRIPT (sección 3) con reintentos/timeout y devuelve el dict `data`
# usado por todo el resto (formulario, guardado, HTML).
# ============================================================================
def get_inventory_fast(max_retries=3, timeout_per_attempt=45):
    last_error = None

    for attempt in range(1, max_retries + 1):
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".ps1", delete=False, encoding="utf-8"
            ) as tmp:
                tmp.write(PS_SCRIPT)
                tmp_path = tmp.name

            si = subprocess.STARTUPINFO()
            si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            si.wShowWindow = 0  # SW_HIDE

            cmd = [
                "powershell",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                tmp_path,
            ]

            result = subprocess.check_output(
                cmd,
                text=True,
                startupinfo=si,
                stderr=subprocess.PIPE,
                timeout=timeout_per_attempt,
            )

            json_match = re.search(r"\{.*\}", result, flags=re.DOTALL)
            if not json_match:
                raise ValueError(
                    f"Salida de PowerShell no contiene JSON válido "
                    f"(intento {attempt}/{max_retries}). "
                    f"Salida recibida: {result[:200]!r}"
                )

            data = json.loads(json_match.group(0))

            try:
                os.unlink(tmp_path)
                tmp_path = None
            except Exception:
                pass

            raw_model = data.get("model", "Desconocido")
            cln_model = re.sub(
                r'(?i)\b(11\.6|12\.5|13\.3|14|15\.6|16|17\.3)\s*(inch|")?\b',
                "",
                raw_model,
            )
            cln_model = re.sub(r"(?i)\bnotebook\s*pc\b", "", cln_model)
            cln_model = re.sub(r"\s+", " ", cln_model).strip()

            ram_list = data.get("ram_rows", [])
            if isinstance(ram_list, str):
                ram_list = [ram_list]
            ram_html = (
                "".join(f"<tr><th>Módulo RAM</th><td>{r}</td></tr>" for r in ram_list)
                or "<tr><th>RAM</th><td>No detectada</td></tr>"
            )

            fabricantes_ram = []
            for r in ram_list:
                partes = r.split("|")
                if len(partes) >= 4:
                    fabricantes_ram.append(partes[3].strip())
            fabricantes_unicos = list(set(fabricantes_ram))
            fabricante_final = (
                " / ".join(fabricantes_unicos) if fabricantes_unicos else "Desconocida"
            )

            partnumbers_ram = []
            for r in ram_list:
                partes = r.split("|")
                if len(partes) >= 5:
                    partnumbers_ram.append(partes[4].strip())
            partnumber_final = (
                " / ".join(set(partnumbers_ram)) if partnumbers_ram else "Desconocida"
            )

            disks = data.get("disks_data", [])
            if isinstance(disks, dict):
                disks = [disks]
            disk_html = ""
            disk_serials = []
            for d in disks:
                desc = d.get("desc", "Desconocido")
                serie = re.sub(r"[^a-zA-Z0-9]", "", d.get("serial", "NA"))[-16:]
                disk_serials.append(serie)
                disk_html += (
                    f"<tr><th>Discos Internos</th><td>{desc}</td></tr>"
                    f"<tr><th>Serie Disco Duro</th><td class='serial-tag'>{serie}</td></tr>"
                )

            net_list = data.get("net_rows", [])
            if isinstance(net_list, str):
                net_list = [net_list]
            net_html = (
                "".join(
                    f"<tr><td>Red Activa</td><td><span class='net-tag'>{n}</span></td></tr>"
                    for n in net_list
                )
                or "<tr><td colspan='2'>Sin red activa</td></tr>"
            )

            return {
                "fecha": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "host": socket.gethostname(),
                "model": cln_model,
                "serial": data.get("serial", "SN_DESCONOCIDO").strip(),
                "key": data.get("win_key", "N/A"),
                "cpu": data.get("cpu", "Desconocido"),
                "ram_html": ram_html,
                "ram_raw": ram_list,
                "disk_html": disk_html,
                "disk_serials": disk_serials,
                "net_rows": net_html,
                "Manufacturer": fabricante_final,
                "Partnumber": partnumber_final,
                "Battery": data.get("battery", "No detectada"),
            }

        except subprocess.TimeoutExpired:
            last_error = (
                f"Timeout tras {timeout_per_attempt}s (intento {attempt}/{max_retries})"
            )
            print(f"⚠️  {last_error}")
        except json.JSONDecodeError as e:
            last_error = f"JSON inválido (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        except Exception as e:
            last_error = f"Error inesperado (intento {attempt}/{max_retries}): {e}"
            print(f"⚠️  {last_error}")
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

        if attempt < max_retries:
            time.sleep(2)

    print(f"❌ Escaneo falló después de {max_retries} intentos. Último: {last_error}")
    return None


_GUIA_CONFIG = os.path.join(os.getcwd(), "Reportes_Guardados", ".guia_retiro.json")


# ============================================================================
# 7. PERSISTENCIA (config de guía, guardado de equipos, empaque, envío a red)
# ============================================================================
def guardar_guia_config(guia: str):
    os.makedirs(os.path.dirname(_GUIA_CONFIG), exist_ok=True)
    with open(_GUIA_CONFIG, "w", encoding="utf-8") as f:
        json.dump({"guia": guia}, f)


def cargar_guia_config() -> str:
    try:
        with open(_GUIA_CONFIG, "r", encoding="utf-8") as f:
            return json.load(f).get("guia", "")
    except Exception:
        return ""


def limpiar_guia_config():
    try:
        if os.path.exists(_GUIA_CONFIG):
            os.remove(_GUIA_CONFIG)
    except Exception:
        pass


# ============================================================================
# PLANTILLA COMPARTIDA DE LA ENTRADA <details> DEL REPORTE HTML
# ----------------------------------------------------------------------------
# Un único constructor usado tanto al guardar en vivo (guardar_equipo_general)
# como al reconstruir desde JSON (_reconstruir_entry_desde_json). Cada caller
# resuelve sus propios valores y los pasa como strings, de modo que la salida
# HTML de cada ruta queda EXACTAMENTE igual que antes (incluida la diferencia
# histórica: la fila de batería solo aparece en la ruta de guardado en vivo,
# vía el parámetro `bateria_html`).
# ============================================================================
def construir_entry_html(
    *,
    safe_id,
    model,
    serial,
    fecha,
    key,
    cpu,
    mov,
    guia_final,
    obs_final,
    ram_html,
    disk_html,
    net_rows,
    comps_json,
    tag_falla_html="",
    falla_html="",
    btn_imprimir="",
    office_html="",
    bateria_html="",
):
    return f"""<details id="{safe_id}" data-comps='{comps_json}' open>
<summary>
  <span class="model-name">{model}</span>
  <span style="font-size:12px;opacity:.85">S/N: {serial} &nbsp;|&nbsp; 📅 {fecha} {tag_falla_html}</span>
</summary>
<div class="content">
  <div class="content-header">
    {btn_imprimir}
    <button class="btn btn-delete" onclick="eliminarEquipo('{safe_id}')">🗑️ Quitar del Reporte</button>
  </div>
  <table>
    <tr class="logistics-row"><th>Tipo Movimiento</th><td><b>{mov}</b></td></tr>
    <tr class="logistics-row"><th>N° de Guía</th><td><b>{guia_final}</b></td></tr>
    {falla_html}
    <tr class="obs-row"><th>⚠️ Obs. General</th><td>{obs_final}</td></tr>
    {office_html}
    <tr><th>Número de Serie</th><td class="serial-tag">{serial}</td></tr>
    <tr><th>Licencia Windows (OA3)</th><td class="win-key">{key}</td></tr>
    <tr><th>Procesador</th><td>{cpu}</td></tr>{bateria_html}
    {ram_html}{disk_html}{net_rows}
  </table>
</div></details>"""


def fijar_total_count(content, total):
    """Reescribe el contador <span id="total-count"> del reporte HTML."""
    return re.sub(
        r'<span id="total-count">\d+</span>',
        f'<span id="total-count">{total}</span>',
        content,
    )


def guardar_equipo_general(
    data, obs, tiene_office, office_key, version_office, mov, guia, detalle_componentes
):
    ruta_base = os.path.join(os.getcwd(), "Reportes_Guardados")
    os.makedirs(ruta_base, exist_ok=True)

    obs_final = obs or "Sin observaciones"
    guia_final = guia or "Sin Guía"
    tipo = "Entregas" if mov == "Entrega" else "Retiros"
    report_file = os.path.join(ruta_base, f"Reporte_{tipo}.html")
    json_file = os.path.join(
        ruta_base, f"data_{re.sub(r'[^a-zA-Z0-9]', '', data['serial'])}.json"
    )
    safe_id = f"dev-{re.sub(r'[^a-zA-Z0-9]', '', data['serial'])}"

    office_html = (
        (
            f'<tr class="office-row"><th>🔑 Office Instalado</th>'
            f"<td><b>Office {version_office}</b> (Key: {office_key})</td></tr>"
        )
        if tiene_office
        else ""
    )

    tag_cls = "tag-entrega" if mov == "Entrega" else "tag-retiro"

    tiene_fallas = any(c["estado"] != "OK" for c in detalle_componentes)
    fallas_str = ", ".join(
        [
            f"{c['nombre']} ({c['estado']})"
            for c in detalle_componentes
            if c["estado"] != "OK"
        ]
    )

    falla_html = ""
    tag_falla_html = ""
    btn_imprimir = ""

    if tiene_fallas:
        tag_falla_html = '<span class="tag-falla">CON FALLAS</span>'
        falla_html = f'<tr class="falla-row"><th>Fallas Detectadas</th><td>{fallas_str}</td></tr>'
        btn_imprimir = f'<button class="btn" style="background:#475569; padding:5px 12px; font-size:11px;" onclick="imprimirEtiqueta(\'{safe_id}\')">🖨️ Etiqueta Individual</button>'

    comps_json = html.escape(json.dumps(detalle_componentes))

    bateria_html = (
        f'\n    <tr><th>🔋 Salud Batería</th>'
        f'<td>{data.get("Battery", "No detectada")}</td></tr>'
    )
    new_entry = construir_entry_html(
        safe_id=safe_id,
        model=data["model"],
        serial=data["serial"],
        fecha=data["fecha"],
        key=data["key"],
        cpu=data["cpu"],
        mov=mov,
        guia_final=guia_final,
        obs_final=obs_final,
        ram_html=data["ram_html"],
        disk_html=data["disk_html"],
        net_rows=data["net_rows"],
        comps_json=comps_json,
        tag_falla_html=tag_falla_html,
        falla_html=falla_html,
        btn_imprimir=btn_imprimir,
        office_html=office_html,
        bateria_html=bateria_html,
    )

    if os.path.exists(report_file):
        with open(report_file, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        content = (
            HTML_START.replace("Inventario Maestro de Equipos", f"Reporte de {tipo}")
            + HTML_END
        )

    content = re.sub(
        rf'<details id="{safe_id}".*?</details>', "", content, flags=re.DOTALL
    )
    content = content.replace(
        '<div id="main-list">', '<div id="main-list">\n' + new_entry
    )

    total = len(re.findall(r"<details ", content))
    content = fijar_total_count(content, total)

    with open(report_file, "w", encoding="utf-8") as f:
        f.write(content)

    sn_hdd = "_".join(data.get("disk_serials", [])) or "PENDIENTE"
    sn_app = (
        office_key.strip()
        if (tiene_office and office_key and office_key.strip())
        else "PENDIENTE"
    )

    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(
            {
                "CLIENTE": "PENDIENTE",
                "SERIAL": data["serial"],
                "SN_Win": data["key"],
                "SN_HDD": sn_hdd,
                "SN_Transf": "PENDIENTE",
                "SN_APP": sn_app,
                "MODELO": data["model"],
                "TIPO_MOVIMIENTO": mov,
                "GUIA_ID": guia_final,
                "OBS": obs_final,
                "DETALLE_COMPONENTES": detalle_componentes,
                "TIENE_FALLAS": tiene_fallas,
                "DATA": data,
                "OFFICE": f"Office {version_office}" if tiene_office else "No",
            },
            f,
            ensure_ascii=False,
            indent=4,
        )


def accion_agregar_lote(
    ventana, data, obs, tiene_office, office_key, version_office, mov, guia, comps_data
):
    detalle_componentes = []
    for c in comps_data:
        est = c["estado"].get()
        ob = c["obs"].get().strip()
        detalle_componentes.append({"nombre": c["nombre"], "estado": est, "obs": ob})

    guardar_equipo_general(
        data,
        obs,
        tiene_office,
        office_key,
        version_office,
        mov,
        guia,
        detalle_componentes,
    )

    if mov == "Retiro" and guia:
        guardar_guia_config(guia)

    tipo = "Entregas" if mov == "Entrega" else "Retiros"
    messagebox.showinfo("✅ Guardado", f"Equipo agregado al archivo 'Reporte_{tipo}'.")
    mostrar_menu_principal(ventana)


def empacar_lote(cliente_clean, ruta_base):
    """
    Mueve los reportes HTML y los data_*.json pendientes a la carpeta del
    cliente (renombrándolos con cliente + fecha), etiqueta cada JSON con el
    cliente y, si hay equipos con fallas, genera el informe Word.

    Lógica pura de archivos/negocio, SIN interacción de UI. Devuelve la lista
    de equipos con fallas para que el handler arme su mensaje.
    """
    ruta_cliente = os.path.join(ruta_base, cliente_clean)
    os.makedirs(ruta_cliente, exist_ok=True)
    fecha_hoy = datetime.now().strftime("%Y%m%d")

    html_e = os.path.join(ruta_base, "Reporte_Entregas.html")
    html_r = os.path.join(ruta_base, "Reporte_Retiros.html")
    for archivo, tipo_doc in [(html_e, "Entregas"), (html_r, "Retiros")]:
        if not os.path.exists(archivo):
            continue

        with open(archivo, "r", encoding="utf-8") as f:
            content = f.read()

        content = content.replace(
            f"Reporte de {tipo_doc}", f"{tipo_doc} - {cliente_clean}"
        )
        content = content.replace(
            f"<title>Reporte de {tipo_doc}</title>",
            f"<title>{tipo_doc} - {cliente_clean}</title>",
        )
        nuevo = os.path.join(
            ruta_cliente, f"Reporte_{tipo_doc}_{cliente_clean}_{fecha_hoy}.html"
        )

        with open(nuevo, "w", encoding="utf-8") as f:
            f.write(content)
        os.remove(archivo)

    equipos_malos = []
    todos_equipos = []
    for archivo in os.listdir(ruta_base):
        if not (archivo.endswith(".json") and archivo.startswith("data_")):
            continue
        ruta_orig = os.path.join(ruta_base, archivo)
        serial_part = archivo[5:-5]
        nuevo_json = os.path.join(
            ruta_cliente, f"Reporte_{cliente_clean}_{fecha_hoy}_{serial_part}.json"
        )
        try:
            with open(ruta_orig, "r", encoding="utf-8") as f:
                jdata = json.load(f)

            jdata["CLIENTE"] = cliente_clean
            todos_equipos.append(jdata)
            if jdata.get("TIENE_FALLAS", False):
                equipos_malos.append(jdata)

            with open(nuevo_json, "w", encoding="utf-8") as f:
                json.dump(jdata, f, ensure_ascii=False, indent=4)

            os.remove(ruta_orig)
        except Exception:
            shutil.move(ruta_orig, nuevo_json)

    # JSON fusionado: un único archivo con todos los equipos del lote (además
    # de los individuales, que siguen usándose en Auditoría y Unir JSON).
    # Solo lleva el resumen logístico; SN_APP y OBS quedan en null si están vacíos.
    if todos_equipos:
        def _resumen_fusionado(eq):
            sn_app = eq.get("SN_APP", "PENDIENTE")
            obs = eq.get("OBS", "")
            return {
                "SERIAL": eq.get("SERIAL", ""),
                "SN_Win": eq.get("SN_Win", ""),
                "SN_HDD": eq.get("SN_HDD", ""),
                "SN_Transf": eq.get("SN_Transf", "PENDIENTE"),
                "SN_APP": None if sn_app in ("PENDIENTE", "", None) else sn_app,
                "OBS": None if obs in ("Sin observaciones", "", None) else obs,
            }

        resumen_equipos = [_resumen_fusionado(eq) for eq in todos_equipos]
        merged_json = os.path.join(
            ruta_cliente, f"Reporte_{cliente_clean}_{fecha_hoy}_FUSIONADO.json"
        )
        with open(merged_json, "w", encoding="utf-8") as f:
            json.dump(resumen_equipos, f, ensure_ascii=False, indent=4)

    if equipos_malos:
        generar_informe_word(cliente_clean, equipos_malos, ruta_cliente)

    return equipos_malos


def accion_finalizar_lote(ventana):
    ruta_base = os.path.join(os.getcwd(), "Reportes_Guardados")
    html_e = os.path.join(ruta_base, "Reporte_Entregas.html")
    html_r = os.path.join(ruta_base, "Reporte_Retiros.html")

    if not os.path.exists(html_e) and not os.path.exists(html_r):
        messagebox.showwarning(
            "Lote Vacío",
            "No hay ningún reporte pendiente.\n\nPrimero escanea y agrega equipos.",
        )
        return

    cliente = simpledialog.askstring(
        "Empacar Lote", "¿A qué CLIENTE pertenecen estos equipos?", parent=ventana
    )
    if not cliente:
        return

    cliente_clean = (
        re.sub(r"[^a-zA-Z0-9\s_\-]", "", cliente).strip().replace(" ", "_").upper()
        or "GENERAL"
    )

    equipos_malos = empacar_lote(cliente_clean, ruta_base)

    if equipos_malos:
        aviso_extra = f"\n\n📄 Se generó el 'Informe_Tecnico_{cliente_clean}.docx' con {len(equipos_malos)} equipo(s) defectuoso(s)."
    else:
        aviso_extra = ""

    messagebox.showinfo(
        "📦 Lote Finalizado",
        f"¡Lote empacado con éxito!\n\nCarpeta: {cliente_clean}{aviso_extra}",
    )
    limpiar_guia_config()


def accion_enviar_red(ventana):
    carpeta_local = os.path.join(os.getcwd(), "Reportes_Guardados")
    
    if not os.path.exists(carpeta_local):
        messagebox.showwarning("Aviso", "No hay reportes locales para enviar.")
        return

    # Detectar solo las carpetas de clientes
    carpetas = [d for d in os.listdir(carpeta_local) if os.path.isdir(os.path.join(carpeta_local, d))]
    
    if not carpetas:
        pendiente = any(os.path.exists(os.path.join(carpeta_local, f)) for f in ("Reporte_Entregas.html", "Reporte_Retiros.html"))
        if pendiente:
            messagebox.showwarning("Lote Incompleto", "Hay equipos pendientes.\n\nPresiona '📦 EMPACAR LOTE' antes de enviar a la red.")
        else:
            messagebox.showwarning("Aviso", "No hay carpetas de clientes empacadas para enviar.")
        return

    # Validar si el servidor de red es accesible antes de intentar copiar
    if not os.path.exists(CARPETA_DESTINO):
        messagebox.showerror("Error de Red", f"No se pudo acceder al servidor:\n{CARPETA_DESTINO}\n\nVerifica tu conexión VPN o Wi-Fi.")
        return

    respuesta = messagebox.askyesno("Confirmar Envío", f"¿Deseas enviar las {len(carpetas)} carpetas de clientes al servidor?\n\nDestino: {CARPETA_DESTINO}")
    if not respuesta:
        return

    # --- Ventana de Carga Visual ---
    win_carga = tk.Toplevel(ventana)
    win_carga.title("Enviando...")
    win_carga.geometry("350x150")
    win_carga.configure(bg=COLORS["fondo"])
    win_carga.grab_set() # Bloquea la ventana principal para evitar que sigan usando el programa
    
    tk.Label(win_carga, text="🚀", font=("Segoe UI", 24), bg=COLORS["fondo"]).pack(pady=(15, 5))
    tk.Label(win_carga, text="Copiando archivos a la red...", font=("Segoe UI", 10, "bold"), bg=COLORS["fondo"], fg=COLORS["azul"]).pack()
    tk.Label(win_carga, text="Por favor, no cierres el programa.", font=("Segoe UI", 9), bg=COLORS["fondo"], fg=COLORS["gris"]).pack()

    # --- Función que hace el trabajo pesado en segundo plano ---
    def tarea_envio():
        errores = 0
        enviados = 0
        
        for cliente_folder in carpetas:
            ruta_origen = os.path.join(carpeta_local, cliente_folder)
            ruta_destino = os.path.join(CARPETA_DESTINO, cliente_folder)

            try:
                # dirs_exist_ok=True es vital: mezcla el contenido si el cliente ya existía en el NAS
                shutil.copytree(ruta_origen, ruta_destino, dirs_exist_ok=True)
                enviados += 1
            except Exception as e:
                errores += 1
                print(f"Error copiando {cliente_folder}: {e}")

        # Una vez terminado, avisamos a la interfaz gráfica principal
        ventana.after(0, finalizar_envio, enviados, errores)

    # --- Función que cierra la carga y muestra el resultado ---
    def finalizar_envio(enviados, errores):
        if win_carga.winfo_exists():
            win_carga.destroy()
            
        if errores > 0:
            messagebox.showwarning("Proceso con Observaciones", f"Se copiaron {enviados} carpetas, pero hubo {errores} error(es).\n\nRevisa si hay archivos abiertos por otro usuario en la red.")
        else:
            messagebox.showinfo("✅ Envío Exitoso", f"Se enviaron {enviados} carpetas a la red correctamente.")

    # Lanzamos el copiado en un hilo (Thread) para que la pantalla no se congele
    threading.Thread(target=tarea_envio, daemon=True).start()

# ============================================================================
# 8. RECONSTRUCCIÓN / MERGE — rearmar el HTML a partir de los JSON guardados
# ============================================================================
def _reconstruir_entry_desde_json(jdata):
    data = jdata.get("DATA", {})
    obs_final = jdata.get("OBS", "Sin observaciones")
    guia_final = jdata.get("GUIA_ID", "Sin Guía")
    mov = jdata.get("TIPO_MOVIMIENTO", "Entrega")
    detalle_componentes = jdata.get("DETALLE_COMPONENTES", [])
    office_str = jdata.get("OFFICE", "No")

    office_html = ""
    if office_str and office_str != "No":
        office_html = (
            f'<tr class="office-row"><th>🔑 Office Instalado</th>'
            f"<td><b>{office_str}</b></td></tr>"
        )

    safe_id = f"dev-{re.sub(r'[^a-zA-Z0-9]', '', data.get('serial', 'DESCONOCIDO'))}"
    tag_cls = "tag-entrega" if mov == "Entrega" else "tag-retiro"

    tiene_fallas = jdata.get("TIENE_FALLAS", False)
    fallas_str = ", ".join(
        [
            f"{c['nombre']} ({c['estado']})"
            for c in detalle_componentes
            if c.get("estado", "OK") != "OK"
        ]
    )

    falla_html = ""
    tag_falla_html = ""
    btn_imprimir = ""

    if tiene_fallas:
        tag_falla_html = '<span class="tag-falla">CON FALLAS</span>'
        falla_html = f'<tr class="falla-row"><th>Fallas Detectadas</th><td>{fallas_str}</td></tr>'
        btn_imprimir = (
            f'<button class="btn" style="background:#475569; padding:5px 12px; font-size:11px;" '
            f"onclick=\"imprimirEtiqueta('{safe_id}')\">🖨️ Etiqueta Individual</button>"
        )

    comps_json = html.escape(json.dumps(detalle_componentes))

    ram_html = data.get("ram_html", "<tr><th>RAM</th><td>No detectada</td></tr>")
    disk_html = data.get("disk_html", "")
    net_rows = data.get("net_rows", "<tr><td colspan='2'>Sin red activa</td></tr>")

    return construir_entry_html(
        safe_id=safe_id,
        model=data.get("model", "Desconocido"),
        serial=data.get("serial", "?"),
        fecha=data.get("fecha", ""),
        key=data.get("key", "N/A"),
        cpu=data.get("cpu", "Desconocido"),
        mov=mov,
        guia_final=guia_final,
        obs_final=obs_final,
        ram_html=ram_html,
        disk_html=disk_html,
        net_rows=net_rows,
        comps_json=comps_json,
        tag_falla_html=tag_falla_html,
        falla_html=falla_html,
        btn_imprimir=btn_imprimir,
        office_html=office_html,
    )


def accion_unir_Json(ventana):
    carpeta = filedialog.askdirectory(
        title="Selecciona la carpeta del cliente (con los archivos .json)",
        parent=ventana,
    )
    if not carpeta:
        return

    archivos_json = [
        f for f in os.listdir(carpeta) if f.endswith(".json") and not f.startswith(".")
    ]

    if not archivos_json:
        messagebox.showwarning(
            "Sin JSONs",
            f"No se encontraron archivos .json en:\n{carpeta}",
            parent=ventana,
        )
        return

    equipos_por_tipo = {"Entregas": [], "Retiros": []}
    errores = []

    for archivo in sorted(archivos_json):
        ruta_json = os.path.join(carpeta, archivo)
        try:
            with open(ruta_json, "r", encoding="utf-8") as f:
                jdata = json.load(f)
            # El archivo fusionado (lista de equipos) ya está cubierto por los
            # JSON individuales; se ignora para no duplicar.
            if isinstance(jdata, list):
                continue
            if "DATA" not in jdata or "SERIAL" not in jdata:
                errores.append(f"{archivo}: estructura inválida (falta DATA o SERIAL)")
                continue
            mov = jdata.get("TIPO_MOVIMIENTO", "Entrega")
            tipo = "Entregas" if mov == "Entrega" else "Retiros"
            equipos_por_tipo[tipo].append(jdata)
        except Exception as e:
            errores.append(f"{archivo}: {e}")

    total_equipos = sum(len(v) for v in equipos_por_tipo.values())
    if total_equipos == 0:
        msg = "No se pudo leer ningún JSON válido."
        if errores:
            msg += "\n\nErrores:\n" + "\n".join(errores)
        messagebox.showerror("Error", msg, parent=ventana)
        return

    nombre_carpeta = os.path.basename(carpeta)
    fecha_hoy = datetime.now().strftime("%Y%m%d")
    archivos_generados = []

    for tipo, equipos in equipos_por_tipo.items():
        if not equipos:
            continue

        titulo = f"{tipo} - {nombre_carpeta}"

        contenido = HTML_START.replace("Inventario Maestro de Equipos", titulo).replace(
            "<title>Inventario Maestro</title>", f"<title>{titulo}</title>"
        )

        entries = []
        for jdata in equipos:
            try:
                entries.append(_reconstruir_entry_desde_json(jdata))
            except Exception as e:
                errores.append(f"S/N {jdata.get('SERIAL', '?')}: {e}")

        contenido += "\n".join(entries)
        contenido += HTML_END

        total = len(entries)
        contenido = fijar_total_count(contenido, total)

        nombre_html = f"Reporte_{tipo}_{nombre_carpeta}_{fecha_hoy}_RECONSTRUIDO.html"
        ruta_salida = os.path.join(carpeta, nombre_html)
        with open(ruta_salida, "w", encoding="utf-8") as f:
            f.write(contenido)
        archivos_generados.append(nombre_html)

    resumen = f"✅ HTML reconstruido exitosamente desde {total_equipos} equipo(s).\n\n"
    resumen += "Archivos generados:\n" + "\n".join(
        f"  • {a}" for a in archivos_generados
    )
    resumen += f"\n\nUbicación:\n{carpeta}"
    if errores:
        resumen += f"\n\n⚠️ Advertencias ({len(errores)}):\n" + "\n".join(errores)

    messagebox.showinfo("🔧 Reconstrucción Completa", resumen, parent=ventana)


# ============================================================================
# 9. MÓDULOS (ventanas Toplevel) — auditorías, comparación Excel, etiquetas,
#    cambio de nombre del equipo.
# ============================================================================
def abrir_modulo_auditoria(ventana_padre):
    win = ventana_modal(ventana_padre, "Auditoría — Entregas vs Retiros", "860x520")
    titulo_ui(win, "🔍 Auditoría Diferencial (Entregas vs Retiros)")

    frame_rutas = tk.Frame(win, bg=COLORS["fondo"])
    frame_rutas.pack(fill="x", padx=20, pady=4)
    ruta_entrega = tk.StringVar()
    ruta_retiro = tk.StringVar()

    def sel_carpeta(var, titulo):
        c = filedialog.askdirectory(title=titulo, parent=win)
        if c:
            var.set(c)

    fila_selector(
        frame_rutas,
        0,
        "📂 Carpeta Entregas",
        lambda: sel_carpeta(ruta_entrega, "Carpeta de ENTREGAS"),
        ruta_entrega,
    )
    fila_selector(
        frame_rutas,
        1,
        "📂 Carpeta Retiros",
        lambda: sel_carpeta(ruta_retiro, "Carpeta de RETIROS"),
        ruta_retiro,
    )

    tree = treeview_auditoria(
        win,
        [
            ("estado", 105, "ESTADO"),
            ("serial", 130, "S/N"),
            ("modelo", 190, "MODELO"),
            ("detalle", 400, "DETALLE"),
        ],
        height=13,
        con_alterado=True,
    )

    def cargar_jsons(ruta):
        datos = {}
        for arch in os.listdir(ruta):
            if arch.endswith(".json"):
                try:
                    with open(os.path.join(ruta, arch), "r", encoding="utf-8") as f:
                        j = json.load(f)
                    if "SERIAL" in j:
                        datos[j["SERIAL"]] = j
                except Exception:
                    pass
        return datos

    def ejecutar():
        tree.delete(*tree.get_children())
        if not ruta_entrega.get() or not ruta_retiro.get():
            return messagebox.showwarning(
                "Faltan Datos", "Selecciona ambas carpetas.", parent=win
            )
        entregas = cargar_jsons(ruta_entrega.get())
        retiros = cargar_jsons(ruta_retiro.get())
        if not entregas:
            return messagebox.showerror(
                "Error", "Carpeta de Entregas sin JSONs válidos.", parent=win
            )
        if not retiros:
            return messagebox.showerror(
                "Error", "Carpeta de Retiros sin JSONs válidos.", parent=win
            )

        for serial, d_e in entregas.items():
            modelo = d_e.get("MODELO", "Desconocido")
            if serial not in retiros:
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❌ FALTANTE",
                        serial,
                        modelo,
                        "Equipo de entrega no está en retiro.",
                    ),
                    tags=("falta",),
                )
                continue
            d_r = retiros[serial]
            alertas = []
            if len(d_e["DATA"].get("ram_raw", [])) != len(
                d_r["DATA"].get("ram_raw", [])
            ):
                alertas.append(f"RAM alterada.")
            s_e = {
                d.get("serial", "NA")
                for d in d_e["DATA"].get("disks_data", [])
                if isinstance(d, dict)
            }
            s_r = {
                d.get("serial", "NA")
                for d in d_r["DATA"].get("disks_data", [])
                if isinstance(d, dict)
            }
            if s_e != s_r:
                alertas.append("DISCO serial distinto.")
            tag = "alterado" if alertas else "ok"
            tree.insert(
                "",
                tk.END,
                values=(
                    "⚠️ ALTERADO" if alertas else "✅ OK",
                    serial,
                    modelo,
                    " | ".join(alertas) if alertas else "Hardware coincide.",
                ),
                tags=(tag,),
            )
        for serial, d_r in retiros.items():
            if serial not in entregas:
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❓ SOBRANTE",
                        serial,
                        d_r.get("MODELO", "Desconocido"),
                        "Equipo en retiro no figura en entregas.",
                    ),
                    tags=("sobra",),
                )

    boton_accion(win, "⚙️  INICIAR AUDITORÍA", ejecutar).pack(pady=10)


def abrir_modulo_auditoria_mixta(ventana_padre):
    win = ventana_modal(ventana_padre, "Auditoría — Sistema vs Físico", "900x550")
    titulo_ui(win, "📊 Auditoría Cruzada (Teórico vs Físico)")

    frame_rutas = tk.Frame(win, bg=COLORS["fondo"])
    frame_rutas.pack(fill="x", padx=20, pady=4)
    ruta_base = tk.StringVar()
    ruta_fisico = tk.StringVar()

    def sel_archivo(var, titulo, tipos):
        arch = filedialog.askopenfilename(title=titulo, parent=win, filetypes=tipos)
        if arch:
            var.set(arch)

    fila_selector(
        frame_rutas,
        0,
        "📂 Cargar Excel Base",
        lambda: sel_archivo(
            ruta_base, "Seleccionar Excel Base", [("Archivos Excel", "*.xlsx *.xls")]
        ),
        ruta_base,
    )
    fila_selector(
        frame_rutas,
        1,
        "📂 Cargar Físico",
        lambda: sel_archivo(
            ruta_fisico,
            "Seleccionar Reporte Físico",
            [("HTML o Excel", "*.html *.xlsx *.xls")],
        ),
        ruta_fisico,
    )

    tree = treeview_auditoria(
        win,
        [
            ("estado", 110, "ESTADO"),
            ("serial", 150, "N° SERIE"),
            ("modelo", 250, "MODELO / INFO"),
            ("origen", 280, "DETALLE"),
        ],
        height=12,
        con_alterado=False,
    )

    def extraer_excel(ruta):
        equipos = {}
        try:
            df = pd.read_excel(ruta)
            df.columns = df.columns.str.strip().str.upper()
            col_serie = None
            for col in df.columns:
                if col in [
                    "NUMERO DE SERIE",
                    "NÚMERO DE SERIE",
                    "SERIAL",
                    "S/N",
                    "SN",
                    "SERIE",
                ]:
                    col_serie = col
                    break
            if not col_serie:
                for col in df.columns:
                    if "SERIE" in str(col) and "DISCO" not in str(col):
                        col_serie = col
                        break
            if col_serie:
                col_modelo = next(
                    (c for c in df.columns if "MODEL" in str(c) or "EQUIPO" in str(c)),
                    None,
                )
                for _, row in df.iterrows():
                    sn = str(row[col_serie]).strip()
                    if sn and sn.lower() != "nan":
                        equipos[sn.upper()] = (
                            str(row[col_modelo]).strip() if col_modelo else "Excel"
                        )
        except Exception as e:
            messagebox.showerror(
                "Error Excel", f"No se pudo leer {ruta}\n{e}", parent=win
            )
        return equipos

    def extraer_html(ruta):
        equipos = {}
        try:
            with open(ruta, "r", encoding="utf-8") as f:
                content = f.read()
            bloques = re.findall(r"<details.*?</details>", content, flags=re.DOTALL)
            for bloque in bloques:
                s_match = re.search(
                    r'<tr><th>Número de Serie</th><td class="serial-tag">(.*?)</td></tr>',
                    bloque,
                )
                m_match = re.search(r'<span class="model-name">(.*?)</span>', bloque)
                if s_match:
                    equipos[s_match.group(1).strip().upper()] = (
                        m_match.group(1).strip() if m_match else "HTML"
                    )
        except Exception as e:
            messagebox.showerror(
                "Error HTML", f"No se pudo leer {ruta}\n{e}", parent=win
            )
        return equipos

    def ejecutar():
        tree.delete(*tree.get_children())
        r_base, r_fis = ruta_base.get(), ruta_fisico.get()
        if not r_base or not r_fis:
            return messagebox.showwarning(
                "Faltan Datos", "Selecciona ambos archivos.", parent=win
            )
        dict_base = extraer_excel(r_base)
        if not dict_base:
            return
        dict_fis = (
            extraer_html(r_fis) if r_fis.endswith(".html") else extraer_excel(r_fis)
        )
        if not dict_fis:
            return

        for sn, mod in dict_base.items():
            if sn in dict_fis:
                tree.insert(
                    "",
                    tk.END,
                    values=("✅ MATCH", sn, mod, "Encontrado en ambos."),
                    tags=("ok",),
                )
            else:
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❌ FALTANTE",
                        sn,
                        mod,
                        "Está en la Base, pero NO en el Físico.",
                    ),
                    tags=("falta",),
                )
        for sn, mod in dict_fis.items():
            if sn not in dict_base:
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❓ SOBRANTE",
                        sn,
                        mod,
                        "Físico no figura en el Excel Base.",
                    ),
                    tags=("sobra",),
                )

    boton_accion(win, "⚙️  CRUZAR INVENTARIOS", ejecutar).pack(pady=10)


def abrir_modulo_comparacion_excel(ventana_padre):
    if not PANDAS_AVAILABLE:
        messagebox.showerror(
            "Falta Librería",
            "Este módulo requiere 'pandas' y 'openpyxl'.\n\n"
            "Instálalos con:\n  pip install pandas openpyxl",
            parent=ventana_padre,
        )
        return

    win = ventana_modal(
        ventana_padre, "Comparación Excel — Entregas vs Retiros", "1020x640"
    )
    titulo_ui(win, "📊 Comparación de Excel: Entregas vs Retiros", pady=(14, 4))
    tk.Label(
        win,
        text="Detecta equipos faltantes, sobrantes y diferencias de hardware entre ambos reportes.",
        font=("Segoe UI", 9),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(0, 10))

    frame_rutas = tk.Frame(win, bg=COLORS["fondo"])
    frame_rutas.pack(fill="x", padx=24, pady=4)

    ruta_entrega = tk.StringVar()
    ruta_retiro = tk.StringVar()

    style_btn = dict(
        bg="#1a3c6e",
        fg="white",
        font=("Segoe UI", 9, "bold"),
        cursor="hand2",
        relief="flat",
        padx=8,
        pady=4,
    )

    def sel_excel(var, titulo):
        arch = filedialog.askopenfilename(
            title=titulo,
            parent=win,
            filetypes=[("Archivos Excel", "*.xlsx *.xls"), ("Todos", "*.*")],
        )
        if arch:
            var.set(arch)

    tk.Button(
        frame_rutas,
        text="📂 Excel ENTREGAS",
        command=lambda: sel_excel(ruta_entrega, "Seleccionar Excel de Entregas"),
        **style_btn,
    ).grid(row=0, column=0, padx=(0, 8), pady=4, sticky="ew")
    tk.Entry(
        frame_rutas,
        textvariable=ruta_entrega,
        width=60,
        state="readonly",
        font=("Segoe UI", 9),
        bg="#fff",
    ).grid(row=0, column=1, pady=4, sticky="ew")

    tk.Button(
        frame_rutas,
        text="📂 Excel RETIROS",
        command=lambda: sel_excel(ruta_retiro, "Seleccionar Excel de Retiros"),
        **style_btn,
    ).grid(row=1, column=0, padx=(0, 8), pady=4, sticky="ew")
    tk.Entry(
        frame_rutas,
        textvariable=ruta_retiro,
        width=60,
        state="readonly",
        font=("Segoe UI", 9),
        bg="#fff",
    ).grid(row=1, column=1, pady=4, sticky="ew")

    frame_rutas.columnconfigure(1, weight=1)

    frame_opciones = tk.Frame(win, bg=COLORS["fondo"])
    frame_opciones.pack(fill="x", padx=24, pady=(4, 0))

    tk.Label(
        frame_opciones, text="Columna N° Serie:", bg=COLORS["fondo"], font=("Segoe UI", 9)
    ).pack(side="left")
    entry_col_serie = tk.Entry(frame_opciones, width=22, font=("Segoe UI", 9))
    entry_col_serie.insert(0, "NUMERO DE SERIE")
    entry_col_serie.pack(side="left", padx=6)

    tk.Label(
        frame_opciones,
        text="(si está vacío, se detecta automáticamente)",
        bg=COLORS["fondo"],
        font=("Segoe UI", 8),
        fg="#94a3b8",
    ).pack(side="left")

    frame_resumen = tk.Frame(win, bg=COLORS["fondo"])
    frame_resumen.pack(fill="x", padx=24, pady=(8, 0))

    lbl_ok = tk.Label(
        frame_resumen,
        text="✅ Match: —",
        bg="#dcfce7",
        fg="#15803d",
        font=("Segoe UI", 9, "bold"),
        padx=10,
        pady=4,
        relief="flat",
    )
    lbl_falta = tk.Label(
        frame_resumen,
        text="❌ Faltantes: —",
        bg="#fee2e2",
        fg="#dc2626",
        font=("Segoe UI", 9, "bold"),
        padx=10,
        pady=4,
        relief="flat",
    )
    lbl_sobra = tk.Label(
        frame_resumen,
        text="❓ Sobrantes: —",
        bg="#fef3c7",
        fg="#d97706",
        font=("Segoe UI", 9, "bold"),
        padx=10,
        pady=4,
        relief="flat",
    )
    lbl_diff = tk.Label(
        frame_resumen,
        text="⚠️ Diferencias: —",
        bg="#fce7f3",
        fg="#be185d",
        font=("Segoe UI", 9, "bold"),
        padx=10,
        pady=4,
        relief="flat",
    )
    for w in (lbl_ok, lbl_falta, lbl_sobra, lbl_diff):
        w.pack(side="left", padx=4)

    style = ttk.Style()
    style.configure("Cmp.Treeview", rowheight=22, font=("Segoe UI", 9))
    style.configure("Cmp.Treeview.Heading", font=("Segoe UI", 9, "bold"))

    cols = ("estado", "serial", "modelo_e", "modelo_r", "diferencias")
    frame_tree = tk.Frame(win, bg=COLORS["fondo"])
    frame_tree.pack(fill="both", expand=True, padx=20, pady=8)

    tree = ttk.Treeview(
        frame_tree, columns=cols, show="headings", height=14, style="Cmp.Treeview"
    )

    col_cfg = [
        ("estado", "ESTADO", 115, "center"),
        ("serial", "N° SERIE", 145, "center"),
        ("modelo_e", "MODELO (Entrega)", 200, "w"),
        ("modelo_r", "MODELO (Retiro)", 200, "w"),
        ("diferencias", "DIFERENCIAS", 310, "w"),
    ]
    for cid, txt, w, anchor in col_cfg:
        tree.heading(cid, text=txt, command=lambda c=cid: _ordenar_columna(tree, c))
        tree.column(cid, width=w, anchor=anchor, minwidth=60)

    tree.tag_configure("ok", foreground="#15803d", background="#f0fdf4")
    tree.tag_configure("falta", foreground="#dc2626", background="#fef2f2")
    tree.tag_configure("sobra", foreground="#d97706", background="#fffbeb")
    tree.tag_configure("diff", foreground="#be185d", background="#fdf4ff")

    sb_v = ttk.Scrollbar(frame_tree, orient="vertical", command=tree.yview)
    sb_h = ttk.Scrollbar(frame_tree, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=sb_v.set, xscrollcommand=sb_h.set)
    tree.grid(row=0, column=0, sticky="nsew")
    sb_v.grid(row=0, column=1, sticky="ns")
    sb_h.grid(row=1, column=0, sticky="ew")
    frame_tree.rowconfigure(0, weight=1)
    frame_tree.columnconfigure(0, weight=1)

    _resultado_cache = []

    def _ordenar_columna(tv, col):
        datos = [(tv.set(k, col), k) for k in tv.get_children("")]
        datos.sort()
        for i, (_, k) in enumerate(datos):
            tv.move(k, "", i)

    COLS_SERIE_CANDIDATAS = [
        "NUMERO DE SERIE",
        "NÚMERO DE SERIE",
        "N° SERIE",
        "N°SERIE",
        "SERIAL",
        "S/N",
        "SN",
        "SERIE",
        "NRO SERIE",
    ]
    COLS_MODELO_CANDIDATAS = [
        "MODELO",
        "NOMBRE EQUIPO",
        "EQUIPO",
        "DESCRIPCION",
        "DESCRIPCIÓN",
        "NOMBRE",
        "DEVICE",
    ]
    COLS_HARDWARE = {
        "RAM": ["RAM", "MEMORIA", "MEMORIA RAM"],
        "DISCO": ["DISCO", "HDD", "SSD", "ALMACENAMIENTO", "DISCO DURO"],
        "CPU": ["CPU", "PROCESADOR", "PROCESSOR"],
    }

    def _detectar_columna(cols_df, candidatas):
        cols_upper = {c.strip().upper(): c for c in cols_df}
        for cand in candidatas:
            if cand in cols_upper:
                return cols_upper[cand]
        for cand in candidatas:
            for col in cols_df:
                if cand in col.strip().upper():
                    return col
        return None

    def _leer_excel(ruta, col_serie_hint=""):
        try:
            df = pd.read_excel(ruta, dtype=str)
        except Exception as e:
            raise RuntimeError(f"No se pudo leer '{os.path.basename(ruta)}':\n{e}")

        df.columns = [str(c).strip() for c in df.columns]

        col_serie = None
        if col_serie_hint.strip():
            col_serie = _detectar_columna(df.columns, [col_serie_hint.strip().upper()])
        if not col_serie:
            col_serie = _detectar_columna(df.columns, COLS_SERIE_CANDIDATAS)

        if not col_serie:
            raise RuntimeError(
                f"No se encontró la columna de N° de Serie en:\n{os.path.basename(ruta)}\n\n"
                f"Columnas disponibles:\n{', '.join(df.columns[:15])}"
            )

        col_modelo = _detectar_columna(df.columns, COLS_MODELO_CANDIDATAS)

        hw_cols = {}
        for hw_key, candidatas in COLS_HARDWARE.items():
            c = _detectar_columna(df.columns, candidatas)
            if c:
                hw_cols[hw_key] = c

        equipos = {}
        for _, row in df.iterrows():
            sn = str(row.get(col_serie, "")).strip()
            if not sn or sn.lower() in ("nan", "", "none"):
                continue
            sn_key = re.sub(r"\s+", "", sn).upper()
            fila = {
                "SERIAL_ORIG": sn,
                "MODELO": str(row.get(col_modelo, "")).strip() if col_modelo else "",
            }
            for hw_key, hw_col in hw_cols.items():
                fila[hw_key] = str(row.get(hw_col, "")).strip()
            equipos[sn_key] = fila

        return equipos, list(hw_cols.keys())

    def ejecutar_comparacion():
        nonlocal _resultado_cache
        _resultado_cache = []
        tree.delete(*tree.get_children())

        r_e = ruta_entrega.get()
        r_r = ruta_retiro.get()
        col_hint = entry_col_serie.get().strip()

        if not r_e or not r_r:
            messagebox.showwarning(
                "Faltan archivos", "Selecciona ambos archivos Excel.", parent=win
            )
            return

        try:
            dict_e, hw_e = _leer_excel(r_e, col_hint)
            dict_r, hw_r = _leer_excel(r_r, col_hint)
        except RuntimeError as err:
            messagebox.showerror("Error al leer Excel", str(err), parent=win)
            return

        hw_comunes = [h for h in hw_e if h in hw_r]
        cnt_ok, cnt_falta, cnt_sobra, cnt_diff = 0, 0, 0, 0

        for sn, d_e in dict_e.items():
            if sn not in dict_r:
                cnt_falta += 1
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❌ FALTANTE",
                        d_e["SERIAL_ORIG"],
                        d_e.get("MODELO", ""),
                        "—",
                        "No figura en el Excel de Retiros",
                    ),
                    tags=("falta",),
                )
                _resultado_cache.append(
                    {
                        "ESTADO": "FALTANTE",
                        "N° SERIE": d_e["SERIAL_ORIG"],
                        "MODELO (Entrega)": d_e.get("MODELO", ""),
                        "MODELO (Retiro)": "",
                        "DIFERENCIAS": "No figura en el Excel de Retiros",
                    }
                )
                continue

            d_r = dict_r[sn]
            diffs = []

            mod_e = d_e.get("MODELO", "").strip().upper()
            mod_r = d_r.get("MODELO", "").strip().upper()
            if mod_e and mod_r and mod_e != mod_r:
                diffs.append(f"MODELO: '{d_e['MODELO']}' → '{d_r['MODELO']}'")

            for hw in hw_comunes:
                val_e = d_e.get(hw, "").strip().upper()
                val_r = d_r.get(hw, "").strip().upper()
                if val_e and val_r and val_e != val_r:
                    diffs.append(f"{hw}: '{d_e.get(hw, '')}' → '{d_r.get(hw, '')}'")

            if diffs:
                cnt_diff += 1
                tag = "diff"
                estado = "⚠️ DIFERENCIA"
            else:
                cnt_ok += 1
                tag = "ok"
                estado = "✅ MATCH"

            diff_str = " | ".join(diffs) if diffs else "Hardware coincide"
            tree.insert(
                "",
                tk.END,
                values=(
                    estado,
                    d_e["SERIAL_ORIG"],
                    d_e.get("MODELO", ""),
                    d_r.get("MODELO", ""),
                    diff_str,
                ),
                tags=(tag,),
            )
            _resultado_cache.append(
                {
                    "ESTADO": estado.replace("✅ ", "").replace("⚠️ ", ""),
                    "N° SERIE": d_e["SERIAL_ORIG"],
                    "MODELO (Entrega)": d_e.get("MODELO", ""),
                    "MODELO (Retiro)": d_r.get("MODELO", ""),
                    "DIFERENCIAS": diff_str,
                }
            )

        for sn, d_r in dict_r.items():
            if sn not in dict_e:
                cnt_sobra += 1
                tree.insert(
                    "",
                    tk.END,
                    values=(
                        "❓ SOBRANTE",
                        d_r["SERIAL_ORIG"],
                        "—",
                        d_r.get("MODELO", ""),
                        "No figura en el Excel de Entregas",
                    ),
                    tags=("sobra",),
                )
                _resultado_cache.append(
                    {
                        "ESTADO": "SOBRANTE",
                        "N° SERIE": d_r["SERIAL_ORIG"],
                        "MODELO (Entrega)": "",
                        "MODELO (Retiro)": d_r.get("MODELO", ""),
                        "DIFERENCIAS": "No figura en el Excel de Entregas",
                    }
                )

        lbl_ok.config(text=f"✅ Match: {cnt_ok}")
        lbl_falta.config(text=f"❌ Faltantes: {cnt_falta}")
        lbl_sobra.config(text=f"❓ Sobrantes: {cnt_sobra}")
        lbl_diff.config(text=f"⚠️ Diferencias: {cnt_diff}")

        total = cnt_ok + cnt_falta + cnt_sobra + cnt_diff
        if total == 0:
            messagebox.showinfo(
                "Sin resultados",
                "No se encontraron equipos válidos en los archivos.",
                parent=win,
            )

    def exportar_resultado():
        if not _resultado_cache:
            messagebox.showwarning(
                "Sin datos", "Primero ejecuta la comparación.", parent=win
            )
            return

        ruta_out = filedialog.asksaveasfilename(
            title="Guardar resultado como Excel",
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile=f"Comparacion_Entregas_Retiros_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            parent=win,
        )
        if not ruta_out:
            return

        try:
            import openpyxl
            from openpyxl.styles import Alignment, Font, PatternFill
        except ImportError:
            messagebox.showerror(
                "Falta librería",
                "Se requiere 'openpyxl' para exportar.\n\npip install openpyxl",
                parent=win,
            )
            return

        try:
            df_out = pd.DataFrame(_resultado_cache)
            with pd.ExcelWriter(ruta_out, engine="openpyxl") as writer:
                df_out.to_excel(writer, index=False, sheet_name="Comparacion")
                ws = writer.sheets["Comparacion"]

                FILLS = {
                    "MATCH": PatternFill("solid", fgColor="D1FAE5"),
                    "FALTANTE": PatternFill("solid", fgColor="FEE2E2"),
                    "SOBRANTE": PatternFill("solid", fgColor="FEF3C7"),
                    "DIFERENCIA": PatternFill("solid", fgColor="FAE8FF"),
                }
                header_fill = PatternFill("solid", fgColor=MARCA_HEX)
                header_font = Font(color="FFFFFF", bold=True)

                for cell in ws[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal="center")

                for row in ws.iter_rows(min_row=2):
                    estado_val = str(row[0].value or "").upper()
                    fill = FILLS.get(estado_val, PatternFill())
                    for cell in row:
                        cell.fill = fill

                for col_cells in ws.columns:
                    max_len = max(
                        (len(str(c.value or "")) for c in col_cells), default=10
                    )
                    ws.column_dimensions[col_cells[0].column_letter].width = min(
                        max_len + 4, 60
                    )

            messagebox.showinfo(
                "✅ Exportado",
                f"Resultado guardado en:\n{ruta_out}",
                parent=win,
            )
        except Exception as e:
            messagebox.showerror("Error al exportar", str(e), parent=win)

    frame_bts = tk.Frame(win, bg=COLORS["fondo"])
    frame_bts.pack(pady=(0, 14))

    tk.Button(
        frame_bts,
        text="⚙️  COMPARAR AHORA",
        command=ejecutar_comparacion,
        bg="#16a34a",
        fg="white",
        font=("Segoe UI", 10, "bold"),
        padx=20,
        pady=7,
        cursor="hand2",
        relief="flat",
    ).pack(side="left", padx=8)

    tk.Button(
        frame_bts,
        text="💾 EXPORTAR A EXCEL",
        command=exportar_resultado,
        bg="#0078d4",
        fg="white",
        font=("Segoe UI", 10, "bold"),
        padx=20,
        pady=7,
        cursor="hand2",
        relief="flat",
    ).pack(side="left", padx=8)

    tk.Button(
        frame_bts,
        text="🗑️  Limpiar",
        command=lambda: [
            tree.delete(*tree.get_children()),
            lbl_ok.config(text="✅ Match: —"),
            lbl_falta.config(text="❌ Faltantes: —"),
            lbl_sobra.config(text="❓ Sobrantes: —"),
            lbl_diff.config(text="⚠️ Diferencias: —"),
        ],
        bg="#94a3b8",
        fg="white",
        font=("Segoe UI", 9),
        padx=12,
        pady=7,
        cursor="hand2",
        relief="flat",
    ).pack(side="left", padx=8)

def abrir_modulo_etiquetas_manual(ventana_padre):
    win = ventana_modal(
        ventana_padre, "Generador Manual de Etiquetas (Equipos Defectuosos)", "750x700"
    )
    titulo_ui(win, "🏷️ Generador Manual de Etiquetas", pady=(15, 5))
    tk.Label(
        win,
        text="Crea e imprime una etiqueta de diagnóstico sin alterar los reportes de inventario.",
        font=("Segoe UI", 9),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(0, 15))

    # --- DATOS DEL EQUIPO ---
    frame_datos = ttk.LabelFrame(win, text=" Datos del Equipo ")
    frame_datos.pack(fill="x", padx=25, pady=5)
    
    tk.Label(frame_datos, text="N° de Serie:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=0, column=0, padx=10, pady=8, sticky="e")
    ent_serial = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_serial.grid(row=0, column=1, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="Modelo:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=0, column=2, padx=10, pady=8, sticky="e")
    ent_modelo = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_modelo.grid(row=0, column=3, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="Procesador (CPU):", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=1, column=0, padx=10, pady=8, sticky="e")
    ent_cpu = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_cpu.grid(row=1, column=1, padx=5, pady=8, sticky="w")

    tk.Label(frame_datos, text="N° de Guía:", font=("Segoe UI", 9), bg=COLORS["fondo"]).grid(row=1, column=2, padx=10, pady=8, sticky="e")
    ent_guia = tk.Entry(frame_datos, width=25, font=("Segoe UI", 9))
    ent_guia.grid(row=1, column=3, padx=5, pady=8, sticky="w")

    # --- REVISIÓN DE COMPONENTES ---
    lf_fallas = ttk.LabelFrame(win, text=" Revisión de Componentes ")
    lf_fallas.pack(fill="both", expand=True, padx=25, pady=10)

    componentes_etiqueta = [
        "Carcaza", "Pantalla", "Teclado", "Touchpad", "Puertos USB", 
        "Puertos Video", "Ethernet/Wi-Fi", "Placa base", "Memoria", "Disco duro", "Batería"
    ]

    comps_data = []

    frame_grilla = tk.Frame(lf_fallas, bg=COLORS["fondo"])
    frame_grilla.pack(padx=10, pady=10)

    tk.Label(frame_grilla, text="Componente", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=0, sticky="w", padx=5)
    tk.Label(frame_grilla, text="Estado", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=1, padx=5)
    tk.Label(frame_grilla, text="Observación (Si es OBS/MALO)", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]).grid(row=0, column=2, sticky="w", padx=5)

    for i, comp_name in enumerate(componentes_etiqueta, start=1):
        tk.Label(frame_grilla, text=comp_name, bg=COLORS["fondo"], font=("Segoe UI", 9)).grid(row=i, column=0, sticky="w", padx=5, pady=2)
        
        cmb_estado = ttk.Combobox(frame_grilla, values=["OK", "OBS", "MALO"], width=6, state="readonly", font=("Segoe UI", 9))
        cmb_estado.set("OK")
        cmb_estado.grid(row=i, column=1, padx=5, pady=2)

        ent_obs = tk.Entry(frame_grilla, width=45, font=("Segoe UI", 9), state=tk.DISABLED, bg="#f1f5f9")
        ent_obs.grid(row=i, column=2, padx=5, pady=2, sticky="w")

        def toggle_entry(event, e=ent_obs, c=cmb_estado):
            if c.get() != "OK":
                e.config(state=tk.NORMAL, bg="#ffffff")
                e.focus()
            else:
                e.delete(0, tk.END)
                e.config(state=tk.DISABLED, bg="#f1f5f9")

        cmb_estado.bind("<<ComboboxSelected>>", toggle_entry)
        comps_data.append({"nombre": comp_name, "estado": cmb_estado, "obs": ent_obs})

    def generar_y_imprimir():
        serial = ent_serial.get().strip() or "S/N"
        modelo = ent_modelo.get().strip() or "Desconocido"
        cpu = ent_cpu.get().strip() or "Desconocido"
        guia = ent_guia.get().strip() or "S/N"
        fecha = datetime.now().strftime('%d/%m/%Y')

        trs = ""
        for c in comps_data:
            nombre = c["nombre"]
            estado = c["estado"].get()
            obs = c["obs"].get().strip()
            
            iconOK = "☑" if estado == "OK" else "□"
            iconOBS = "☑" if estado == "OBS" else "□"
            iconMalo = "☑" if estado == "MALO" else "□"
            
            trs += f"""<tr>
              <td>{nombre}</td>
              <td style="text-align: center;">{iconOK} OK &nbsp;&nbsp;&nbsp; {iconOBS} OBS &nbsp;&nbsp;&nbsp; {iconMalo} MALO</td>
              <td>{obs}</td>
            </tr>"""

        html_etiqueta = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Etiqueta_{serial}</title>
            <style>
              @media print {{ @page {{ margin: 15mm; size: auto; }} body {{ margin: 0; }} }}
              body {{ font-family: 'Arial', sans-serif; font-size: 11px; color: #000; max-width: 700px; margin: auto; padding: 20px; line-height: 1.4; }}
              .header-title {{ display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 20px; font-size: 13px; }}
              .row {{ margin-bottom: 10px; font-size: 11px;}}
              .row span {{ margin-right: 25px; }}
              table {{ width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 20px; }}
              th, td {{ border: 1px solid #000; padding: 5px 8px; text-align: left; }}
              .checkboxes {{ font-family: 'Segoe UI Symbol', sans-serif; }}
              .bold {{ font-weight: bold; }}
            </style>
        </head>
        <body onload="setTimeout(() => {{ window.print(); }}, 500);">
            <div class="header-title">
              <span>ETIQUETAS NOTEBOOK</span>
              <span>FECHA: {fecha}</span>
            </div>
            <div class="row">
              <span><span class="bold">MODELO:</span> {modelo}</span>
              <span><span class="bold">PROCESADOR:</span> {cpu}</span>
              <span><span class="bold">GUIA:</span> {guia}</span>
              <span><span class="bold">SERIAL:</span> {serial}</span>
            </div>
            <div class="row checkboxes" style="margin-bottom:20px;">
              <span class="bold">ESTADO GLOBAL:</span> &nbsp;&nbsp;&nbsp;&nbsp;
              □ BUENA &nbsp;&nbsp;&nbsp; ☑ POR REPARAR &nbsp;&nbsp;&nbsp; □ MALA &nbsp;&nbsp;&nbsp; □ REPUESTO
            </div>
            <table>
              <thead>
                <tr>
                  <th width="20%">REVISIÓN DE COMPONENTES</th>
                  <th width="35%" style="text-align: center;">ESTADO</th>
                  <th width="45%">OBSERVACIÓN</th>
                </tr>
              </thead>
              <tbody class="checkboxes">
                {trs}
              </tbody>
            </table>
            <div class="row checkboxes bold" style="margin-top:20px;">
              CHECKLIST: &nbsp;&nbsp;&nbsp;&nbsp; □ PRUEBA S/O &nbsp;&nbsp;&nbsp;&nbsp; □ HYDRA &nbsp;&nbsp;&nbsp;&nbsp; □ HARDWARE DEFAULT
            </div>
        </body>
        </html>
        """

        temp_path = os.path.join(tempfile.gettempdir(), "etiqueta_manual_temp.html")
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write(html_etiqueta)

        webbrowser.open(f"file://{temp_path}")

    tk.Button(
        win,
        text="🖨️ GENERAR E IMPRIMIR ETIQUETA",
        command=generar_y_imprimir,
        bg="#0078d4",
        fg="white",
        font=("Segoe UI", 10, "bold"),
        padx=20,
        pady=8,
        cursor="hand2",
        relief="flat"
    ).pack(pady=10)

def abrir_modulo_cambiar_nombre(ventana_padre):
    # 1. Obtener el número de serie de la BIOS
    try:
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = 0
        
        cmd = ["powershell", "-NoProfile", "-Command", "(Get-CimInstance Win32_Bios).SerialNumber"]
        serial = subprocess.check_output(cmd, text=True, startupinfo=si).strip()
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo obtener el número de serie de la BIOS.\n{e}", parent=ventana_padre)
        return

    if not serial or serial.lower() in ("to be filled by o.e.m.", "default string"):
        messagebox.showwarning("Aviso", "No se detectó un número de serie válido en la BIOS de este equipo.", parent=ventana_padre)
        return

    # 2. Limpiar y formatear el nombre (Máximo 15 caracteres para NetBIOS en Windows)
    serial_limpio = re.sub(r"[^a-zA-Z0-9\-]", "", serial)
    nuevo_nombre = f"ARR-{serial_limpio}"
    
    if len(nuevo_nombre) > 15:
        nuevo_nombre = nuevo_nombre[:15] # Truncar a 15 caracteres si es muy largo
        
    respuesta = messagebox.askyesno(
        "Confirmar Cambio de Nombre", 
        f"El número de serie detectado es: {serial}\n\n"
        f"¿Estás seguro de cambiar el nombre de este equipo a:\n\n{nuevo_nombre}?\n\n"
        "(Debes aceptar la ventana de permisos de Administrador que aparecerá a continuación).",
        parent=ventana_padre
    )
    
    if not respuesta:
        return
        
    # 3. Aplicar el cambio de nombre
    try:
        ps_command = f'Rename-Computer -NewName "{nuevo_nombre}"'
        cmd = [
            "powershell",
            "-NoProfile",
            "-Command",
            f"Start-Process powershell -ArgumentList '-NoProfile -WindowStyle Hidden -Command {ps_command}' -Verb RunAs"
        ]
        
        subprocess.run(cmd, check=True)
        
        messagebox.showinfo(
            "✅ Solicitud Enviada", 
            f"Se ha ordenado el cambio de nombre a '{nuevo_nombre}'.\n\nPor favor, REINICIA EL EQUIPO manualmente para que los cambios surtan efecto.", 
            parent=ventana_padre
        )
        
    except subprocess.CalledProcessError:
        messagebox.showerror(
            "Error", 
            "No se pudo cambiar el nombre del equipo.\n\nAsegúrate de aceptar la ventana de permisos de Administrador.", 
            parent=ventana_padre
        )
    except Exception as e:
        messagebox.showerror(
            "Error Inesperado", 
            f"Ocurrió un error al intentar cambiar el nombre: {e}", 
            parent=ventana_padre
        )








# ─────────────────────────────────────────────────────────
#  Navegación e Interfaz Principal
# ─────────────────────────────────────────────────────────
# ============================================================================
# 10. UI PRINCIPAL — navegación en la ventana raíz: menú, pantalla de escaneo
#     y formulario de registro.
# ============================================================================
def limpiar_ventana(ventana):
    for widget in ventana.winfo_children():
        widget.destroy()


def mostrar_menu_principal(ventana):
    limpiar_ventana(ventana)
    ventana.geometry("520x650")

    titulo_ui(
        ventana, "🖥️ Sistema de Control de Inventario", size=16, pady=(30, 20)
    )
    frame_btns = tk.Frame(ventana, bg=COLORS["fondo"])
    frame_btns.pack(fill="both", expand=True, padx=40)

    boton_menu(
        frame_btns,
        "🔍 ESCANEAR ESTE EQUIPO\n(Crear registro)",
        lambda: iniciar_escaneo(ventana),
        COLORS["azul_btn"],
        size=12,
    )
    boton_menu(
        frame_btns,
        "📦 EMPACAR LOTE PENDIENTE",
        lambda: accion_finalizar_lote(ventana),
        COLORS["verde"],
    )
    boton_menu(
        frame_btns,
        "🚀 ENVIAR CARPETAS A RED",
        lambda: accion_enviar_red(ventana),
        COLORS["naranja"],
    )
    boton_menu(
        frame_btns,
        "🔧 RECONSTRUIR HTML DESDE JSONs",
        lambda: accion_unir_Json(ventana),
        COLORS["morado"],
    )

    ttk.Separator(frame_btns, orient="horizontal").pack(fill="x", pady=15)
    tk.Label(
        frame_btns,
        text="Herramientas de Auditoría",
        font=fuente(10, True),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack(pady=(0, 5))
    boton_menu(
        frame_btns,
        "📑 COMPARAR EXCEL (Entregas vs Retiros)",
        lambda: abrir_modulo_comparacion_excel(ventana),
        COLORS["celeste"],
        size=9,
        bold="normal",
    )

    # --- BOTÓN PARA CAMBIAR NOMBRE DEL EQUIPO ---
    boton_menu(
        frame_btns,
        "💻 CAMBIAR NOMBRE DEL EQUIPO",
        lambda: abrir_modulo_cambiar_nombre(ventana),
        COLORS["verde_esmeralda"],  # Color esmeralda para diferenciarlo
        size=9,
        bold="bold",
    )

    # --- BOTÓN DE ETIQUETAS MANUALES ---
    boton_menu(
        frame_btns,
        "🏷️ GENERAR ETIQUETA MANUAL (Equipos Malos)",
        lambda: abrir_modulo_etiquetas_manual(ventana),
        COLORS["pizarra"],
        size=9,
        bold="bold",
    )


def iniciar_escaneo(ventana):
    limpiar_ventana(ventana)
    ventana.geometry("750x850")

    frame_carga = tk.Frame(ventana, bg=COLORS["fondo"])
    frame_carga.place(relx=0.5, rely=0.5, anchor="center")

    tk.Label(frame_carga, text="⏳", font=fuente(48), bg=COLORS["fondo"]).pack()
    lbl_scan = tk.Label(
        frame_carga,
        text="Escaneando hardware...",
        font=fuente(15, True),
        bg=COLORS["fondo"],
        fg=COLORS["azul"],
    )
    lbl_scan.pack(pady=6)
    tk.Label(
        frame_carga,
        text="Esto tomará unos segundos.",
        font=fuente(10),
        bg=COLORS["fondo"],
        fg=COLORS["gris"],
    ).pack()

    dots = ["", ".", "..", "..."]
    dot_idx = [0]

    def animar():
        if frame_carga.winfo_exists():
            lbl_scan.config(text=f"Escaneando hardware{dots[dot_idx[0] % 4]}")
            dot_idx[0] += 1
            ventana.after(450, animar)

    animar()

    resultado = [None]

    def escanear():
        resultado[0] = get_inventory_fast()
        ventana.after(0, lambda: construir_ui_formulario(ventana, resultado[0]))

    threading.Thread(target=escanear, daemon=True).start()


def construir_ui_formulario(ventana, data):
    limpiar_ventana(ventana)

    if not data:
        messagebox.showerror("Error", "No se pudo obtener información del hardware.")
        mostrar_menu_principal(ventana)
        return

    top_frame = tk.Frame(ventana, bg=COLORS["fondo"])
    top_frame.pack(fill="x", pady=(14, 4), padx=30)

    tk.Button(
        top_frame,
        text="⬅ Volver",
        command=lambda: mostrar_menu_principal(ventana),
        bg="#cbd5e1",
        fg=COLORS["gris_oscuro"],
        font=("Segoe UI", 9, "bold"),
        cursor="hand2",
        relief="flat",
        padx=10,
    ).pack(side="left")
    tk.Label(
        top_frame,
        text="💻 Asistente de Entrega / Retiro",
        font=("Segoe UI", 15, "bold"),
        bg=COLORS["fondo"],
        fg=COLORS["azul"],
    ).pack(side="left", padx=50)

    info_frame = tk.Frame(ventana, bg="#dbeafe", bd=1, relief="solid")
    info_frame.pack(padx=30, fill="x")
    tk.Label(
        info_frame,
        text=f"  {data['model']}   |   S/N: {data['serial']}  ",
        font=("Consolas", 10),
        bg="#dbeafe",
        fg="#1d4ed8",
        pady=6,
    ).pack()

    marco = tk.Frame(ventana, bg=COLORS["fondo"])
    marco.pack(pady=10, fill="x", padx=30)

    lf_log = ttk.LabelFrame(marco, text=" 🚚 Datos Logísticos ")
    lf_log.pack(fill="x", pady=(0, 10))
    frame_log_int = tk.Frame(lf_log, bg=COLORS["fondo"])
    frame_log_int.pack(pady=5)

    tk.Label(
        frame_log_int, text="Movimiento:", bg=COLORS["fondo"], font=("Segoe UI", 9)
    ).grid(row=0, column=0, padx=8, pady=6, sticky="e")
    combo_mov = ttk.Combobox(
        frame_log_int,
        values=["Entrega", "Retiro"],
        width=13,
        state="readonly",
        font=("Segoe UI", 9),
    )
    combo_mov.set("Entrega")
    combo_mov.grid(row=0, column=1, padx=6, pady=6, sticky="w")

    tk.Label(
        frame_log_int, text="N° de Guía:", bg=COLORS["fondo"], font=("Segoe UI", 9)
    ).grid(row=0, column=2, padx=8, sticky="e")
    entry_guia = tk.Entry(
        frame_log_int, width=13, font=("Segoe UI", 9), state=tk.DISABLED, fg="#aaa"
    )
    entry_guia.grid(row=0, column=3, padx=6, pady=6, sticky="w")

    lf_fallas = ttk.LabelFrame(marco, text=" ⚠️ Revisión de Componentes (Solo Retiros) ")

    componentes_etiqueta = [
        "Carcaza",
        "Pantalla",
        "Teclado",
        "Touchpad",
        "Puertos USB",
        "Puertos Video",
        "Ethernet/Wi-Fi",
        "Placa base",
        "Memoria",
        "Disco duro",
        "Batería",
    ]

    comps_data = []

    frame_grilla = tk.Frame(lf_fallas, bg=COLORS["fondo"])
    frame_grilla.pack(padx=10, pady=10)

    tk.Label(
        frame_grilla, text="Componente", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]
    ).grid(row=0, column=0, sticky="w", padx=5)
    tk.Label(
        frame_grilla, text="Estado", font=("Segoe UI", 9, "bold"), bg=COLORS["fondo"]
    ).grid(row=0, column=1, padx=5)
    tk.Label(
        frame_grilla,
        text="Observación (Obligatorio si es OBS/MALO)",
        font=("Segoe UI", 9, "bold"),
        bg=COLORS["fondo"],
    ).grid(row=0, column=2, sticky="w", padx=5)

    for i, comp_name in enumerate(componentes_etiqueta, start=1):
        tk.Label(frame_grilla, text=comp_name, bg=COLORS["fondo"], font=("Segoe UI", 9)).grid(
            row=i, column=0, sticky="w", padx=5, pady=2
        )

        cmb_estado = ttk.Combobox(
            frame_grilla,
            values=["OK", "OBS", "MALO"],
            width=6,
            state="readonly",
            font=("Segoe UI", 9),
        )
        cmb_estado.set("OK")
        cmb_estado.grid(row=i, column=1, padx=5, pady=2)

        ent_obs = tk.Entry(
            frame_grilla,
            width=40,
            font=("Segoe UI", 9),
            state=tk.DISABLED,
            bg="#f1f5f9",
        )
        ent_obs.grid(row=i, column=2, padx=5, pady=2, sticky="w")

        def toggle_entry(event, e=ent_obs, c=cmb_estado):
            if c.get() != "OK":
                e.config(state=tk.NORMAL, bg="#ffffff")
                e.focus()
            else:
                e.delete(0, tk.END)
                e.config(state=tk.DISABLED, bg="#f1f5f9")

        cmb_estado.bind("<<ComboboxSelected>>", toggle_entry)

        comps_data.append({"nombre": comp_name, "estado": cmb_estado, "obs": ent_obs})

    def toggle_guia(event=None):
        if combo_mov.get() == "Retiro":
            entry_guia.config(state=tk.NORMAL, fg="#000")
            lf_fallas.pack(fill="x", pady=(0, 10))
            guia_guardada = cargar_guia_config()
            if guia_guardada:
                entry_guia.delete(0, tk.END)
                entry_guia.insert(0, guia_guardada)
        else:
            entry_guia.delete(0, tk.END)
            entry_guia.config(state=tk.DISABLED, fg="#aaa")
            lf_fallas.pack_forget()
            for c in comps_data:
                c["estado"].set("OK")
                c["obs"].config(state=tk.NORMAL)
                c["obs"].delete(0, tk.END)
                c["obs"].config(state=tk.DISABLED, bg="#f1f5f9")

    combo_mov.bind("<<ComboboxSelected>>", toggle_guia)

    tk.Label(
        marco,
        text="Observación General para el Acta:",
        bg=COLORS["fondo"],
        font=("Segoe UI", 10, "bold"),
        fg=COLORS["gris_oscuro"],
    ).pack(anchor="w", pady=(4, 2))
    entry_obs_gen = tk.Entry(marco, width=50, font=("Segoe UI", 10))
    entry_obs_gen.pack(fill="x", ipady=4, pady=(0, 8))

    frame_off = tk.Frame(marco, bg=COLORS["fondo"])
    frame_off.pack(fill="x", pady=(4, 0))
    var_office = tk.IntVar()

    def toggle_office():
        state = tk.NORMAL if var_office.get() else tk.DISABLED
        entry_office.config(state=state)
        combo_ver.config(state="readonly" if var_office.get() else tk.DISABLED)
        if not var_office.get():
            entry_office.delete(0, tk.END)

    tk.Checkbutton(
        frame_off,
        text="📦 Registrar Office",
        variable=var_office,
        command=toggle_office,
        bg=COLORS["fondo"],
        font=("Segoe UI", 10, "bold"),
        fg="#1e40af",
        activebackground="#eef2f7",
        cursor="hand2",
    ).pack(side="left")
    combo_ver = ttk.Combobox(
        frame_off,
        values=["2013", "2016", "2019", "2021", "2024", "365"],
        width=8,
        state=tk.DISABLED,
        font=("Segoe UI", 10),
    )
    combo_ver.set("2016")
    combo_ver.pack(side="left", padx=6)

    entry_office = tk.Entry(
        marco, font=("Consolas", 10), justify="center", state=tk.DISABLED, fg="#555"
    )
    entry_office.pack(fill="x", ipady=4, pady=6)

    frame_bts = tk.Frame(ventana, bg=COLORS["fondo"])
    frame_bts.pack(pady=10)

    tk.Button(
        frame_bts,
        text="💾 GUARDAR EQUIPO EN LOTE",
        command=lambda: accion_agregar_lote(
            ventana,
            data,
            entry_obs_gen.get().strip(),
            var_office.get() == 1,
            entry_office.get().strip(),
            combo_ver.get(),
            combo_mov.get(),
            entry_guia.get().strip(),
            comps_data,
        ),
        bg="#0078d4",
        fg="white",
        font=("Segoe UI", 12, "bold"),
        padx=20,
        pady=10,
        cursor="hand2",
        relief="flat",
    ).pack()


# ============================================================================
# 11. ENTRYPOINT
# ============================================================================
def iniciar_interfaz_principal():
    ventana = tk.Tk()
    ventana.title("Asistente de Entrega — Arrienda.cl")
    ventana.resizable(False, False)
    ventana.attributes("-topmost", True)
    ventana.configure(bg=COLORS["fondo"])

    style = ttk.Style()
    style.theme_use("clam")
    style.configure(
        "TCombobox", fieldbackground="#fff", background="#fff", font=("Segoe UI", 9)
    )
    style.configure("TEntry", fieldbackground="#fff", font=("Segoe UI", 9))
    style.configure("TLabelframe", background="#eef2f7")
    style.configure(
        "TLabelframe.Label",
        background="#eef2f7",
        font=("Segoe UI", 10, "bold"),
        foreground="#334155",
    )

    mostrar_menu_principal(ventana)
    ventana.mainloop()


if __name__ == "__main__":
    iniciar_interfaz_principal()
